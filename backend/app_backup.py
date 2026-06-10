import os
import re 
import fitz
import docx as docxlib
import json
import sqlite3
import requests
from dotenv import load_dotenv
load_dotenv()
from datetime import datetime
from io import BytesIO

# ── top of file ──────────────────────────────
import os
import uuid
from flask import Flask, render_template, request, jsonify, redirect, url_for
from werkzeug.utils import secure_filename
from werkzeug.exceptions import RequestEntityTooLarge
from PIL import Image

app = Flask(__name__)

# ── config (right after app = Flask) ─────────
UPLOAD_FOLDER   = 'static/uploads/profile_photos'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'webp'}
MAX_FILE_SIZE   = 2 * 1024 * 1024
MAX_DIMENSION   = (400, 400)

app.config['UPLOAD_FOLDER']       = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH']  = MAX_FILE_SIZE

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def resize_image(image_path, max_size=MAX_DIMENSION):
    with Image.open(image_path) as img:
        if img.mode in ('RGBA', 'P'):
            img = img.convert('RGB')
        img.thumbnail(max_size, Image.LANCZOS)
        img.save(image_path, optimize=True, quality=85)

def delete_old_photo(old_path):
    if old_path:
        local_path = old_path.lstrip('/')
        if os.path.exists(local_path):
            os.remove(local_path)
    
from flask import (
    Flask,
    render_template,
    render_template_string,
    request,
    redirect,
    session,
    jsonify,
    send_file,
)
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
try:
    from pdfminer.high_level import extract_text as pdf_extract_text
    PDF_EXTRACT_OK = True
except ImportError:
    PDF_EXTRACT_OK = False

app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", "dev-secret-key-change-in-production")
app.jinja_env.filters['fromjson'] = lambda s: json.loads(s) if s else {}

GROQ_KEY = os.environ.get("GROQ_KEY", "")
DB_PATH = "screening.db"

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_tables():
    conn = get_db()
    conn.execute("""
    CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT UNIQUE,
        password TEXT,
        role TEXT DEFAULT 'jobseeker',
        full_name TEXT,
        email TEXT,
        phone TEXT,
        city TEXT,
        state TEXT,
        job_title TEXT,
        experience TEXT,
        education TEXT,
        company_name TEXT,
        company_size TEXT,
        industry TEXT,
        designation TEXT,
        created_at TEXT
    )""")
    conn.execute("""
    CREATE TABLE IF NOT EXISTS resume_drafts (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT,
        draft_json TEXT,
        updated_at TEXT
    )""")
    conn.execute("""
    CREATE TABLE IF NOT EXISTS resume_analytics (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT UNIQUE,
        views INTEGER DEFAULT 0,
        downloads INTEGER DEFAULT 0,
        ats_score INTEGER DEFAULT 0
    )""")
    conn.execute("""
    CREATE TABLE IF NOT EXISTS resume_versions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT,
        version_name TEXT,
        resume_json TEXT,
        created_at TEXT
    )""")
    conn.execute("""
    CREATE TABLE IF NOT EXISTS jobs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        recruiter TEXT,
        title TEXT,
        company TEXT,
        location TEXT,
        job_type TEXT DEFAULT 'Full-time',
        description TEXT,
        skills_required TEXT,
        salary TEXT,
        posted_at TEXT,
        active INTEGER DEFAULT 1
    )""")
    conn.execute("""
    CREATE TABLE IF NOT EXISTS applications (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT,
        job_id INTEGER,
        job_title TEXT,
        company TEXT,
        status TEXT DEFAULT 'Applied',
        applied_at TEXT,
        resume_json TEXT
    )""")
    conn.execute("""
    CREATE TABLE IF NOT EXISTS screening_history (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        recruiter TEXT,
        candidate_name TEXT,
        filename TEXT,
        ats_score INTEGER DEFAULT 0,
        result_json TEXT,
        screened_at TEXT
    )""")
    conn.execute("""
    CREATE TABLE IF NOT EXISTS notifications (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user TEXT,
        type TEXT,
        message TEXT,
        link TEXT,
        is_read INTEGER DEFAULT 0,
        created_at TEXT
    )""")
    conn.commit()
    conn.close()

def push_notification(user, msg_type, message, link="/"):
    """Insert a notification for a user."""
    try:
        conn = get_db()
        conn.execute(
            "INSERT INTO notifications (user, type, message, link, is_read, created_at) VALUES (?,?,?,?,0,?)",
            (user, msg_type, message, link, datetime.now().isoformat())
        )
        conn.commit()
        conn.close()
    except Exception:
        pass


def gemini(prompt):
    if not GROQ_KEY:
        return "AI unavailable (no GROQ_KEY set)"
    try:
        from groq import Groq
        client = Groq(api_key=GROQ_KEY)
        chat_completion = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama-3.3-70b-versatile",
        )
        return chat_completion.choices[0].message.content
    except Exception as e:
        return f"AI error: {str(e)}"
    
@app.route('/')
def index():
    if 'user' not in session:
        return render_template('public/landing_page_1.html')  # show landing page
    return redirect(url_for('candidate_dashboard'))  # logged in → go to dashboard

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "").strip()
        role = request.form.get("role", "jobseeker")
        conn = get_db()
        user = conn.execute(
            "SELECT * FROM users WHERE username=? AND password=? AND role=?",
            (username, password, role)
        ).fetchone()
        conn.close()
        if user:
            session["user"] = username
            session["role"] = user["role"]
            return redirect("/recruiter/dashboard" if user["role"] == "recruiter" else "/dashboard")
        return render_template_string(LOGIN_HTML, error="Invalid credentials or wrong role selected")
    return render_template_string(LOGIN_HTML, error=None)

@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "").strip()
        role = request.form.get("role", "jobseeker")
        if not username or not password:
            return render_template_string(REGISTER_HTML, error="All fields required")
        conn = get_db()
        if conn.execute("SELECT id FROM users WHERE username=?", (username,)).fetchone():
            conn.close()
            return render_template_string(REGISTER_HTML, error="Username already taken")
        conn.execute("""
            INSERT INTO users (username, password, role, full_name, email, phone,
            city, state, job_title, experience, education,
            company_name, company_size, industry, designation, created_at)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
        """, (
            username, password, role,
            request.form.get("full_name"),
            request.form.get("email"),
            request.form.get("phone"),
            request.form.get("city"),
            request.form.get("state"),
            request.form.get("job_title"),
            request.form.get("experience"),
            request.form.get("education"),
            request.form.get("company_name"),
            request.form.get("company_size"),
            request.form.get("industry"),
            request.form.get("designation"),
            datetime.now().isoformat()
        ))
        conn.commit()
        conn.close()
        session["user"] = username
        session["role"] = role
        return redirect("/recruiter/dashboard" if role == "recruiter" else "/dashboard")
    return render_template_string(REGISTER_HTML, error=None)

@app.route("/logout")
def logout():
    session.clear()
    return redirect("/login")

# CANDIDATE DASHBOARD
@app.route("/dashboard")
def candidate_dashboard():
    if "user" not in session: return redirect("/login")
    if session.get("role") == "recruiter": return redirect("/recruiter/dashboard")
    conn = get_db()
    analytics = conn.execute("SELECT * FROM resume_analytics WHERE username=?", (session["user"],)).fetchone()
    applications = conn.execute("SELECT * FROM applications WHERE username=? ORDER BY id DESC LIMIT 5", (session["user"],)).fetchall()
    total_apps = conn.execute("SELECT COUNT(*) as c FROM applications WHERE username=?", (session["user"],)).fetchone()["c"]
    shortlisted = conn.execute("SELECT COUNT(*) as c FROM applications WHERE username=? AND status='Shortlisted'", (session["user"],)).fetchone()["c"]
    conn.close()
    stats = dict(analytics) if analytics else {"views": 0, "downloads": 0, "ats_score": 0}
    return render_template_string(CANDIDATE_DASHBOARD_HTML,
        username=session["user"],
        stats=stats,
        applications=applications,
        total_apps=total_apps,
        shortlisted=shortlisted
    )

# RESUME BUILDER
@app.route("/resume-builder")
def resume_builder():
    if "user" not in session: return redirect("/login")
    if session.get("role") == "recruiter": return redirect("/recruiter/dashboard")
    return render_template("jobseeker/resume-builder.html")

@app.route('/upload-photo', methods=['POST'])
def upload_photo():
    if 'photo' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['photo']

    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    if not allowed_file(file.filename):
        return jsonify({'error': 'Invalid file type. Use JPG, PNG, or WEBP'}), 400

    ext = file.filename.rsplit('.', 1)[1].lower()
    unique_filename = f"{uuid.uuid4().hex}.{ext}"

    os.makedirs('static/uploads/profile_photos', exist_ok=True)
    save_path = os.path.join('static/uploads/profile_photos', unique_filename)

    old_photo = request.form.get('old_photo', '')
    delete_old_photo(old_photo)

    file.save(save_path)

    try:
        resize_image(save_path)
    except Exception as e:
        os.remove(save_path)
        return jsonify({'error': f'Image processing failed: {str(e)}'}), 500

    photo_url = f'/static/uploads/profile_photos/{unique_filename}'
    return jsonify({'success': True, 'path': photo_url})


@app.errorhandler(RequestEntityTooLarge)
def file_too_large(e):
    return jsonify({'error': 'File too large. Max size is 2 MB'}), 413

@app.route("/recruiter/dashboard")
def recruiter_dashboard():
    if "user" not in session: return redirect("/login")
    if session.get("role") != "recruiter": return redirect("/dashboard")
    conn = get_db()
    jobs = conn.execute("SELECT * FROM jobs WHERE recruiter=? ORDER BY id DESC", (session["user"],)).fetchall()
    total_apps = conn.execute(
        "SELECT COUNT(*) as c FROM applications WHERE job_id IN (SELECT id FROM jobs WHERE recruiter=?)",
        (session["user"],)).fetchone()["c"]
    pending = conn.execute(
        "SELECT COUNT(*) as c FROM applications WHERE status='Applied' AND job_id IN (SELECT id FROM jobs WHERE recruiter=?)",
        (session["user"],)).fetchone()["c"]
    shortlisted = conn.execute(
        "SELECT COUNT(*) as c FROM applications WHERE status='Shortlisted' AND job_id IN (SELECT id FROM jobs WHERE recruiter=?)",
        (session["user"],)).fetchone()["c"]
    hired = conn.execute(
        "SELECT COUNT(*) as c FROM applications WHERE status='Hired' AND job_id IN (SELECT id FROM jobs WHERE recruiter=?)",
        (session["user"],)).fetchone()["c"]
    rejected = conn.execute(
        "SELECT COUNT(*) as c FROM applications WHERE status='Rejected' AND job_id IN (SELECT id FROM jobs WHERE recruiter=?)",
        (session["user"],)).fetchone()["c"]
    recent_apps = conn.execute(
        """SELECT a.*, j.title as job_title FROM applications a
           JOIN jobs j ON a.job_id=j.id
           WHERE j.recruiter=? ORDER BY a.id DESC LIMIT 5""",
        (session["user"],)).fetchall()
    # top applicants = most recent shortlisted
    top_apps = conn.execute(
        """SELECT a.*, j.title as job_title FROM applications a
           JOIN jobs j ON a.job_id=j.id
           WHERE j.recruiter=? AND a.status IN ('Shortlisted','Hired')
           ORDER BY a.id DESC LIMIT 5""",
        (session["user"],)).fetchall()
    conn.close()
    return render_template_string(RECRUITER_DASHBOARD_HTML,
        username=session["user"], jobs=jobs, total_jobs=len(jobs),
        total_apps=total_apps, pending=pending, shortlisted=shortlisted,
        hired=hired, rejected=rejected, recent_apps=recent_apps, top_apps=top_apps,
        now=datetime.now().strftime("%A, %d %B %Y"))

# POST JOB
@app.route("/recruiter/post-job", methods=["GET", "POST"])
def post_job():
    if "user" not in session or session.get("role") != "recruiter":
        return redirect("/login")
    if request.method == "POST":
        conn = get_db()
        conn.execute("""
            INSERT INTO jobs (recruiter, title, company, location, job_type, description, skills_required, salary, posted_at)
            VALUES (?,?,?,?,?,?,?,?,?)
        """, (session["user"], request.form.get("title"), request.form.get("company"),
              request.form.get("location"), request.form.get("job_type","Full-time"),
              request.form.get("description"), request.form.get("skills_required"),
              request.form.get("salary"), datetime.now().isoformat()))
        conn.commit(); conn.close()
        return redirect("/recruiter/dashboard")
    return render_template_string(POST_JOB_HTML)

# APPLICATIONS PAGE (all jobs grouped)
@app.route("/recruiter/applications")
def recruiter_applications():
    if "user" not in session or session.get("role") != "recruiter":
        return redirect("/login")
    conn = get_db()
    jobs = conn.execute("SELECT * FROM jobs WHERE recruiter=? ORDER BY id DESC", (session["user"],)).fetchall()
    all_apps = []
    total = pending = shortlisted = rejected = 0
    for job in jobs:
        apps = conn.execute("SELECT * FROM applications WHERE job_id=? ORDER BY id DESC", (job["id"],)).fetchall()
        if apps:
            app_list = [dict(a) for a in apps]
            all_apps.append((dict(job), app_list))
            for a in app_list:
                total += 1
                if a["status"] == "Applied": pending += 1
                elif a["status"] == "Shortlisted": shortlisted += 1
                elif a["status"] == "Rejected": rejected += 1
    conn.close()
    return render_template_string(RECRUITER_APPLICATIONS_HTML,
        username=session["user"], all_apps=all_apps,
        total=total, pending=pending, shortlisted=shortlisted, rejected=rejected)

# VIEW APPLICANTS (single job)
@app.route("/recruiter/job/<int:job_id>/applicants")
def view_applicants(job_id):
    if "user" not in session or session.get("role") != "recruiter":
        return redirect("/login")
    conn = get_db()
    job = conn.execute("SELECT * FROM jobs WHERE id=? AND recruiter=?", (job_id, session["user"])).fetchone()
    if not job: conn.close(); return "Job not found", 404
    applicants = conn.execute("SELECT * FROM applications WHERE job_id=? ORDER BY id DESC", (job_id,)).fetchall()
    conn.close()
    return render_template_string(APPLICANTS_HTML, job=job, applicants=applicants)

# CANDIDATE DETAIL
@app.route("/recruiter/candidate/<int:app_id>")
def candidate_detail(app_id):
    if "user" not in session or session.get("role") != "recruiter":
        return redirect("/login")
    conn = get_db()
    app_row = conn.execute(
        """SELECT a.*, j.title as job_title, j.skills_required, j.description as job_desc
           FROM applications a JOIN jobs j ON a.job_id=j.id
           WHERE a.id=? AND j.recruiter=?""", (app_id, session["user"])).fetchone()
    if not app_row: conn.close(); return "Not found", 404
    # get user profile info
    user_info = conn.execute("SELECT * FROM users WHERE username=?", (app_row["username"],)).fetchone()
    conn.close()
    resume_data = {}
    try: resume_data = json.loads(app_row["resume_json"] or "{}")
    except: pass
    return render_template_string(CANDIDATE_DETAIL_HTML,
        username=session["user"], app=dict(app_row),
        user_info=dict(user_info) if user_info else {},
        resume=resume_data)

# ALL CANDIDATES
@app.route("/recruiter/candidates")
def recruiter_candidates():
    if "user" not in session or session.get("role") != "recruiter":
        return redirect("/login")
    conn = get_db()
    candidates = conn.execute("""
        SELECT a.*, j.title as job_title, u.email, u.full_name, u.phone
        FROM applications a
        JOIN jobs j ON a.job_id=j.id
        LEFT JOIN users u ON a.username=u.username
        WHERE j.recruiter=? ORDER BY a.id DESC
    """, (session["user"],)).fetchall()
    conn.close()
    return render_template_string(CANDIDATES_LIST_HTML, candidates=candidates, username=session["user"])

# UPDATE STATUS
@app.route("/api/recruiter/update-status", methods=["POST"])
def update_status():
    if "user" not in session or session.get("role") != "recruiter":
        return jsonify({"error": "Unauthorized"}), 401
    data = request.get_json()
    conn = get_db()
    conn.execute("UPDATE applications SET status=? WHERE id=?", (data["status"], data["app_id"]))
    # notify the candidate
    app_row = conn.execute("SELECT username, job_title FROM applications WHERE id=?", (data["app_id"],)).fetchone()
    if app_row:
        push_notification(
            app_row["username"], "status",
            f"Your application for '{app_row['job_title']}' was updated to {data['status']}",
            "/analytics"
        )
    conn.commit(); conn.close()
    return jsonify({"success": True})

# NOTIFICATIONS API
@app.route("/api/notifications")
def get_notifications():
    if "user" not in session: return jsonify([])
    conn = get_db()
    rows = conn.execute(
        "SELECT * FROM notifications WHERE user=? ORDER BY id DESC LIMIT 20",
        (session["user"],)
    ).fetchall()
    unread = conn.execute(
        "SELECT COUNT(*) as c FROM notifications WHERE user=? AND is_read=0",
        (session["user"],)
    ).fetchone()["c"]
    conn.close()
    return jsonify({"notifications": [dict(r) for r in rows], "unread": unread})

@app.route("/api/notifications/mark-read", methods=["POST"])
def mark_notifications_read():
    if "user" not in session: return jsonify({"error": "Unauthorized"}), 401
    conn = get_db()
    conn.execute("UPDATE notifications SET is_read=1 WHERE user=?", (session["user"],))
    conn.commit(); conn.close()
    return jsonify({"success": True})

# AI INTERVIEW QUESTIONS
@app.route("/api/recruiter/interview-questions", methods=["POST"])
def interview_questions():
    if "user" not in session or session.get("role") != "recruiter":
        return jsonify({"error": "Unauthorized"}), 401
    data = request.get_json()
    resume = data.get("resume", {})
    job_title = data.get("job_title", "Software Engineer")
    skills = resume.get("skills", [])
    prompt = f"""You are a senior technical interviewer. Generate 8 targeted interview questions for a candidate applying for: {job_title}.
Candidate skills: {', '.join(skills) if skills else 'General software development'}
Mix of: 2 behavioural, 3 technical, 2 situational, 1 culture-fit.
Return ONLY a JSON array of objects, no extra text:
[{{"type":"Technical","question":"Explain how you would..."}}, ...]"""
    result = gemini(prompt)
    try:
        result = result.strip()
        if "```" in result: result = result.split("```")[1].replace("json","").strip()
        return jsonify({"questions": json.loads(result)})
    except:
        return jsonify({"questions": [
            {"type":"Technical","question":"Walk me through your most complex project."},
            {"type":"Behavioural","question":"Describe a time you handled a difficult deadline."},
            {"type":"Technical","question":"How do you approach debugging a production issue?"},
        ]})

# RECRUITER — SCREEN RESUME
@app.route("/recruiter/screen-resume", methods=["GET","POST"])
def recruiter_screen_resume():
    if "user" not in session or session.get("role") != "recruiter":
        return redirect("/login")
    result = None
    if request.method == "POST":
        file = request.files.get("resume_pdf")
        job_desc = request.form.get("job_description","")
        candidate_name = request.form.get("candidate_name","Unknown")
        if file:
            try:
                pdf_bytes = file.read()
                if PDF_EXTRACT_OK:
                    from io import BytesIO as _BIO
                    resume_text = pdf_extract_text(_BIO(pdf_bytes))[:4000]
                else:
                    resume_text = "(PDF extraction unavailable)"
            except Exception as e:
                resume_text = f"(Error reading PDF: {e})"
            prompt = f"""You are a senior ATS expert and recruiter. Analyse this resume.
RESUME TEXT: {resume_text}
JOB DESCRIPTION: {job_desc if job_desc else "General software/tech role"}
Respond ONLY with valid JSON (no markdown):
{{"ats_score":74,"technical_score":68,"overall_fit":71,"grade":"B","summary":"2-3 sentence assessment.",
"strengths":["strength1","strength2"],"improvements":["improve1","improve2"],
"keyword_match":58,"skills_found":["Python","SQL"],"skills_missing":["Docker"],
"sections_found":["Experience","Education","Skills"],"sections_missing":["Summary"],
"recommendation":"Hire/Maybe/Pass"}}"""
            result_raw = gemini(prompt)
            try:
                result_raw = result_raw.strip()
                if "```" in result_raw:
                    result_raw = result_raw.split("```")[1].replace("json","").strip()
                result = json.loads(result_raw)
                result["candidate_name"] = candidate_name
                result["filename"] = file.filename
                # store in screening history
                conn = get_db()
                conn.execute("""
                    INSERT INTO screening_history (recruiter, candidate_name, filename, ats_score, result_json, screened_at)
                    VALUES (?,?,?,?,?,?)
                """, (session["user"], candidate_name, file.filename,
                      result.get("ats_score",0), json.dumps(result), datetime.now().isoformat()))
                conn.commit(); conn.close()
            except Exception as e:
                result = {"error": f"AI parsing failed: {e}", "candidate_name": candidate_name}
    return render_template_string(SCREEN_RESUME_HTML, username=session["user"], result=result)

# RECRUITER — ANALYTICS
@app.route("/recruiter/analytics")
def recruiter_analytics():
    if "user" not in session or session.get("role") != "recruiter":
        return redirect("/login")
    conn = get_db()
    total_apps = conn.execute(
        "SELECT COUNT(*) as c FROM applications WHERE job_id IN (SELECT id FROM jobs WHERE recruiter=?)",
        (session["user"],)).fetchone()["c"]
    shortlisted = conn.execute(
        "SELECT COUNT(*) as c FROM applications WHERE status='Shortlisted' AND job_id IN (SELECT id FROM jobs WHERE recruiter=?)",
        (session["user"],)).fetchone()["c"]
    hired = conn.execute(
        "SELECT COUNT(*) as c FROM applications WHERE status='Hired' AND job_id IN (SELECT id FROM jobs WHERE recruiter=?)",
        (session["user"],)).fetchone()["c"]
    rejected = conn.execute(
        "SELECT COUNT(*) as c FROM applications WHERE status='Rejected' AND job_id IN (SELECT id FROM jobs WHERE recruiter=?)",
        (session["user"],)).fetchone()["c"]
    pending = conn.execute(
        "SELECT COUNT(*) as c FROM applications WHERE status='Applied' AND job_id IN (SELECT id FROM jobs WHERE recruiter=?)",
        (session["user"],)).fetchone()["c"]
    jobs = conn.execute("SELECT * FROM jobs WHERE recruiter=? ORDER BY id DESC", (session["user"],)).fetchall()
    # per-job breakdown
    job_stats = []
    for job in jobs:
        cnt = conn.execute("SELECT COUNT(*) as c FROM applications WHERE job_id=?", (job["id"],)).fetchone()["c"]
        job_stats.append({"title": job["title"], "count": cnt, "id": job["id"]})
    history = conn.execute("""
        SELECT * FROM screening_history WHERE recruiter=? ORDER BY id DESC LIMIT 20
    """, (session["user"],)).fetchall() if conn.execute(
        "SELECT name FROM sqlite_master WHERE type='table' AND name='screening_history'"
    ).fetchone() else []
    conn.close()
    return render_template_string(RECRUITER_ANALYTICS_HTML,
        username=session["user"], total_apps=total_apps, shortlisted=shortlisted,
        hired=hired, rejected=rejected, pending=pending,
        jobs=jobs, job_stats=job_stats, history=history)

# RECRUITER — HISTORY
@app.route("/recruiter/history")
def recruiter_history():
    if "user" not in session or session.get("role") != "recruiter":
        return redirect("/login")
    conn = get_db()
    has_table = conn.execute(
        "SELECT name FROM sqlite_master WHERE type='table' AND name='screening_history'"
    ).fetchone()
    history = []
    if has_table:
        history = conn.execute(
            "SELECT * FROM screening_history WHERE recruiter=? ORDER BY id DESC",
            (session["user"],)).fetchall()
    # also include all application status changes
    apps = conn.execute("""
        SELECT a.*, j.title as job_title FROM applications a
        JOIN jobs j ON a.job_id=j.id
        WHERE j.recruiter=? ORDER BY a.id DESC
    """, (session["user"],)).fetchall()
    conn.close()
    return render_template_string(RECRUITER_HISTORY_HTML,
        username=session["user"], history=history, apps=apps)

# ── RECRUITER APPLICATIONS JSON API ─────────────────────────
@app.route("/api/recruiter/applications")
def api_recruiter_applications():
    if "user" not in session or session.get("role") != "recruiter":
        return jsonify({"error": "Unauthorized"}), 401
    conn = get_db()
    apps = conn.execute("""
        SELECT a.id, a.username as applicant, a.job_id, a.job_title, a.company,
               a.status, a.applied_at as created_at, a.resume_json
        FROM applications a JOIN jobs j ON a.job_id=j.id
        WHERE j.recruiter=? ORDER BY a.id DESC
    """, (session["user"],)).fetchall()
    conn.close()
    result = []
    for a in apps:
        row = dict(a)
        try:
            rj = json.loads(row.get("resume_json") or "{}")
            row["ats_score"] = rj.get("ats_score", 0)
            row["cover_note"] = rj.get("summary", "")
        except:
            row["ats_score"] = 0
            row["cover_note"] = ""
        row.pop("resume_json", None)
        result.append(row)
    return jsonify(result)

@app.route("/api/applications/<int:app_id>/status", methods=["POST"])
def update_application_status(app_id):
    if "user" not in session or session.get("role") != "recruiter":
        return jsonify({"error": "Unauthorized"}), 401
    data = request.get_json()
    status = data.get("status", "")
    conn = get_db()
    row = conn.execute("""
        SELECT a.id, a.username, a.job_title FROM applications a
        JOIN jobs j ON a.job_id=j.id
        WHERE a.id=? AND j.recruiter=?
    """, (app_id, session["user"])).fetchone()
    if not row:
        conn.close()
        return jsonify({"error": "Not found"}), 404
    conn.execute("UPDATE applications SET status=? WHERE id=?", (status, app_id))
    push_notification(row["username"], "status",
        f"Your application for '{row['job_title']}' was updated to {status}", "/analytics")
    conn.commit(); conn.close()
    return jsonify({"success": True})

# ── AI — PROJECT IDEAS ────────────────────────────────────────
@app.route("/api/ai/project-ideas", methods=["POST"])
def ai_project_ideas():
    if "user" not in session: return jsonify({"error": "Unauthorized"}), 401
    data = request.get_json()
    role = data.get("role", "Software Developer")
    prompt = f"""Suggest exactly 5 impressive portfolio projects for a {role}.
For each give: project name, one-line description, and tech stack.
Return ONLY a JSON array of strings. Example:
["1. Sales Dashboard — Real-time analytics using React + Python Flask + SQLite",
 "2. Resume Screener — AI-powered ATS using Gemini API + Flask"]"""
    result = gemini(prompt)
    try:
        result = result.strip()
        if "```" in result:
            result = result.split("```")[1].replace("json","").strip()
        return jsonify({"projects": json.loads(result)})
    except:
        return jsonify({"projects": [
            "1. Portfolio Website — Personal showcase using React + Tailwind",
            "2. Task Manager — Full-stack app with Flask + SQLite",
            "3. Resume Analyzer — AI-powered screening with Gemini API",
            "4. Chat Application — Real-time chat with WebSockets",
            "5. E-Commerce Dashboard — Sales analytics with Chart.js"
        ]})

# ── AI — CAREER ROADMAP ───────────────────────────────────────
@app.route("/api/ai/career-roadmap", methods=["POST"])
def ai_career_roadmap():
    if "user" not in session: return jsonify({"error": "Unauthorized"}), 401
    data = request.get_json()
    role = data.get("role", "Software Developer")
    prompt = f"""Create a concise 6-month career roadmap for someone becoming a {role}.
Structure: Month 1-2, Month 3-4, Month 5-6 with specific skills, projects and milestones.
Keep it practical and actionable. Return as plain text, no JSON."""
    result = gemini(prompt)
    return jsonify({"roadmap": result})

@app.route("/jobs")
def job_board():
    if "user" not in session: return redirect("/login")
    conn = get_db()
    query = request.args.get("q", "")
    if query:
        jobs = conn.execute(
            "SELECT * FROM jobs WHERE active=1 AND (title LIKE ? OR company LIKE ? OR skills_required LIKE ?) ORDER BY id DESC",
            (f"%{query}%", f"%{query}%", f"%{query}%")
        ).fetchall()
    else:
        jobs = conn.execute("SELECT * FROM jobs WHERE active=1 ORDER BY id DESC").fetchall()
    conn.close()
    return render_template_string(JOB_BOARD_HTML, jobs=jobs, query=query)

# APPLY
@app.route("/api/jobs/apply", methods=["POST"])
def apply_to_job():
    if "user" not in session: return jsonify({"error": "Unauthorized"}), 401
    data = request.get_json()
    job_id = data.get("job_id")
    conn = get_db()
    existing = conn.execute("SELECT id FROM applications WHERE username=? AND job_id=?", (session["user"], job_id)).fetchone()
    if existing:
        conn.close()
        return jsonify({"error": "Already applied"})
    job = conn.execute("SELECT * FROM jobs WHERE id=?", (job_id,)).fetchone()
    draft = conn.execute("SELECT draft_json FROM resume_drafts WHERE username=?", (session["user"],)).fetchone()
    conn.execute("""
        INSERT INTO applications (username, job_id, job_title, company, applied_at, resume_json)
        VALUES (?,?,?,?,?,?)
    """, (
        session["user"], job_id,
        job["title"] if job else "Unknown",
        job["company"] if job else "Unknown",
        datetime.now().isoformat(),
        draft["draft_json"] if draft else "{}"
    ))
    conn.commit()
    # notify the recruiter
    if job:
        recruiter = job["recruiter"]
        push_notification(
            recruiter, "application",
            f"New application from {session['user']} for '{job['title']}'",
            "/recruiter/applications"
        )
    conn.close()
    return jsonify({"success": True})

# ANALYTICS
@app.route("/analytics")
def analytics_page():
    if "user" not in session: return redirect("/login")
    conn = get_db()
    analytics = conn.execute("SELECT * FROM resume_analytics WHERE username=?", (session["user"],)).fetchone()
    applications = conn.execute("SELECT * FROM applications WHERE username=? ORDER BY id DESC", (session["user"],)).fetchall()
    conn.close()
    stats = dict(analytics) if analytics else {"views": 0, "downloads": 0, "ats_score": 0}
    return render_template_string(ANALYTICS_HTML, stats=stats, applications=applications, username=session["user"])

@app.route("/api/resume-analytics")
def resume_analytics_api():
    if "user" not in session: return jsonify({"error": "Unauthorized"}), 401
    conn = get_db()
    row = conn.execute("SELECT * FROM resume_analytics WHERE username=?", (session["user"],)).fetchone()
    conn.close()
    return jsonify(dict(row) if row else {"views": 0, "downloads": 0, "ats_score": 0})

# AI ROUTES
@app.route("/api/ai/suggest-skills", methods=["POST"])
def ai_suggest_skills():
    if "user" not in session: return jsonify({"error": "Unauthorized"}), 401
    data = request.get_json()
    job_title = data.get("job_title", "Software Engineer")
    prompt = f"""You are a career coach. Suggest exactly 10 highly relevant technical and soft skills for the job title: "{job_title}".
Return ONLY a JSON array of strings, nothing else. Example: ["Python", "SQL", "Problem Solving"]"""
    result = gemini(prompt)
    try:
        result = result.strip()
        if "```" in result:
            result = result.split("```")[1].replace("json","").strip()
        skills = json.loads(result)
        return jsonify({"skills": skills})
    except:
        return jsonify({"skills": ["Python", "Communication", "Problem Solving", "Teamwork", "SQL"]})

@app.route("/api/ai/ats-score", methods=["POST"])
def ai_ats_score():
    if "user" not in session: return jsonify({"error": "Unauthorized"}), 401
    data = request.get_json()
    resume_text = json.dumps(data.get("resume", {}))
    job_desc = data.get("job_description", "")
    prompt = f"""You are an ATS expert. Score this resume against the job description.
RESUME DATA: {resume_text}
JOB DESCRIPTION: {job_desc if job_desc else "General software engineering role"}
Respond ONLY with valid JSON (no extra text):
{{"score": 78, "grade": "B+", "strengths": ["Clear work experience", "Relevant skills listed"], "improvements": ["Add more keywords", "Quantify achievements"], "keyword_match": 65}}"""
    result = gemini(prompt)
    try:
        result = result.strip()
        if "```" in result:
            result = result.split("```")[1].replace("json","").strip()
        parsed = json.loads(result)
        conn = get_db()
        conn.execute("""
            INSERT INTO resume_analytics (username, ats_score) VALUES (?,?)
            ON CONFLICT(username) DO UPDATE SET ats_score=excluded.ats_score
        """, (session["user"], parsed.get("score", 0)))
        conn.commit()
        conn.close()
        return jsonify(parsed)
    except:
        return jsonify({"score": 60, "grade": "C+", "strengths": ["Resume submitted"], "improvements": ["Add more detail"], "keyword_match": 50})

@app.route("/api/ai/job-match", methods=["POST"])
def ai_job_match():
    if "user" not in session: return jsonify({"error": "Unauthorized"}), 401
    data = request.get_json()
    resume_data = data.get("resume", {})
    job_id = data.get("job_id")
    conn = get_db()
    job = conn.execute("SELECT * FROM jobs WHERE id=?", (job_id,)).fetchone()
    conn.close()
    if not job:
        return jsonify({"error": "Job not found"}), 404
    prompt = f"""You are a recruitment AI. Calculate job match percentage.
CANDIDATE RESUME: {json.dumps(resume_data)}
JOB TITLE: {job['title']}
JOB DESCRIPTION: {job['description']}
REQUIRED SKILLS: {job['skills_required']}
Respond ONLY with valid JSON (no extra text):
{{"match_percent": 82, "matched_skills": ["Python", "Flask"], "missing_skills": ["Docker"], "recommendation": "Strong candidate."}}"""
    result = gemini(prompt)
    try:
        result = result.strip()
        if "```" in result:
            result = result.split("```")[1].replace("json","").strip()
        return jsonify(json.loads(result))
    except:
        return jsonify({"match_percent": 70, "matched_skills": [], "missing_skills": [], "recommendation": "Good fit"})

# RESUME SAVE/LOAD/DOWNLOAD
@app.route("/api/resume-builder/save", methods=["POST"])
def save_resume_draft():
    if "user" not in session: return jsonify({"error": "Unauthorized"}), 401
    data = request.get_json()
    conn = get_db()
    existing = conn.execute("SELECT id FROM resume_drafts WHERE username=?", (session["user"],)).fetchone()
    if existing:
        conn.execute("UPDATE resume_drafts SET draft_json=?, updated_at=? WHERE username=?",
                     (json.dumps(data), datetime.now().isoformat(), session["user"]))
    else:
        conn.execute("INSERT INTO resume_drafts (username, draft_json, updated_at) VALUES (?,?,?)",
                     (session["user"], json.dumps(data), datetime.now().isoformat()))
    conn.commit()
    conn.close()
    return jsonify({"success": True})

@app.route("/api/resume-builder/load")
def load_resume_draft():
    if "user" not in session: return jsonify({})
    conn = get_db()
    row = conn.execute("SELECT draft_json FROM resume_drafts WHERE username=?", (session["user"],)).fetchone()
    conn.close()
    return jsonify(json.loads(row["draft_json"]) if row else {})

@app.route("/api/resume-builder/download", methods=["POST"])
def download_resume_pdf():
    if "user" not in session: return jsonify({"error": "Not logged in"}), 401
    data = request.get_json()
    if 'edu' in data and 'education' not in data:
        data['education'] = [{'degree': e.get('degree',''), 'college': e.get('institution', e.get('college','')), 'from': e.get('start',''), 'to': e.get('end','')} for e in data['edu']]
    if 'exp' in data and 'experiences' not in data:
        data['experiences'] = data['exp']
    if 'skill' in data and 'skills' not in data:
        data['skills'] = [s.get('name','') if isinstance(s, dict) else s for s in data['skill']]
    if 'proj' in data and 'projects' not in data:
        data['projects'] = [{'name': p.get('title', p.get('name','')), 'tech': p.get('tech',''), 'desc': p.get('desc','')} for p in data['proj']]
    if 'lang' in data and 'languages' not in data:
        data['languages'] = [l.get('name','') if isinstance(l, dict) else l for l in data['lang']]
    if 'cert' in data and 'certifications' not in data:
        data['certifications'] = [c.get('title','') if isinstance(c, dict) else c for c in data['cert']]
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 50
    margin_l, margin_r = 50, width - 50

    def nl(amount=15):
        nonlocal y
        y -= amount
        if y < 60:
            p.showPage()
            y = height - 50

    def draw_wrapped(text, x, max_width, font="Helvetica", size=11, line_height=14):
        p.setFont(font, size)
        words = str(text).split()
        line = ""
        for word in words:
            test = line + (" " if line else "") + word
            if p.stringWidth(test, font, size) <= max_width:
                line = test
            else:
                p.drawString(x, y, line)
                nl(line_height)
                line = word
        if line:
            p.drawString(x, y, line)
            nl(line_height)

    def section_header(title):
        nl(6)
        p.setFont("Helvetica-Bold", 11)
        p.setFillColorRGB(0.49, 0.36, 0.96)
        p.drawString(margin_l, y, title.upper())
        p.setFillColorRGB(0, 0, 0)
        nl(4)
        p.setStrokeColorRGB(0.49, 0.36, 0.96)
        p.setLineWidth(0.5)
        p.line(margin_l, y, margin_r, y)
        nl(12)

    # Header
    p.setFont("Helvetica-Bold", 22)
    p.drawString(margin_l, y, data.get("name", "Your Name"))
    nl(22)
    if data.get("headline"):
        p.setFont("Helvetica", 13)
        p.setFillColorRGB(0.3, 0.3, 0.3)
        p.drawString(margin_l, y, data["headline"])
        p.setFillColorRGB(0, 0, 0)
        nl(16)
    contact = "  |  ".join(filter(None, [data.get("email",""), data.get("phone",""), data.get("location","")]))
    if contact:
        p.setFont("Helvetica", 10)
        p.setFillColorRGB(0.4, 0.4, 0.4)
        p.drawString(margin_l, y, contact)
        p.setFillColorRGB(0, 0, 0)
        nl(14)
    p.setStrokeColorRGB(0.48, 0.36, 0.96)
    p.setLineWidth(1.5)
    p.line(margin_l, y, margin_r, y)
    nl(14)

    if data.get("summary"):
        section_header("Summary")
        draw_wrapped(data["summary"], margin_l, margin_r - margin_l, size=10, line_height=13)
        nl(4)

    if data.get("skills"):
        section_header("Skills")
        p.setFont("Helvetica", 10)
        skills_line = "  •  ".join(data["skills"])
        draw_wrapped(skills_line, margin_l, margin_r - margin_l, size=10, line_height=13)
        nl(4)

    if data.get("experiences"):
        section_header("Experience")
        for exp in data["experiences"]:
            p.setFont("Helvetica-Bold", 11)
            p.drawString(margin_l, y, exp.get("title", ""))
            nl(13)
            p.setFont("Helvetica", 10)
            p.setFillColorRGB(0.4, 0.4, 0.4)
            sub = f"{exp.get('company','')}  |  {exp.get('start','')} – {exp.get('end','Present')}"
            p.drawString(margin_l, y, sub)
            p.setFillColorRGB(0, 0, 0)
            nl(12)
            if exp.get("desc"):
                draw_wrapped(exp["desc"], margin_l + 8, margin_r - margin_l - 8, size=10, line_height=13)
            nl(6)

    if data.get("education"):
        section_header("Education")
        for edu in data["education"]:
            p.setFont("Helvetica-Bold", 11)
            p.drawString(margin_l, y, edu.get("degree", ""))
            nl(13)
            p.setFont("Helvetica", 10)
            p.setFillColorRGB(0.4, 0.4, 0.4)
            sub = f"{edu.get('college','')}  |  {edu.get('from','')} – {edu.get('to','')}"
            p.drawString(margin_l, y, sub.strip(" |– "))
            p.setFillColorRGB(0, 0, 0)
            nl(16)

    if data.get("projects"):
        section_header("Projects")
        for proj in data["projects"]:
            p.setFont("Helvetica-Bold", 11)
            p.drawString(margin_l, y, proj.get("name", proj.get("title", "")))
            nl(13)
            if proj.get("tech"):
                p.setFont("Helvetica-Oblique", 10)
                p.setFillColorRGB(0.4, 0.4, 0.4)
                p.drawString(margin_l, y, proj["tech"])
                p.setFillColorRGB(0, 0, 0)
                nl(12)
            if proj.get("desc"):
                draw_wrapped(proj["desc"], margin_l + 8, margin_r - margin_l - 8, size=10, line_height=13)
            nl(6)

    if data.get("certifications"):
        section_header("Certifications")
        for c in data["certifications"]:
            p.setFont("Helvetica", 10)
            p.drawString(margin_l, y, f"• {c}")
            nl(13)

    if data.get("languages"):
        section_header("Languages")
        p.setFont("Helvetica", 10)
        p.drawString(margin_l, y, "  •  ".join(data["languages"]))
        nl(13)

    p.save()
    buffer.seek(0)
    name_slug = data.get("name", "resume").replace(" ", "_")
    return send_file(buffer, as_attachment=True, download_name=f"{name_slug}_resume.pdf", mimetype="application/pdf")
  
def _track_download(username):
    try:
        conn = get_db()
        conn.execute("""
            INSERT INTO resume_analytics (username, downloads) VALUES (?,1)
            ON CONFLICT(username) DO UPDATE SET downloads=downloads+1
        """, (username,))
        conn.commit()
        conn.close()
    except: pass

@app.route("/api/resume/downloaded", methods=["POST"])
def resume_downloaded():
    if "user" not in session: return jsonify({"error": "Unauthorized"}), 401
    _track_download(session["user"])
    return jsonify({"success": True})

@app.route("/api/ai/ats-score-upload", methods=["POST"])
def ai_ats_score_upload():
    """Score an uploaded PDF resume against an optional job description."""
    if "user" not in session:
        return jsonify({"error": "Unauthorized"}), 401
    file = request.files.get("resume_pdf")
    job_desc = request.form.get("job_description", "")
    if not file:
        return jsonify({"error": "No file uploaded"}), 400
    try:
        pdf_bytes = file.read()
        if PDF_EXTRACT_OK:
            from io import BytesIO as _BIO
            resume_text = pdf_extract_text(_BIO(pdf_bytes))
            resume_text = resume_text[:4000]  # trim to avoid token limits
        else:
            resume_text = "(PDF text extraction unavailable)"
    except Exception as e:
        resume_text = f"(Could not extract PDF text: {e})"

    prompt = f"""You are an ATS (Applicant Tracking System) expert. Analyse this resume text and score it.

RESUME TEXT:
{resume_text}

JOB DESCRIPTION: {job_desc if job_desc else "General software/tech role"}

Respond ONLY with a valid JSON object (no extra text, no markdown):
{{
  "score": 74,
  "grade": "B",
  "summary": "2-3 sentence overall assessment of the candidate.",
  "strengths": ["Clear skills section", "Relevant experience listed"],
  "improvements": ["Add quantified achievements", "Include more keywords from job description"],
  "keyword_match": 58,
  "sections_found": ["Experience", "Education", "Skills"],
  "sections_missing": ["Summary", "Certifications"]
}}"""
    result = gemini(prompt)
    try:
        result = result.strip()
        if "```" in result:
            result = result.split("```")[1].replace("json", "").strip()
        parsed = json.loads(result)
        # Save score to analytics
        conn = get_db()
        conn.execute("""
            INSERT INTO resume_analytics (username, ats_score) VALUES (?,?)
            ON CONFLICT(username) DO UPDATE SET ats_score=excluded.ats_score
        """, (session["user"], parsed.get("score", 0)))
        conn.commit()
        conn.close()
        return jsonify(parsed)
    except Exception as e:
        return jsonify({
            "score": 60, "grade": "C+",
            "summary": "Resume received. Please ensure it has clear sections.",
            "strengths": ["Resume submitted successfully"],
            "improvements": ["Add more detail to experience", "Include a summary section"],
            "keyword_match": 50,
            "sections_found": ["Resume"],
            "sections_missing": []
        })


@app.route("/api/resume-builder/save-version", methods=["POST"])
def save_resume_version():
    if "user" not in session: return jsonify({"error": "Unauthorized"}), 401
    data = request.get_json()
    version_name = data.get("version_name", f"Version {datetime.now().strftime('%d %b %Y')}")
    conn = get_db()
    conn.execute(
        "INSERT INTO resume_versions (username, version_name, resume_json, created_at) VALUES (?,?,?,?)",
        (session["user"], version_name, json.dumps(data), datetime.now().isoformat())
    )
    conn.commit()
    conn.close()
    return jsonify({"success": True})

@app.route("/api/resume-builder/versions")
def list_resume_versions():
    if "user" not in session: return jsonify([])
    conn = get_db()
    rows = conn.execute(
        "SELECT id, version_name, created_at FROM resume_versions WHERE username=? ORDER BY id DESC",
        (session["user"],)
    ).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])


# ============================================================
# HTML TEMPLATES
# ============================================================

LOGIN_HTML = """<!DOCTYPE html>
<html><head><title>Login – RecruitAI</title>
<link href="https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;600;700&family=DM+Mono:wght@400;500&display=swap" rel="stylesheet">
<style>
*{margin:0;padding:0;box-sizing:border-box;}
body{
    font-family:'DM Sans',sans-serif;
    background:linear-gradient(
        135deg,
        #0f172a 0%,
        #111827 50%,
        #1e293b 100%
    );
    color:#f8fafc;
    display:flex;
    align-items:center;
    justify-content:center;
    min-height:100vh;
}
body::before{content:'';position:fixed;top:-200px;left:-200px;width:600px;height:600px;background:radial-gradient(circle,rgba(99,76,255,0.12) 0%,transparent 70%);pointer-events:none;}
body::after{content:'';position:fixed;bottom:-200px;right:-200px;width:500px;height:500px;background:radial-gradient(circle,rgba(168,85,247,0.08) 0%,transparent 70%);pointer-events:none;}
.box{background:#0f1018;padding:44px;border-radius:24px;width:400px;border:1px solid rgba(255,255,255,0.07);position:relative;z-index:1;}
.logo{font-size:22px;font-weight:700;color:#a78bfa;margin-bottom:6px;letter-spacing:-0.5px;}
.logo span{color:#6366f1;}
h2{font-size:26px;font-weight:700;margin-bottom:4px;letter-spacing:-0.5px;}
.sub{color:#4a4b5a;font-size:13px;margin-bottom:28px;}
.role-toggle{display:flex;gap:8px;margin-bottom:24px;background:#0a0b12;padding:4px;border-radius:14px;border:1px solid rgba(255,255,255,0.05);}
.role-btn{flex:1;padding:10px;border:none;border-radius:10px;background:transparent;color:#555;font-size:13px;font-weight:600;cursor:pointer;transition:all .2s;font-family:'DM Sans',sans-serif;}
.role-btn.active{background:#1a1b2e;color:#c4b5fd;box-shadow:0 2px 8px rgba(99,76,255,0.15);}
label{display:block;margin-bottom:5px;font-size:12px;color:#555;font-weight:500;}
input{width:100%;padding:12px 14px;border:1px solid rgba(255,255,255,0.07);outline:none;border-radius:12px;background:#0a0b12;color:white;font-size:14px;margin-bottom:14px;transition:border-color .2s;font-family:'DM Sans',sans-serif;}
input:focus{border-color:#6366f1;}
button[type=submit]{width:100%;padding:13px;border:none;border-radius:12px;background:linear-gradient(135deg,#6366f1,#a78bfa);color:white;font-weight:700;cursor:pointer;font-size:14px;font-family:'DM Sans',sans-serif;transition:opacity .2s;}
button[type=submit]:hover{opacity:.9;}
.err{color:#f87171;font-size:13px;margin-bottom:14px;background:rgba(248,113,113,.08);padding:10px 14px;border-radius:10px;border:1px solid rgba(248,113,113,.15);}
a{color:#a78bfa;text-decoration:none;font-size:13px;}
.link-row{margin-top:18px;text-align:center;color:#333;}
</style></head><body>
<div class="box">
  <div class="logo">✦ Recruit<span>AI</span></div>
  <h2>Welcome back</h2>
  <p class="sub">Sign in to your account</p>
  {% if error %}<div class="err">{{ error }}</div>{% endif %}
  <div class="role-toggle">
    <button type="button" class="role-btn active" id="btn-jobseeker" onclick="setRole('jobseeker')">👤 Candidate</button>
    <button type="button" class="role-btn" id="btn-recruiter" onclick="setRole('recruiter')">🏢 Recruiter</button>
  </div>
  <form method="POST">
    <input type="hidden" name="role" id="roleInput" value="jobseeker">
    <label>Username</label><input type="text" name="username" required placeholder="Enter your username">
    <label>Password</label><input type="password" name="password" required placeholder="Enter your password">
    <button type="submit" id="loginBtn">Sign in as Candidate</button>
  </form>
  <div class="link-row"><a href="/register">Don't have an account? <span style="color:#a78bfa;">Create one →</span></a></div>
</div>
<script>
function setRole(r){
  document.getElementById("roleInput").value=r;
  document.getElementById("btn-jobseeker").classList.toggle("active",r==="jobseeker");
  document.getElementById("btn-recruiter").classList.toggle("active",r==="recruiter");
  document.getElementById("loginBtn").textContent=r==="recruiter"?"Sign in as Recruiter":"Sign in as Candidate";
}
</script></body></html>"""

REGISTER_HTML = """<!DOCTYPE html>
<html><head><title>Register – RecruitAI</title>
<link href="https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;600;700&display=swap" rel="stylesheet">
<style>
*{margin:0;padding:0;box-sizing:border-box;}
body{font-family:'DM Sans',sans-serif;background:#08090e;color:white;display:flex;align-items:center;justify-content:center;min-height:100vh;padding:30px 0;}
.box{background:#0f1018;padding:40px;border-radius:24px;width:500px;border:1px solid rgba(255,255,255,0.07);}
.logo{font-size:20px;font-weight:700;color:#a78bfa;margin-bottom:6px;}
h2{font-size:24px;font-weight:700;margin-bottom:4px;letter-spacing:-0.5px;}
.sub{color:#4a4b5a;font-size:13px;margin-bottom:24px;}
.role-toggle{display:flex;gap:8px;margin-bottom:22px;background:#0a0b12;padding:4px;border-radius:14px;border:1px solid rgba(255,255,255,0.05);}
.role-btn{flex:1;padding:10px;border:none;border-radius:10px;background:transparent;color:#555;font-size:13px;font-weight:600;cursor:pointer;transition:all .2s;font-family:'DM Sans',sans-serif;}
.role-btn.active{background:#1a1b2e;color:#c4b5fd;}
.row{display:grid;grid-template-columns:1fr 1fr;gap:12px;}
label{display:block;margin-bottom:4px;font-size:12px;color:#555;font-weight:500;}
input,select{width:100%;padding:11px 13px;border:1px solid rgba(255,255,255,0.07);outline:none;border-radius:11px;background:#0a0b12;color:white;font-size:13px;margin-bottom:13px;transition:border-color .2s;font-family:'DM Sans',sans-serif;}
input:focus,select:focus{border-color:#6366f1;}
select option{background:#0f1018;}
.section-label{font-size:11px;color:#6366f1;text-transform:uppercase;letter-spacing:.8px;margin-bottom:10px;margin-top:2px;font-weight:600;}
button[type=submit]{width:100%;padding:13px;border:none;border-radius:12px;background:linear-gradient(135deg,#6366f1,#a78bfa);color:white;font-weight:700;cursor:pointer;font-size:14px;font-family:'DM Sans',sans-serif;}
.err{color:#f87171;font-size:13px;margin-bottom:14px;background:rgba(248,113,113,.08);padding:10px 14px;border-radius:10px;}
a{color:#a78bfa;text-decoration:none;font-size:13px;}
.link-row{margin-top:16px;text-align:center;}
.divider{border:none;border-top:1px solid rgba(255,255,255,0.05);margin:16px 0;}
.hidden{display:none !important;}
</style></head><body>
<div class="box">
  <div class="logo">✦ RecruitAI</div>
  <h2>Create Account</h2>
  <p class="sub">Register as Candidate or Recruiter</p>
  {% if error %}<div class="err">{{ error }}</div>{% endif %}
  <div class="role-toggle">
    <button type="button" class="role-btn active" id="btn-jobseeker" onclick="setRole('jobseeker')">👤 Candidate</button>
    <button type="button" class="role-btn" id="btn-recruiter" onclick="setRole('recruiter')">🏢 Recruiter</button>
  </div>
  <form method="POST">
    <input type="hidden" name="role" id="roleInput" value="jobseeker">
    <div class="section-label">Account Details</div>
    <div class="row">
      <div><label>Username *</label><input type="text" name="username" required placeholder="e.g. rahul123"></div>
      <div><label>Password *</label><input type="password" name="password" required placeholder="Min 6 characters"></div>
    </div>
    <div class="divider"></div>
    <div class="section-label">Personal Details</div>
    <div class="row">
      <div><label>Full Name *</label><input type="text" name="full_name" required placeholder="Rahul Sharma"></div>
      <div><label>Phone</label><input type="text" name="phone" placeholder="+91 98765 43210"></div>
    </div>
    <label>Email *</label><input type="email" name="email" required placeholder="rahul@email.com">
    <div class="divider"></div>
    <div class="section-label" id="extraLabel">Candidate Details</div>
    <div id="candidateFields">
      <div class="row">
        <div><label>City</label><input type="text" name="city" placeholder="Mumbai"></div>
        <div><label>State</label><input type="text" name="state" placeholder="Maharashtra"></div>
      </div>
      <label>Current Job Title</label><input type="text" name="job_title" placeholder="e.g. Software Developer">
      <label>Experience</label>
      <select name="experience">
        <option value="">Select Experience</option>
        <option>Fresher</option>
        <option>0–1 years</option>
        <option>1–3 years</option>
        <option>3–5 years</option>
        <option>5+ years</option>
      </select>
      <label>Highest Education</label>
      <select name="education">
        <option value="">Select Education</option>
        <option>High School</option>
        <option>Diploma</option>
        <option>B.Tech / B.E.</option>
        <option>BCA / B.Sc</option>
        <option>MBA</option>
        <option>MCA / M.Tech</option>
        <option>PhD</option>
      </select>
    </div>
    <div id="recruiterFields" class="hidden">
      <div class="row">
        <div><label>Company Name *</label><input type="text" name="company_name" placeholder="e.g. Infosys"></div>
        <div><label>Company Size</label>
          <select name="company_size">
            <option value="">Select Size</option>
            <option>1–10</option><option>11–50</option><option>51–200</option><option>201–500</option><option>500+</option>
          </select>
        </div>
      </div>
      <label>Industry</label>
      <select name="industry">
        <option value="">Select Industry</option>
        <option>Information Technology</option><option>Finance & Banking</option>
        <option>Healthcare</option><option>Education</option><option>E-Commerce</option>
        <option>Manufacturing</option><option>Consulting</option><option>Other</option>
      </select>
      <label>Your Designation</label>
      <input type="text" name="designation" placeholder="e.g. HR Manager, Talent Acquisition Lead">
    </div>
    <button type="submit" id="regBtn">Register as Candidate</button>
  </form>
  <div class="link-row"><a href="/login">Already have an account? <span style="color:#a78bfa;">Sign in →</span></a></div>
</div>
<script>
function setRole(r){
  document.getElementById("roleInput").value = r;
  document.getElementById("btn-jobseeker").classList.toggle("active", r==="jobseeker");
  document.getElementById("btn-recruiter").classList.toggle("active", r==="recruiter");
  document.getElementById("regBtn").textContent = r==="recruiter" ? "Register as Recruiter" : "Register as Candidate";
  document.getElementById("extraLabel").textContent = r==="recruiter" ? "Recruiter Details" : "Candidate Details";

  var cf = document.getElementById("candidateFields");
  var rf = document.getElementById("recruiterFields");

  if(r === "recruiter"){
    cf.style.display = "none";
    rf.style.display = "block";
    rf.classList.remove("hidden");
  } else {
    cf.style.display = "block";
    rf.style.display = "none";
    rf.classList.add("hidden");
  }

  /* disable hidden fields so they don't interfere with form validation */
  document.querySelectorAll("#candidateFields input, #candidateFields select").forEach(function(el){
    el.disabled = (r === "recruiter");
  });
  document.querySelectorAll("#recruiterFields input, #recruiterFields select").forEach(function(el){
    el.disabled = (r === "jobseeker");
  });
}

document.addEventListener("DOMContentLoaded", function(){ setRole("jobseeker"); });
</script></body></html>"""


# Shared sidebar CSS
SIDEBAR_CSS = """
*{margin:0;padding:0;box-sizing:border-box;}
body{font-family:'DM Sans',sans-serif;background:#08090e;color:white;display:flex;min-height:100vh;}
.sidebar{width:240px;min-width:240px;background:#0c0d15;border-right:1px solid rgba(255,255,255,0.05);display:flex;flex-direction:column;padding:24px 0;height:100vh;position:fixed;left:0;top:0;}
.logo{font-size:18px;font-weight:700;color:#a78bfa;padding:0 20px 24px;letter-spacing:-0.3px;border-bottom:1px solid rgba(255,255,255,0.05);}
.logo span{color:#6366f1;}
.nav-section{padding:20px 12px 8px;flex:1;}
.nav-label{font-size:10px;color:#2a2b3a;text-transform:uppercase;letter-spacing:1px;font-weight:600;padding:0 8px;margin-bottom:6px;}
.nav-item{display:flex;align-items:center;gap:10px;padding:10px 12px;border-radius:10px;color:#4a4b5e;font-size:13px;font-weight:500;text-decoration:none;transition:all .18s;margin-bottom:2px;cursor:pointer;}
.nav-item:hover{background:#12131f;color:#9ca3c0;}
.nav-item.active{background:#16172a;color:#c4b5fd;}
.nav-item .icon{width:18px;text-align:center;font-size:15px;}
.sidebar-bottom{padding:16px 12px;border-top:1px solid rgba(255,255,255,0.05);}
.user-chip{display:flex;align-items:center;gap:10px;padding:10px 12px;border-radius:10px;background:#12131f;}
.avatar{width:32px;height:32px;border-radius:8px;background:linear-gradient(135deg,#6366f1,#a78bfa);display:flex;align-items:center;justify-content:center;font-weight:700;font-size:13px;flex-shrink:0;}
.user-info .name{font-size:13px;font-weight:600;color:#ddd;}
.user-info .role{font-size:11px;color:#444;}
.logout-link{display:block;text-align:center;margin-top:8px;font-size:12px;color:#333;text-decoration:none;padding:8px;border-radius:8px;transition:all .2s;}
.logout-link:hover{background:#12131f;color:#f87171;}
.main{margin-left:240px;flex:1;min-height:100vh;}
.topbar{padding:20px 32px;border-bottom:1px solid rgba(255,255,255,0.04);display:flex;justify-content:space-between;align-items:center;background:#08090e;position:sticky;top:0;z-index:10;}
.topbar h1{font-size:20px;font-weight:700;letter-spacing:-0.3px;}
.topbar .sub{font-size:13px;color:#3a3b4a;margin-top:2px;}
.content{padding:28px 32px;}
.stats-grid{display:grid;grid-template-columns:repeat(auto-fit,minmax(160px,1fr));gap:14px;margin-bottom:28px;}
.stat-card{background:#0f1018;border-radius:16px;padding:20px;border:1px solid rgba(255,255,255,0.05);transition:border-color .2s;}
.stat-card:hover{border-color:rgba(99,102,241,0.2);}
.stat-card .s-icon{font-size:20px;margin-bottom:10px;}
.stat-card .s-val{font-size:28px;font-weight:700;letter-spacing:-1px;color:#e2e8f0;}
.stat-card .s-label{font-size:11px;color:#3a3b4a;margin-top:3px;font-weight:500;text-transform:uppercase;letter-spacing:.5px;}
.stat-card .s-hint{font-size:11px;color:#2a2b38;margin-top:6px;}
.card{background:#0f1018;border-radius:16px;border:1px solid rgba(255,255,255,0.05);margin-bottom:20px;}
.card-header{padding:18px 22px;border-bottom:1px solid rgba(255,255,255,0.04);display:flex;justify-content:space-between;align-items:center;}
.card-header h2{font-size:14px;font-weight:600;color:#c4c8d8;}
.card-body{padding:4px 0;}
table{width:100%;border-collapse:collapse;}
th{padding:10px 22px;text-align:left;font-size:11px;color:#2a2b3a;text-transform:uppercase;letter-spacing:.5px;font-weight:600;}
td{padding:12px 22px;border-top:1px solid rgba(255,255,255,0.03);font-size:13px;color:#9ca3c0;}
td strong{color:#dde1ee;}
tr:hover td{background:rgba(255,255,255,0.01);}
.badge{display:inline-flex;align-items:center;padding:3px 10px;border-radius:99px;font-size:11px;font-weight:600;}
.badge-purple{background:rgba(99,102,241,0.12);color:#818cf8;}
.badge-green{background:rgba(52,211,153,0.1);color:#34d399;}
.badge-red{background:rgba(248,113,113,0.1);color:#f87171;}
.badge-yellow{background:rgba(251,191,36,0.1);color:#fbbf24;}
.badge-blue{background:rgba(96,165,250,0.1);color:#60a5fa;}
.btn{display:inline-flex;align-items:center;gap:6px;padding:8px 16px;border-radius:10px;font-size:12px;font-weight:600;border:none;cursor:pointer;transition:all .18s;text-decoration:none;font-family:'DM Sans',sans-serif;}
.btn-primary{background:linear-gradient(135deg,#6366f1,#8b5cf6);color:white;}
.btn-primary:hover{opacity:.9;}
.btn-ghost{background:#12131f;color:#6b7280;border:1px solid rgba(255,255,255,0.05);}
.btn-ghost:hover{color:#c4b5fd;border-color:rgba(99,102,241,0.3);}
.btn-sm{padding:5px 12px;font-size:11px;}
.empty-state{text-align:center;padding:60px 40px;color:#2a2b3a;}
.empty-state .e-icon{font-size:36px;margin-bottom:12px;opacity:.4;}
.empty-state h3{font-size:15px;color:#3a3b4a;margin-bottom:6px;}
.empty-state p{font-size:13px;}
.modal-overlay{display:none;position:fixed;inset:0;background:rgba(0,0,0,0.75);z-index:200;align-items:center;justify-content:center;backdrop-filter:blur(4px);}
.modal-overlay.show{display:flex;}
.modal{background:#0f1018;border-radius:20px;padding:28px;width:480px;max-width:92vw;border:1px solid rgba(255,255,255,0.08);}
.modal-header{display:flex;justify-content:space-between;align-items:center;margin-bottom:18px;}
.modal-header h2{font-size:16px;font-weight:700;color:#c4b5fd;}
.modal-close{background:none;border:none;color:#444;font-size:22px;cursor:pointer;line-height:1;}
.modal-close:hover{color:#aaa;}
.score-ring{display:inline-flex;align-items:center;justify-content:center;width:80px;height:80px;border-radius:50%;border:3px solid #6366f1;font-size:24px;font-weight:700;color:#a78bfa;}
.tag{display:inline-flex;align-items:center;padding:4px 10px;border-radius:99px;font-size:11px;margin:3px;}
.tag-g{background:rgba(52,211,153,0.1);color:#34d399;}
.tag-r{background:rgba(248,113,113,0.1);color:#f87171;}

.notif-bell{position:relative;cursor:pointer;background:#12131f;border:1px solid rgba(255,255,255,0.07);border-radius:10px;width:38px;height:38px;display:flex;align-items:center;justify-content:center;font-size:16px;transition:all .2s;flex-shrink:0;}
.notif-bell:hover{background:#1a1b2e;border-color:rgba(99,102,241,0.3);}
.notif-badge{position:absolute;top:-5px;right:-5px;background:#f87171;color:white;border-radius:99px;font-size:10px;font-weight:700;min-width:18px;height:18px;display:flex;align-items:center;justify-content:center;padding:0 4px;border:2px solid #08090e;}
.notif-panel{position:absolute;top:48px;right:0;width:340px;background:#0f1018;border:1px solid rgba(255,255,255,0.08);border-radius:16px;box-shadow:0 16px 48px rgba(0,0,0,0.5);z-index:1000;display:none;overflow:hidden;}
.notif-panel.open{display:block;}
.notif-header{padding:14px 18px;border-bottom:1px solid rgba(255,255,255,0.05);display:flex;justify-content:space-between;align-items:center;}
.notif-header h3{font-size:13px;font-weight:700;color:#c4b5fd;}
.notif-clear{font-size:11px;color:#6366f1;cursor:pointer;background:none;border:none;font-family:'DM Sans',sans-serif;font-weight:600;}
.notif-clear:hover{color:#a78bfa;}
.notif-list{max-height:320px;overflow-y:auto;}
.notif-item{padding:12px 18px;border-bottom:1px solid rgba(255,255,255,0.03);display:flex;gap:10px;align-items:flex-start;transition:background .15s;cursor:pointer;}
.notif-item:hover{background:rgba(255,255,255,0.02);}
.notif-item.unread{background:rgba(99,102,241,0.04);}
.notif-icon{width:32px;height:32px;border-radius:8px;display:flex;align-items:center;justify-content:center;font-size:14px;flex-shrink:0;}
.notif-msg{font-size:12px;color:#9ca3c0;line-height:1.5;flex:1;}
.notif-time{font-size:10px;color:#2a2b3a;margin-top:3px;}
.notif-unread-dot{width:6px;height:6px;border-radius:50%;background:#6366f1;flex-shrink:0;margin-top:5px;}
.notif-empty{padding:32px;text-align:center;color:#2a2b3a;font-size:13px;}

.toast{position:fixed;bottom:24px;right:24px;background:#1a1b2e;color:#c4b5fd;padding:12px 18px;border-radius:12px;font-size:13px;font-weight:500;opacity:0;transition:opacity .3s;pointer-events:none;z-index:9999;border:1px solid rgba(99,102,241,0.3);box-shadow:0 8px 32px rgba(0,0,0,0.4);}
.toast.show{opacity:1;}
"""


CANDIDATE_DASHBOARD_HTML = """<!DOCTYPE html>
<html><head><title>My Dashboard – RecruitAI</title>
<link href="https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;600;700&display=swap" rel="stylesheet">
<style>""" + SIDEBAR_CSS + """
.progress-track{display:flex;align-items:center;gap:0;margin:16px 0;}
.prog-step{display:flex;flex-direction:column;align-items:center;flex:1;position:relative;}
.prog-step:not(:last-child)::after{content:'';position:absolute;top:14px;left:50%;width:100%;height:2px;background:#1a1b2e;z-index:0;}
.prog-step.done:not(:last-child)::after{background:#6366f1;}
.prog-dot{width:28px;height:28px;border-radius:50%;background:#12131f;border:2px solid #1a1b2e;display:flex;align-items:center;justify-content:center;font-size:11px;font-weight:700;position:relative;z-index:1;transition:all .3s;}
.prog-step.done .prog-dot{background:#6366f1;border-color:#6366f1;color:white;}
.prog-step.active .prog-dot{background:#1a1b2e;border-color:#a78bfa;color:#a78bfa;box-shadow:0 0 12px rgba(99,102,241,0.3);}
.prog-label{font-size:10px;color:#2a2b3a;margin-top:6px;text-align:center;font-weight:500;}
.prog-step.done .prog-label,.prog-step.active .prog-label{color:#6b7280;}
.quick-actions{display:grid;grid-template-columns:repeat(2,1fr);gap:10px;margin-bottom:20px;}
.qa-btn{background:#0f1018;border:1px solid rgba(255,255,255,0.05);border-radius:14px;padding:18px;cursor:pointer;transition:all .2s;text-decoration:none;display:flex;align-items:center;gap:12px;}
.qa-btn:hover{border-color:rgba(99,102,241,0.3);background:#12131f;}
.qa-icon{width:38px;height:38px;border-radius:10px;display:flex;align-items:center;justify-content:center;font-size:18px;}
.qa-text .qa-title{font-size:13px;font-weight:600;color:#c4c8d8;}
.qa-text .qa-sub{font-size:11px;color:#2a2b3a;margin-top:2px;}
.upload-zone{border:2px dashed rgba(99,102,241,0.3);border-radius:14px;padding:24px;text-align:center;cursor:pointer;transition:all .2s;background:#0a0b12;position:relative;}
.upload-zone:hover,.upload-zone.drag{border-color:#6366f1;background:rgba(99,102,241,0.05);}
.upload-zone input[type=file]{position:absolute;inset:0;opacity:0;cursor:pointer;width:100%;height:100%;}
.upload-zone .uz-icon{font-size:28px;margin-bottom:6px;}
.upload-zone .uz-title{font-size:13px;font-weight:600;color:#c4b5fd;margin-bottom:3px;}
.upload-zone .uz-sub{font-size:11px;color:#3a3b4a;}
.tab-row{display:flex;gap:0;background:#0a0b12;border-radius:11px;padding:3px;border:1px solid rgba(255,255,255,0.05);margin-bottom:14px;}
.tab-btn{flex:1;padding:8px;border:none;border-radius:8px;background:transparent;color:#555;font-size:12px;font-weight:600;cursor:pointer;transition:all .2s;font-family:'DM Sans',sans-serif;}
.tab-btn.active{background:#16172a;color:#c4b5fd;}
.score-ring{display:inline-flex;align-items:center;justify-content:center;width:72px;height:72px;border-radius:50%;border:3px solid #6366f1;font-size:22px;font-weight:700;color:#a78bfa;flex-shrink:0;}
</style></head><body>

<div class="sidebar">
  <div class="logo">✦ Recruit<span>AI</span></div>
  <div class="nav-section">
    <div class="nav-label">Main</div>
    <a class="nav-item active" href="/dashboard"><span class="icon">🏠</span> Dashboard</a>
    <a class="nav-item" href="/resume-builder"><span class="icon">📄</span> Resume Builder</a>
    <a class="nav-item" href="/jobs"><span class="icon">💼</span> Browse Jobs</a>
    <a class="nav-item {% if request.path == '/job-finder' %}active{% endif %}" href="/job-finder"><span class="icon">🔍</span> Job Finder</a>
    <a class="nav-item" href="/analytics"><span class="icon">📊</span> Analytics</a>
  </div>
  <div class="sidebar-bottom">
    <div class="user-chip">
      <div class="avatar">{{ username[0].upper() }}</div>
      <div class="user-info"><div class="name">{{ username }}</div><div class="role">Candidate</div></div>
    </div>
    <a href="/logout" class="logout-link">← Sign out</a>
  </div>
</div>

<div class="main">
  <div class="topbar">
    <div>
      <h1>Welcome back, {{ username }} 👋</h1>
      <div class="sub">Here's your job search overview</div>
    </div>
    <a href="/resume-builder" class="btn btn-primary">+ Build Resume</a>
  </div>
  <div class="content">

    <div class="stats-grid">
      <div class="stat-card"><div class="s-icon">📄</div><div class="s-val">{{ stats.downloads }}</div><div class="s-label">Resume Downloads</div></div>
      <div class="stat-card"><div class="s-icon">🎯</div><div class="s-val">{{ stats.ats_score or '—' }}</div><div class="s-label">ATS Score</div></div>
      <div class="stat-card"><div class="s-icon">📨</div><div class="s-val">{{ total_apps }}</div><div class="s-label">Applications</div></div>
      <div class="stat-card"><div class="s-icon">⭐</div><div class="s-val">{{ shortlisted }}</div><div class="s-label">Shortlisted</div></div>
    </div>

    <div style="display:grid;grid-template-columns:1fr 1fr;gap:16px;margin-bottom:20px;">
      <div>
        <div style="font-size:13px;font-weight:600;color:#555;margin-bottom:10px;text-transform:uppercase;letter-spacing:.5px;">Quick Actions</div>
        <div class="quick-actions">
          <a href="/resume-builder" class="qa-btn">
            <div class="qa-icon" style="background:rgba(99,102,241,0.12);">✏️</div>
            <div class="qa-text"><div class="qa-title">Edit Resume</div><div class="qa-sub">Update your profile</div></div>
          </a>
          <a href="/jobs" class="qa-btn">
            <div class="qa-icon" style="background:rgba(52,211,153,0.1);">🔍</div>
            <div class="qa-text"><div class="qa-title">Browse Jobs</div><div class="qa-sub">Find opportunities</div></div>
          </a>
          <a href="/analytics" class="qa-btn">
            <div class="qa-icon" style="background:rgba(251,191,36,0.1);">📊</div>
            <div class="qa-text"><div class="qa-title">ATS Score</div><div class="qa-sub">Check resume score</div></div>
          </a>
          <a href="/analytics" class="qa-btn">
            <div class="qa-icon" style="background:rgba(248,113,113,0.1);">🤖</div>
            <div class="qa-text"><div class="qa-title">AI Skills</div><div class="qa-sub">Get suggestions</div></div>
          </a>
        </div>
      </div>
      <div>
        <div style="font-size:13px;font-weight:600;color:#555;margin-bottom:10px;text-transform:uppercase;letter-spacing:.5px;">Application Progress</div>
        <div class="card" style="padding:20px;">
          {% if applications %}
          {% set app = applications[0] %}
          <div style="font-size:13px;font-weight:600;color:#c4c8d8;margin-bottom:4px;">{{ app.job_title }}</div>
          <div style="font-size:11px;color:#333;margin-bottom:14px;">{{ app.company }}</div>
          {% set stages = ['Applied','Reviewing','Shortlisted','Interview','Offer'] %}
          {% set status_map = {'Applied':1,'Reviewing':2,'Shortlisted':3,'Interviewing':3,'Interview':4,'Hired':5,'Offer':5} %}
          {% set current = status_map.get(app.status, 1) %}
          <div class="progress-track">
            {% for i, stage in [(1,'Applied'),(2,'Review'),(3,'Shortlist'),(4,'Interview'),(5,'Offer')] %}
            <div class="prog-step {{ 'done' if current > i else ('active' if current == i else '') }}">
              <div class="prog-dot">{{ '✓' if current > i else i }}</div>
              <div class="prog-label">{{ stage }}</div>
            </div>
            {% endfor %}
          </div>
          <div style="margin-top:10px;display:inline-flex;align-items:center;gap:6px;">
            {% if app.status == 'Applied' %}<span class="badge badge-purple">{{ app.status }}</span>
            {% elif app.status in ['Shortlisted','Reviewing'] %}<span class="badge badge-green">{{ app.status }}</span>
            {% elif app.status == 'Rejected' %}<span class="badge badge-red">{{ app.status }}</span>
            {% elif app.status in ['Interviewing','Interview'] %}<span class="badge badge-yellow">{{ app.status }}</span>
            {% else %}<span class="badge badge-blue">{{ app.status }}</span>{% endif %}
          </div>
          {% else %}
          <div style="text-align:center;padding:20px;color:#2a2b3a;">
            <div style="font-size:24px;margin-bottom:8px;">📭</div>
            <div style="font-size:13px;">No applications yet</div>
            <a href="/jobs" style="color:#a78bfa;font-size:12px;text-decoration:none;">Browse jobs →</a>
          </div>
          {% endif %}
        </div>
      </div>
    </div>

    <!-- ===== RESUME UPLOAD + ATS + AI FEEDBACK ===== -->
    <div class="card" style="margin-bottom:20px;">
      <div class="card-header">
        <h2>📄 Resume Screening &amp; AI Feedback</h2>
        <span style="font-size:11px;color:#3a3b4a;">Upload your PDF resume or use your Builder resume</span>
      </div>
      <div style="padding:20px;">
        <div class="tab-row">
          <button class="tab-btn active" id="dtab-upload" onclick="dSwitchTab('upload')">📤 Upload PDF</button>
          <button class="tab-btn" id="dtab-builder" onclick="dSwitchTab('builder')">✏️ Builder Resume</button>
        </div>

        <!-- Upload panel -->
        <div id="dpanel-upload">
          <div class="upload-zone" id="dDropZone">
            <input type="file" id="dResumeFile" accept=".pdf" onchange="dOnFileSelect(this)">
            <div class="uz-icon">📄</div>
            <div class="uz-title" id="dUzTitle">Drop your resume PDF here or click to browse</div>
            <div class="uz-sub">PDF only &nbsp;•&nbsp; Max 5MB</div>
          </div>
          <textarea id="dJobDescUpload" rows="2" placeholder="Paste a job description for tailored ATS score (optional)..."
            style="width:100%;margin-top:10px;padding:10px 13px;border:1px solid rgba(255,255,255,0.07);outline:none;border-radius:11px;background:#0a0b12;color:white;font-size:12px;font-family:'DM Sans',sans-serif;resize:vertical;"></textarea>
          <button class="btn btn-primary" style="width:100%;justify-content:center;margin-top:10px;padding:12px;" onclick="dScoreUpload()">🎯 Analyse &amp; Get AI Feedback</button>
        </div>

        <!-- Builder panel -->
        <div id="dpanel-builder" style="display:none;">
          <p style="font-size:12px;color:#3a3b4a;margin-bottom:10px;">Scores your saved Resume Builder data.</p>
          <textarea id="dJobDescBuilder" rows="2" placeholder="Paste a job description for tailored ATS score (optional)..."
            style="width:100%;padding:10px 13px;border:1px solid rgba(255,255,255,0.07);outline:none;border-radius:11px;background:#0a0b12;color:white;font-size:12px;font-family:'DM Sans',sans-serif;resize:vertical;"></textarea>
          <button class="btn btn-primary" style="width:100%;justify-content:center;margin-top:10px;padding:12px;" onclick="dScoreBuilder()">🎯 Analyse &amp; Get AI Feedback</button>
        </div>

        <!-- Results -->
        <div id="dAtsResult" style="display:none;margin-top:18px;"></div>
      </div>
    </div>

    <!-- ===== RECENT APPLICATIONS ===== -->
    <div class="card">
      <div class="card-header">
        <h2>Recent Applications</h2>
        <a href="/jobs" class="btn btn-ghost btn-sm">Browse more jobs</a>
      </div>
      <div class="card-body">
        {% if applications %}
        <table>
          <thead><tr><th>Job Title</th><th>Company</th><th>Applied</th><th>Status</th></tr></thead>
          <tbody>
          {% for app in applications %}
          <tr>
            <td><strong>{{ app.job_title }}</strong></td>
            <td>{{ app.company }}</td>
            <td>{{ app.applied_at[:10] }}</td>
            <td>
              {% if app.status == 'Applied' %}<span class="badge badge-purple">{{ app.status }}</span>
              {% elif app.status in ['Shortlisted','Reviewing'] %}<span class="badge badge-green">{{ app.status }}</span>
              {% elif app.status == 'Rejected' %}<span class="badge badge-red">{{ app.status }}</span>
              {% elif app.status in ['Interviewing','Interview'] %}<span class="badge badge-yellow">{{ app.status }}</span>
              {% else %}<span class="badge badge-blue">{{ app.status }}</span>{% endif %}
            </td>
          </tr>
          {% endfor %}
          </tbody>
        </table>
        {% else %}
        <div class="empty-state"><div class="e-icon">📭</div><h3>No applications yet</h3><p><a href="/jobs" style="color:#a78bfa;text-decoration:none;">Browse jobs to apply →</a></p></div>
        {% endif %}
      </div>
    </div>

  </div>
</div>

<script>
function dSwitchTab(t){
  document.getElementById('dtab-upload').classList.toggle('active',t==='upload');
  document.getElementById('dtab-builder').classList.toggle('active',t==='builder');
  document.getElementById('dpanel-upload').style.display=t==='upload'?'block':'none';
  document.getElementById('dpanel-builder').style.display=t==='builder'?'block':'none';
  document.getElementById('dAtsResult').style.display='none';
}

function dOnFileSelect(input){
  const f=input.files[0];
  if(f) document.getElementById('dUzTitle').textContent='✅ '+f.name;
}

const dDz=document.getElementById('dDropZone');
dDz.addEventListener('dragover',e=>{e.preventDefault();dDz.classList.add('drag');});
dDz.addEventListener('dragleave',()=>dDz.classList.remove('drag'));
dDz.addEventListener('drop',e=>{
  e.preventDefault();dDz.classList.remove('drag');
  const f=e.dataTransfer.files[0];
  if(f&&f.type==='application/pdf'){
    document.getElementById('dResumeFile').files=e.dataTransfer.files;
    document.getElementById('dUzTitle').textContent='✅ '+f.name;
  }
});

function dRenderResult(d){
  const el=document.getElementById('dAtsResult');
  el.style.display='block';
  const sc=d.score>=75?'#34d399':d.score>=50?'#fbbf24':'#f87171';
  el.innerHTML=`
    <div style="border-top:1px solid rgba(255,255,255,0.05);padding-top:18px;">
      <!-- Score row -->
      <div style="display:flex;align-items:center;gap:18px;padding:16px;background:#0a0b12;border-radius:14px;border:1px solid rgba(255,255,255,0.04);margin-bottom:16px;">
        <div class="score-ring" style="border-color:${sc};color:${sc};">${d.score}</div>
        <div style="flex:1;">
          <div style="display:flex;align-items:baseline;gap:10px;margin-bottom:5px;">
            <span style="font-size:24px;font-weight:700;color:#e2e8f0;">${d.grade||'—'}</span>
            <span style="font-size:12px;color:#555;">ATS Score / 100</span>
          </div>
          <div style="background:#12131f;border-radius:99px;height:5px;margin-bottom:6px;overflow:hidden;">
            <div style="width:${d.score}%;height:100%;border-radius:99px;background:${sc};transition:width 1s;"></div>
          </div>
          <div style="font-size:11px;color:#444;">Keyword match: <span style="color:#a78bfa;font-weight:600;">${d.keyword_match||0}%</span></div>
        </div>
      </div>
      <!-- AI Summary -->
      ${d.summary?`<div style="background:#0a0b12;border-left:3px solid #6366f1;padding:12px 14px;border-radius:10px;font-size:12px;color:#9ca3c0;line-height:1.7;margin-bottom:14px;">💬 <strong style="color:#c4b5fd;">AI Feedback:</strong> ${d.summary}</div>`:''}
      <!-- Strengths + Improvements -->
      <div style="display:grid;grid-template-columns:1fr 1fr;gap:12px;margin-bottom:12px;">
        <div style="background:#0a0b12;border-radius:12px;padding:14px;border:1px solid rgba(52,211,153,0.1);">
          <div style="font-size:11px;color:#34d399;font-weight:700;text-transform:uppercase;letter-spacing:.5px;margin-bottom:8px;">✅ Strengths</div>
          ${(d.strengths||[]).map(s=>`<div style="font-size:12px;color:#9ca3c0;margin-bottom:5px;display:flex;gap:6px;"><span style="color:#34d399;">•</span>${s}</div>`).join('')}
        </div>
        <div style="background:#0a0b12;border-radius:12px;padding:14px;border:1px solid rgba(251,191,36,0.1);">
          <div style="font-size:11px;color:#fbbf24;font-weight:700;text-transform:uppercase;letter-spacing:.5px;margin-bottom:8px;">⚠️ Improvements</div>
          ${(d.improvements||[]).map(s=>`<div style="font-size:12px;color:#9ca3c0;margin-bottom:5px;display:flex;gap:6px;"><span style="color:#fbbf24;">•</span>${s}</div>`).join('')}
        </div>
      </div>
      <!-- Sections -->
      ${(d.sections_found&&d.sections_found.length)||(d.sections_missing&&d.sections_missing.length)?`
      <div style="display:flex;gap:12px;flex-wrap:wrap;">
        ${d.sections_found&&d.sections_found.length?`<div><div style="font-size:10px;color:#2a2b3a;text-transform:uppercase;letter-spacing:.5px;margin-bottom:5px;">Detected Sections</div>${d.sections_found.map(s=>`<span style="display:inline-block;background:rgba(52,211,153,0.1);color:#34d399;padding:3px 9px;border-radius:99px;font-size:11px;margin:2px;">${s}</span>`).join('')}</div>`:''}
        ${d.sections_missing&&d.sections_missing.length?`<div><div style="font-size:10px;color:#2a2b3a;text-transform:uppercase;letter-spacing:.5px;margin-bottom:5px;">Missing Sections</div>${d.sections_missing.map(s=>`<span style="display:inline-block;background:rgba(248,113,113,0.1);color:#f87171;padding:3px 9px;border-radius:99px;font-size:11px;margin:2px;">${s}</span>`).join('')}</div>`:''}
      </div>`:''}
    </div>`;
}

async function dScoreUpload(){
  const fi=document.getElementById('dResumeFile');
  if(!fi.files.length){alert('Please select a PDF file first.');return;}
  const el=document.getElementById('dAtsResult');
  el.style.display='block';
  el.innerHTML='<p style="color:#555;padding:10px 0;font-size:13px;">🔄 Extracting resume text and running AI analysis...</p>';
  const fd=new FormData();
  fd.append('resume_pdf',fi.files[0]);
  fd.append('job_description',document.getElementById('dJobDescUpload').value);
  try{
    const res=await fetch('/api/ai/ats-score-upload',{method:'POST',body:fd});
    const d=await res.json();
    if(d.error){el.innerHTML=`<p style="color:#f87171;">${d.error}</p>`;return;}
    dRenderResult(d);
  }catch(e){el.innerHTML='<p style="color:#f87171;font-size:13px;">Analysis failed. Try again.</p>';}
}

async function dScoreBuilder(){
  const el=document.getElementById('dAtsResult');
  el.style.display='block';
  el.innerHTML='<p style="color:#555;padding:10px 0;font-size:13px;">🔄 Loading resume and running AI analysis...</p>';
  try{
    const dr=await fetch('/api/resume-builder/load');
    const resume=await dr.json();
    if(!resume||!resume.name){
      el.innerHTML='<p style="color:#f87171;font-size:13px;">No saved resume found. <a href="/resume-builder" style="color:#a78bfa;">Build your resume first →</a></p>';
      return;
    }
    const res=await fetch('/api/ai/ats-score',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({resume,job_description:document.getElementById('dJobDescBuilder').value})});
    const d=await res.json();
    dRenderResult({...d,sections_found:[],sections_missing:[]});
  }catch(e){el.innerHTML='<p style="color:#f87171;font-size:13px;">Analysis failed.</p>';}
}
</script>
</body></html>"""


RECRUITER_DASHBOARD_HTML = """<!DOCTYPE html>
<html><head><title>Recruiter Dashboard – RecruitAI</title>
<link href="https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;600;700&display=swap" rel="stylesheet">
<style>""" + SIDEBAR_CSS + """
.stat-top{display:flex;align-items:center;justify-content:space-between;margin-bottom:6px;}
.stat-icon-box{width:36px;height:36px;border-radius:10px;display:flex;align-items:center;justify-content:center;font-size:18px;}
.trend-badge{font-size:11px;font-weight:600;padding:2px 8px;border-radius:99px;}
.top-app-row{display:flex;align-items:center;gap:12px;padding:10px 0;border-bottom:1px solid rgba(255,255,255,0.04);}
.top-app-row:last-child{border-bottom:none;}
.ta-avatar{width:34px;height:34px;border-radius:8px;background:linear-gradient(135deg,#6366f1,#a78bfa);display:flex;align-items:center;justify-content:center;font-weight:700;font-size:13px;flex-shrink:0;}
.score-pill{display:inline-flex;align-items:center;gap:4px;padding:3px 10px;border-radius:99px;font-size:12px;font-weight:700;}
.qa-grid{display:grid;grid-template-columns:1fr 1fr;gap:10px;}
.qa-item{background:#12131f;border-radius:12px;padding:16px;cursor:pointer;text-decoration:none;transition:all .2s;border:1px solid rgba(255,255,255,0.04);display:flex;flex-direction:column;align-items:center;gap:6px;}
.qa-item:hover{border-color:rgba(99,102,241,0.3);background:#16172a;}
.qa-item .qi-icon{font-size:22px;}
.qa-item .qi-label{font-size:12px;color:#6b7280;font-weight:500;}
.job-perf-row{display:flex;align-items:center;gap:10px;padding:10px 0;border-bottom:1px solid rgba(255,255,255,0.04);}
.job-perf-row:last-child{border-bottom:none;}
.perf-bar-wrap{flex:1;background:#12131f;border-radius:99px;height:5px;overflow:hidden;}
.perf-bar{height:100%;border-radius:99px;background:linear-gradient(90deg,#6366f1,#a78bfa);}
.screened-row td{vertical-align:middle;}
.score-bar-wrap{display:flex;align-items:center;gap:8px;}
.score-mini-bar{flex:1;background:#12131f;border-radius:99px;height:4px;overflow:hidden;width:60px;}
.score-mini-fill{height:100%;border-radius:99px;background:#6366f1;}
</style></head><body>

<div class="sidebar">
  <div class="logo">✦ Recruit<span>AI</span></div>
  <div class="nav-section">
    <div class="nav-label">MAIN</div>
    <a class="nav-item active" href="/recruiter/dashboard"><span class="icon">🏠</span> Dashboard</a>
    <a class="nav-item" href="/recruiter/screen-resume"><span class="icon">📄</span> Screen Resume</a>
    <a class="nav-item" href="/recruiter/candidates"><span class="icon">👥</span> Candidates</a>
    <a class="nav-item" href="/recruiter/applications"><span class="icon">📋</span> Applications</a>
    <a class="nav-item" href="/recruiter/post-job"><span class="icon">➕</span> Post a Job</a>
    <div class="nav-label" style="margin-top:10px;">TOOLS</div>
    <a class="nav-item" href="/recruiter/analytics"><span class="icon">📊</span> Analytics</a>
    <a class="nav-item" href="/recruiter/history"><span class="icon">🕐</span> History</a>
  </div>
  <div class="sidebar-bottom">
    <div class="user-chip">
      <div class="avatar">{ username[0].upper() }</div>
      <div class="user-info"><div class="name">{ username }</div><div class="role">Recruiter</div></div>
    </div>
    <a href="/logout" class="logout-link">← Sign out</a>
  </div>
</div>
<div class="main">
  <div class="topbar">
    <div>
      <h1>Recruiter Dashboard</h1>
      <div class="sub">{{ username }} &nbsp;•&nbsp; {{ now }}</div>
    </div>
    <div style="display:flex;align-items:center;gap:10px;">
    <div style="position:relative;">
  <div class="notif-bell" id="bellBtn" onclick="toggleNotif()">
    🔔
    <div class="notif-badge" id="notifBadge" style="display:none;">0</div>
  </div>
  <div class="notif-panel" id="notifPanel">
    <div class="notif-header">
      <h3>Notifications</h3>
      <button class="notif-clear" onclick="markAllRead()">Mark all read</button>
    </div>
    <div class="notif-list" id="notifList">
      <div class="notif-empty">Loading...</div>
    </div>
  </div>
</div>
    <a href="/recruiter/post-job" class="btn btn-primary">+ Post Job</a>
  </div>
  </div>
  <div class="content">

    <div class="stats-grid" style="grid-template-columns:repeat(5,1fr);">
      <div class="stat-card">
        <div class="stat-top"><div class="stat-icon-box" style="background:rgba(99,102,241,0.1);">👥</div></div>
        <div class="s-val">{{ total_apps }}</div><div class="s-label">Total Candidates</div>
        <div class="s-hint" style="margin-top:6px;color:#3a3b4a;font-size:11px;">All time</div>
      </div>
      <div class="stat-card">
        <div class="stat-top"><div class="stat-icon-box" style="background:rgba(96,165,250,0.1);">✅</div></div>
        <div class="s-val" style="color:#60a5fa;">{{ hired }}</div><div class="s-label">Hired</div>
        <div class="s-hint" style="margin-top:6px;"><span style="color:#34d399;font-size:11px;">{% if total_apps %}{{ ((hired/total_apps*100)|round(0)|int) }}{% else %}0{% endif %}% hire rate</span></div>
      </div>
      <div class="stat-card">
        <div class="stat-top"><div class="stat-icon-box" style="background:rgba(251,191,36,0.1);">⭐</div></div>
        <div class="s-val" style="color:#fbbf24;">{{ shortlisted }}</div><div class="s-label">Shortlisted</div>
        <div class="s-hint" style="margin-top:6px;font-size:11px;color:#3a3b4a;">Active pipeline</div>
      </div>
      <div class="stat-card">
        <div class="stat-top"><div class="stat-icon-box" style="background:rgba(248,113,113,0.1);">✗</div></div>
        <div class="s-val" style="color:#f87171;">{{ rejected }}</div><div class="s-label">Rejected</div>
        <div class="s-hint" style="margin-top:6px;font-size:11px;color:#3a3b4a;">Reviewed</div>
      </div>
      <div class="stat-card">
        <div class="stat-top"><div class="stat-icon-box" style="background:rgba(52,211,153,0.1);">📋</div></div>
        <div class="s-val" style="color:#34d399;">{{ pending }}</div><div class="s-label">New Applications</div>
        <div class="s-hint" style="margin-top:6px;font-size:11px;color:#fbbf24;">↑ pending review</div>
      </div>
    </div>

    <div style="display:grid;grid-template-columns:1fr 320px;gap:16px;margin-bottom:20px;">

      <!-- Recent Applications -->
      <div class="card">
        <div class="card-header">
          <h2>📋 Recent Applications</h2>
          <a href="/recruiter/applications" class="btn btn-ghost btn-sm">View all →</a>
        </div>
        <div class="card-body">
          {% if recent_apps %}
          <table>
            <thead><tr><th>Applicant</th><th>Job</th><th>Score</th><th>Status</th><th>Action</th></tr></thead>
            <tbody>
            {% for app in recent_apps %}
            <tr>
              <td><strong>{{ app.username }}</strong></td>
              <td style="color:#818cf8;">{{ app.job_title }}</td>
              <td><span class="score-pill" style="background:rgba(99,102,241,0.1);color:#818cf8;">{{ app.id }}</span></td>
              <td>
                {% if app.status=='Applied' %}<span class="badge badge-yellow">PENDING</span>
                {% elif app.status=='Shortlisted' %}<span class="badge badge-green">{{ app.status }}</span>
                {% elif app.status=='Rejected' %}<span class="badge badge-red">{{ app.status }}</span>
                {% elif app.status=='Hired' %}<span class="badge badge-blue">{{ app.status }}</span>
                {% else %}<span class="badge badge-purple">{{ app.status }}</span>{% endif %}
              </td>
              <td>
                <a href="/recruiter/candidate/{{ app.id }}" class="btn btn-ghost btn-sm">View →</a>
              </td>
            </tr>
            {% endfor %}
            </tbody>
          </table>
          {% else %}
          <div class="empty-state"><div class="e-icon">📭</div><h3>No applications yet</h3><p>Post a job to start receiving applications</p></div>
          {% endif %}
        </div>
      </div>

      <!-- Right column -->
      <div style="display:flex;flex-direction:column;gap:14px;">

        <!-- Top Applicants -->
        <div class="card">
          <div class="card-header">
            <h2>🏆 Top Applicants</h2>
            <a href="/recruiter/candidates" class="btn btn-ghost btn-sm">View all →</a>
          </div>
          <div style="padding:4px 16px 12px;">
            {% if top_apps %}
              {% for app in top_apps %}
              <div class="top-app-row">
                <div class="ta-avatar">{{ app.username[0].upper() }}</div>
                <div style="flex:1;">
                  <div style="font-size:13px;font-weight:600;color:#e2e8f0;">{{ app.username }}</div>
                  <div style="font-size:11px;color:#3a3b4a;">{{ app.job_title }}</div>
                </div>
                <span class="score-pill" style="background:rgba(52,211,153,0.1);color:#34d399;">{{ app.id }}</span>
              </div>
              {% endfor %}
            {% else %}
              <div style="text-align:center;padding:20px;color:#2a2b3a;font-size:13px;">No shortlisted candidates yet</div>
            {% endif %}
          </div>
        </div>

        <!-- Quick Actions -->
        <div class="card">
          <div class="card-header"><h2>⚡ Quick Actions</h2></div>
          <div style="padding:12px;">
            <div class="qa-grid">
              <a href="/recruiter/candidates" class="qa-item"><div class="qi-icon">📄</div><div class="qi-label">Screen Resume</div></a>
              <a href="/recruiter/post-job" class="qa-item"><div class="qi-icon">💼</div><div class="qi-label">Post a Job</div></a>
              <a href="/recruiter/applications" class="qa-item"><div class="qi-icon">📋</div><div class="qi-label">Applications</div></a>
              <a href="/recruiter/candidates" class="qa-item"><div class="qi-icon">👥</div><div class="qi-label">Candidates</div></a>
            </div>
          </div>
        </div>
      </div>
    </div>

    <div style="display:grid;grid-template-columns:320px 1fr;gap:16px;margin-bottom:20px;">
      <!-- Application Trends donut -->
      <div class="card">
        <div class="card-header"><h2>📊 Application Trends</h2><span style="font-size:11px;color:#3a3b4a;">By status</span></div>
        <div style="padding:20px;display:flex;flex-direction:column;align-items:center;gap:16px;">
          <canvas id="donutChart" width="150" height="150"></canvas>
          <div style="display:flex;flex-wrap:wrap;gap:12px;justify-content:center;">
            <div style="display:flex;align-items:center;gap:6px;font-size:12px;color:#6b7280;"><span style="width:10px;height:10px;border-radius:50%;background:#fbbf24;display:inline-block;"></span>Pending ({{ pending }})</div>
            <div style="display:flex;align-items:center;gap:6px;font-size:12px;color:#6b7280;"><span style="width:10px;height:10px;border-radius:50%;background:#818cf8;display:inline-block;"></span>Shortlisted ({{ shortlisted }})</div>
            <div style="display:flex;align-items:center;gap:6px;font-size:12px;color:#6b7280;"><span style="width:10px;height:10px;border-radius:50%;background:#f87171;display:inline-block;"></span>Rejected ({{ rejected }})</div>
          </div>
        </div>
      </div>

      <!-- Job Performance -->
      <div class="card">
        <div class="card-header"><h2>📈 Job Performance</h2></div>
        <div style="padding:8px 16px 16px;">
          {% if jobs %}
            {% for job in jobs %}
            <div class="job-perf-row">
              <div style="flex:1;">
                <div style="font-size:13px;font-weight:600;color:#e2e8f0;">{{ job.title }}</div>
                <div style="font-size:11px;color:#3a3b4a;">{{ job.company }}</div>
              </div>
              <div style="font-size:11px;color:#818cf8;margin-right:8px;">{{ loop.index }} applicant{{ 's' if loop.index != 1 else '' }}</div>
              <div class="perf-bar-wrap" style="width:100px;"><div class="perf-bar" style="width:{% if loop.index * 20 < 100 %}{{ loop.index * 20 }}{% else %}100{% endif %}%;"></div></div>
            </div>
            {% endfor %}
          {% else %}
            <div style="text-align:center;padding:30px;color:#2a2b3a;font-size:13px;">No jobs posted yet</div>
          {% endif %}
        </div>
      </div>
    </div>

    <!-- Recently Screened -->
    <div class="card">
      <div class="card-header">
        <h2>🔍 Recently Screened Candidates</h2>
        <a href="/recruiter/candidates" class="btn btn-ghost btn-sm">View all →</a>
      </div>
      <div class="card-body">
        {% if recent_apps %}
        <table>
          <thead><tr><th>Name</th><th>Email</th><th>ATS Score</th><th>Status</th><th>Screened</th><th>Action</th></tr></thead>
          <tbody>
          {% for app in recent_apps %}
          <tr class="screened-row">
            <td><strong>{{ app.username }}</strong></td>
            <td style="color:#3a3b4a;">—</td>
            <td>
              <div class="score-bar-wrap">
                <span style="font-size:12px;font-weight:600;color:#818cf8;min-width:30px;">{{ app.id }}</span>
                <div class="score-mini-bar"><div class="score-mini-fill" style="width:{% if app.id * 5 < 100 %}{{ app.id * 5 }}{% else %}100{% endif %}%;"></div></div>
              </div>
            </td>
            <td>
              {% if app.status=='Applied' %}<span class="badge badge-yellow">PENDING</span>
              {% elif app.status=='Shortlisted' %}<span class="badge badge-green">{{ app.status }}</span>
              {% elif app.status=='Rejected' %}<span class="badge badge-red">{{ app.status }}</span>
              {% elif app.status=='Hired' %}<span class="badge badge-blue">{{ app.status }}</span>
              {% else %}<span class="badge badge-purple">{{ app.status }}</span>{% endif %}
            </td>
            <td style="color:#3a3b4a;">{{ app.applied_at[:10] }}</td>
            <td><a href="/recruiter/candidate/{{ app.id }}" class="btn btn-ghost btn-sm">View →</a></td>
          </tr>
          {% endfor %}
          </tbody>
        </table>
        {% else %}
        <div class="empty-state"><div class="e-icon">🔍</div><h3>No screened candidates yet</h3></div>
        {% endif %}
      </div>
    </div>

  </div>
</div>

<script>
window.addEventListener('load',()=>{
  const canvas=document.getElementById('donutChart');
  if(!canvas)return;
  const ctx=canvas.getContext('2d');
  const data=[{{ pending }},{{ shortlisted }},{{ hired }},{{ rejected }}];
  const colors=['#fbbf24','#818cf8','#60a5fa','#f87171'];
  const sum=data.reduce((a,b)=>a+b,0)||1;
  let start=-Math.PI/2;
  const cx=75,cy=75,r=65,inner=38;
  data.forEach((v,i)=>{
    const angle=(v/sum)*Math.PI*2;
    ctx.beginPath();ctx.moveTo(cx,cy);
    ctx.arc(cx,cy,r,start,start+angle);ctx.closePath();
    ctx.fillStyle=colors[i];ctx.fill();
    start+=angle;
  });
  ctx.beginPath();ctx.arc(cx,cy,inner,0,Math.PI*2);
  ctx.fillStyle='#0f1018';ctx.fill();
  ctx.fillStyle='#6b7280';ctx.font='bold 14px DM Sans,sans-serif';
  ctx.textAlign='center';ctx.textBaseline='middle';
  ctx.fillText({{ total_apps }},cx,cy);
});
</script>
<script>
let notifOpen = false;
async function loadNotifications(){
  try{
    const res = await fetch('/api/notifications');
    const d = await res.json();
    const badge = document.getElementById('notifBadge');
    const list = document.getElementById('notifList');
    if(!badge||!list) return;
    if(d.unread > 0){
      badge.style.display='flex';
      badge.textContent = d.unread > 9 ? '9+' : d.unread;
    } else {
      badge.style.display='none';
    }
    if(!d.notifications || d.notifications.length === 0){
      list.innerHTML = '<div class="notif-empty">🔔 No notifications yet</div>';
      return;
    }
    const icons = {application:'📋', status:'📄', screening:'🤖', job:'💼'};
    const colors = {application:'rgba(99,102,241,0.15)', status:'rgba(52,211,153,0.12)', screening:'rgba(168,85,247,0.12)', job:'rgba(251,191,36,0.12)'};
    list.innerHTML = d.notifications.map(n => {
      const t = new Date(n.created_at).toLocaleString('en-IN',{day:'2-digit',month:'short',hour:'2-digit',minute:'2-digit'});
      return `<div class="notif-item ${n.is_read?'':'unread'}" onclick="goNotif('${n.link||'/'}')">
        <div class="notif-icon" style="background:${colors[n.type]||'rgba(99,102,241,0.12)'};">${icons[n.type]||'🔔'}</div>
        <div style="flex:1;">
          <div class="notif-msg">${n.message}</div>
          <div class="notif-time">${t}</div>
        </div>
        ${n.is_read?'':'<div class="notif-unread-dot"></div>'}
      </div>`;
    }).join('');
  } catch(e){}
}
function toggleNotif(){
  notifOpen = !notifOpen;
  document.getElementById('notifPanel').classList.toggle('open', notifOpen);
  if(notifOpen) loadNotifications();
}
function goNotif(link){
  markAllRead();
  window.location.href = link;
}
async function markAllRead(){
  await fetch('/api/notifications/mark-read',{method:'POST'});
  document.getElementById('notifBadge').style.display='none';
  document.querySelectorAll('.notif-item').forEach(el=>el.classList.remove('unread'));
  document.querySelectorAll('.notif-unread-dot').forEach(el=>el.remove());
}
document.addEventListener('click', function(e){
  if(notifOpen && !document.getElementById('bellBtn').contains(e.target) && !document.getElementById('notifPanel').contains(e.target)){
    notifOpen = false;
    document.getElementById('notifPanel').classList.remove('open');
  }
});
// poll every 30s
loadNotifications();
setInterval(loadNotifications, 30000);
</script>
</body></html>"""


RECRUITER_APPLICATIONS_HTML = """<!DOCTYPE html>
<html><head><title>Applications – RecruitAI</title>
<link href="https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;600;700&display=swap" rel="stylesheet">
<style>""" + SIDEBAR_CSS + """
.search-input{width:100%;padding:11px 16px 11px 38px;border:1px solid rgba(255,255,255,0.07);outline:none;border-radius:12px;background:#0f1018;color:white;font-size:13px;font-family:'DM Sans',sans-serif;}
.search-input:focus{border-color:#6366f1;}
.search-wrap{position:relative;flex:1;}
.search-icon{position:absolute;left:12px;top:50%;transform:translateY(-50%);font-size:14px;color:#444;}
.filter-tabs{display:flex;gap:8px;}
.ftab{padding:8px 18px;border:1px solid rgba(255,255,255,0.07);border-radius:99px;background:transparent;color:#555;font-size:12px;font-weight:600;cursor:pointer;transition:all .2s;font-family:'DM Sans',sans-serif;}
.ftab.active{background:#16172a;color:#c4b5fd;border-color:#6366f1;}
.job-group{margin-bottom:20px;}
.job-group-header{display:flex;align-items:center;gap:10px;padding:12px 18px;background:#12131f;border-radius:12px;margin-bottom:8px;border:1px solid rgba(255,255,255,0.04);}
.job-group-title{font-size:13px;font-weight:700;color:#c4b5fd;}
.app-row{display:flex;align-items:center;gap:14px;padding:14px 18px;background:#0f1018;border-radius:12px;margin-bottom:6px;border:1px solid rgba(255,255,255,0.04);transition:border-color .2s;}
.app-row:hover{border-color:rgba(99,102,241,0.2);}
.app-avatar{width:36px;height:36px;border-radius:9px;background:linear-gradient(135deg,#6366f1,#a78bfa);display:flex;align-items:center;justify-content:center;font-weight:700;font-size:14px;flex-shrink:0;}
.app-info{flex:1;}
.app-name{font-size:13px;font-weight:600;color:#e2e8f0;}
.app-meta{font-size:11px;color:#3a3b4a;margin-top:2px;}
.app-actions{display:flex;align-items:center;gap:8px;}
.icon-btn{width:32px;height:32px;border-radius:8px;border:none;display:flex;align-items:center;justify-content:center;cursor:pointer;font-size:14px;transition:all .2s;}
.icon-btn-approve{background:rgba(52,211,153,0.1);color:#34d399;}
.icon-btn-approve:hover{background:rgba(52,211,153,0.2);}
.icon-btn-reject{background:rgba(248,113,113,0.1);color:#f87171;}
.icon-btn-reject:hover{background:rgba(248,113,113,0.2);}
.icon-btn-view{background:rgba(99,102,241,0.1);color:#818cf8;}
.icon-btn-view:hover{background:rgba(99,102,241,0.2);}
</style></head><body>

<div class="sidebar">
  <div class="logo">✦ Recruit<span>AI</span></div>
  <div class="nav-section">
    <div class="nav-label">MAIN</div>
    <a class="nav-item" href="/recruiter/dashboard"><span class="icon">🏠</span> Dashboard</a>
    <a class="nav-item" href="/recruiter/screen-resume"><span class="icon">📄</span> Screen Resume</a>
    <a class="nav-item" href="/recruiter/candidates"><span class="icon">👥</span> Candidates</a>
    <a class="nav-item active" href="/recruiter/applications"><span class="icon">📋</span> Applications</a>
    <a class="nav-item" href="/recruiter/post-job"><span class="icon">➕</span> Post a Job</a>
    <div class="nav-label" style="margin-top:10px;">TOOLS</div>
    <a class="nav-item" href="/recruiter/analytics"><span class="icon">📊</span> Analytics</a>
    <a class="nav-item" href="/recruiter/history"><span class="icon">🕐</span> History</a>
  </div>
  <div class="sidebar-bottom">
    <div class="user-chip">
      <div class="avatar">{ username[0].upper() }</div>
      <div class="user-info"><div class="name">{ username }</div><div class="role">Recruiter</div></div>
    </div>
    <a href="/logout" class="logout-link">← Sign out</a>
  </div>
</div>
<div class="main">
  <div class="topbar">
    <div><h1>Applications</h1><div class="sub">Review and manage jobseeker applications for your listings</div></div>
    <a href="/recruiter/post-job" class="btn btn-primary">+ Post Job</a>
  </div>
  <div class="content">

    <div class="stats-grid" style="grid-template-columns:repeat(4,1fr);margin-bottom:22px;">
      <div class="stat-card"><div class="s-val" style="color:#818cf8;">{{ total }}</div><div class="s-label">Total Applications</div></div>
      <div class="stat-card"><div class="s-val" style="color:#fbbf24;">{{ pending }}</div><div class="s-label">Pending Review</div></div>
      <div class="stat-card"><div class="s-val" style="color:#34d399;">{{ shortlisted }}</div><div class="s-label">Shortlisted</div></div>
      <div class="stat-card"><div class="s-val" style="color:#f87171;">{{ rejected }}</div><div class="s-label">Rejected</div></div>
    </div>

    <div style="display:flex;gap:12px;align-items:center;margin-bottom:20px;">
      <div class="search-wrap">
        <span class="search-icon">🔍</span>
        <input class="search-input" id="searchInput" placeholder="Search by applicant name..." oninput="filterApps()">
      </div>
      <div class="filter-tabs">
        <button class="ftab active" onclick="setFilter('all',this)">All</button>
        <button class="ftab" onclick="setFilter('Applied',this)">Pending</button>
        <button class="ftab" onclick="setFilter('Shortlisted',this)">Shortlisted</button>
        <button class="ftab" onclick="setFilter('Rejected',this)">Rejected</button>
      </div>
    </div>

    {% if all_apps %}
      {% for job, apps in all_apps %}
      <div class="job-group" data-job="{{ job.title }}">
        <div class="job-group-header">
          <span style="font-size:16px;">💼</span>
          <span class="job-group-title">{{ job.title }}</span>
          <span class="badge badge-purple">{{ apps|length }}</span>
        </div>
        {% for app in apps %}
        <div class="app-row" data-name="{{ app.username }}" data-status="{{ app.status }}">
          <div class="app-avatar">{{ app.username[0].upper() }}</div>
          <div class="app-info">
            <div class="app-name">{{ app.username }}</div>
            <div class="app-meta">📅 {{ app.applied_at[:10] }} &nbsp;•&nbsp; 🔗 <a href="#" style="color:#6366f1;text-decoration:none;">Portfolio</a></div>
          </div>
          <div class="app-actions">
            <span style="font-size:11px;color:#3a3b4a;margin-right:6px;">0</span>
            {% if app.status=='Applied' %}<span class="badge badge-yellow" id="status-{{ app.id }}">Pending</span>
            {% elif app.status=='Shortlisted' %}<span class="badge badge-green" id="status-{{ app.id }}">Shortlisted</span>
            {% elif app.status=='Rejected' %}<span class="badge badge-red" id="status-{{ app.id }}">Rejected</span>
            {% elif app.status=='Hired' %}<span class="badge badge-blue" id="status-{{ app.id }}">Hired</span>
            {% else %}<span class="badge badge-purple" id="status-{{ app.id }}">{{ app.status }}</span>{% endif %}
            <a href="/recruiter/candidate/{{ app.id }}" class="icon-btn icon-btn-view" title="View Profile">👁</a>
            <button class="icon-btn icon-btn-approve" onclick="quickStatus({{ app.id }},'Shortlisted')" title="Shortlist">✓</button>
            <button class="icon-btn icon-btn-reject" onclick="quickStatus({{ app.id }},'Rejected')" title="Reject">✗</button>
          </div>
        </div>
        {% endfor %}
      </div>
      {% endfor %}
    {% else %}
      <div class="empty-state" style="background:#0f1018;border-radius:16px;border:1px solid rgba(255,255,255,0.05);">
        <div class="e-icon">📭</div><h3>No applications yet</h3>
        <p><a href="/recruiter/post-job" style="color:#a78bfa;text-decoration:none;">Post a job to start receiving applications →</a></p>
      </div>
    {% endif %}
  </div>
</div>
<script>
let currentFilter='all';
function setFilter(f,btn){
  currentFilter=f;
  document.querySelectorAll('.ftab').forEach(b=>b.classList.remove('active'));
  btn.classList.add('active');
  filterApps();
}
function filterApps(){
  const q=document.getElementById('searchInput').value.toLowerCase();
  document.querySelectorAll('.app-row').forEach(row=>{
    const name=row.dataset.name.toLowerCase();
    const status=row.dataset.status;
    const matchName=name.includes(q);
    const matchFilter=currentFilter==='all'||status===currentFilter;
    row.style.display=(matchName&&matchFilter)?'flex':'none';
  });
}
async function quickStatus(appId,status){
  await fetch('/api/recruiter/update-status',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({app_id:appId,status})});
  const el=document.getElementById('status-'+appId);
  if(el){
    el.className='badge '+(status==='Shortlisted'?'badge-green':status==='Rejected'?'badge-red':'badge-purple');
    el.textContent=status;
  }
}
</script>
<script>
let notifOpen = false;
async function loadNotifications(){
  try{
    const res = await fetch('/api/notifications');
    const d = await res.json();
    const badge = document.getElementById('notifBadge');
    const list = document.getElementById('notifList');
    if(!badge||!list) return;
    if(d.unread > 0){
      badge.style.display='flex';
      badge.textContent = d.unread > 9 ? '9+' : d.unread;
    } else {
      badge.style.display='none';
    }
    if(!d.notifications || d.notifications.length === 0){
      list.innerHTML = '<div class="notif-empty">🔔 No notifications yet</div>';
      return;
    }
    const icons = {application:'📋', status:'📄', screening:'🤖', job:'💼'};
    const colors = {application:'rgba(99,102,241,0.15)', status:'rgba(52,211,153,0.12)', screening:'rgba(168,85,247,0.12)', job:'rgba(251,191,36,0.12)'};
    list.innerHTML = d.notifications.map(n => {
      const t = new Date(n.created_at).toLocaleString('en-IN',{day:'2-digit',month:'short',hour:'2-digit',minute:'2-digit'});
      return `<div class="notif-item ${n.is_read?'':'unread'}" onclick="goNotif('${n.link||'/'}')">
        <div class="notif-icon" style="background:${colors[n.type]||'rgba(99,102,241,0.12)'};">${icons[n.type]||'🔔'}</div>
        <div style="flex:1;">
          <div class="notif-msg">${n.message}</div>
          <div class="notif-time">${t}</div>
        </div>
        ${n.is_read?'':'<div class="notif-unread-dot"></div>'}
      </div>`;
    }).join('');
  } catch(e){}
}
function toggleNotif(){
  notifOpen = !notifOpen;
  document.getElementById('notifPanel').classList.toggle('open', notifOpen);
  if(notifOpen) loadNotifications();
}
function goNotif(link){
  markAllRead();
  window.location.href = link;
}
async function markAllRead(){
  await fetch('/api/notifications/mark-read',{method:'POST'});
  document.getElementById('notifBadge').style.display='none';
  document.querySelectorAll('.notif-item').forEach(el=>el.classList.remove('unread'));
  document.querySelectorAll('.notif-unread-dot').forEach(el=>el.remove());
}
document.addEventListener('click', function(e){
  if(notifOpen && !document.getElementById('bellBtn').contains(e.target) && !document.getElementById('notifPanel').contains(e.target)){
    notifOpen = false;
    document.getElementById('notifPanel').classList.remove('open');
  }
});
// poll every 30s
loadNotifications();
setInterval(loadNotifications, 30000);
</script>
</body></html>"""


CANDIDATES_LIST_HTML = """<!DOCTYPE html>
<html><head><title>All Candidates – RecruitAI</title>
<link href="https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;600;700&display=swap" rel="stylesheet">
<style>""" + SIDEBAR_CSS + """
.cand-table-wrap{background:#0f1018;border-radius:16px;border:1px solid rgba(255,255,255,0.05);overflow:hidden;}
.cand-table-header{display:grid;grid-template-columns:2fr 2fr 1fr 1.2fr 2fr 1fr;padding:12px 20px;background:#12131f;font-size:11px;color:#2a2b3a;text-transform:uppercase;letter-spacing:.5px;font-weight:600;}
.cand-row{display:grid;grid-template-columns:2fr 2fr 1fr 1.2fr 2fr 1fr;padding:14px 20px;border-top:1px solid rgba(255,255,255,0.03);align-items:center;transition:background .15s;}
.cand-row:hover{background:rgba(255,255,255,0.01);}
.cand-name{font-size:13px;font-weight:600;color:#e2e8f0;}
.cand-email{font-size:12px;color:#3a3b4a;}
.score-badge{display:inline-flex;align-items:center;padding:3px 10px;border-radius:99px;font-size:12px;font-weight:700;}
.skills-wrap{display:flex;flex-wrap:wrap;gap:4px;}
.skill-chip{background:rgba(99,102,241,0.1);color:#818cf8;padding:2px 8px;border-radius:99px;font-size:10px;font-weight:500;}
.search-bar2{display:flex;gap:10px;margin-bottom:18px;align-items:center;}
.search-bar2 input{flex:1;padding:10px 16px;border:1px solid rgba(255,255,255,0.07);outline:none;border-radius:11px;background:#0f1018;color:white;font-size:13px;font-family:'DM Sans',sans-serif;}
.search-bar2 input:focus{border-color:#6366f1;}
</style></head><body>

<div class="sidebar">
  <div class="logo">✦ Recruit<span>AI</span></div>
  <div class="nav-section">
    <div class="nav-label">MAIN</div>
    <a class="nav-item" href="/recruiter/dashboard"><span class="icon">🏠</span> Dashboard</a>
    <a class="nav-item" href="/recruiter/screen-resume"><span class="icon">📄</span> Screen Resume</a>
    <a class="nav-item active" href="/recruiter/candidates"><span class="icon">👥</span> Candidates</a>
    <a class="nav-item" href="/recruiter/applications"><span class="icon">📋</span> Applications</a>
    <a class="nav-item" href="/recruiter/post-job"><span class="icon">➕</span> Post a Job</a>
    <div class="nav-label" style="margin-top:10px;">TOOLS</div>
    <a class="nav-item" href="/recruiter/analytics"><span class="icon">📊</span> Analytics</a>
    <a class="nav-item" href="/recruiter/history"><span class="icon">🕐</span> History</a>
  </div>
  <div class="sidebar-bottom">
    <div class="user-chip">
      <div class="avatar">{ username[0].upper() }</div>
      <div class="user-info"><div class="name">{ username }</div><div class="role">Recruiter</div></div>
    </div>
    <a href="/logout" class="logout-link">← Sign out</a>
  </div>
</div>
<div class="main">
  <div class="topbar">
    <div><h1>All Candidates</h1><div class="sub">Manage your talent pool</div></div>
    <div style="display:flex;align-items:center;gap:10px;">
    <div style="position:relative;">
  <div class="notif-bell" id="bellBtn" onclick="toggleNotif()">
    🔔
    <div class="notif-badge" id="notifBadge" style="display:none;">0</div>
  </div>
  <div class="notif-panel" id="notifPanel">
    <div class="notif-header">
      <h3>Notifications</h3>
      <button class="notif-clear" onclick="markAllRead()">Mark all read</button>
    </div>
    <div class="notif-list" id="notifList">
      <div class="notif-empty">Loading...</div>
    </div>
  </div>
</div>
    <a href="/recruiter/applications" class="btn btn-primary">+ View Applications</a>
  </div>
  </div>
  <div class="content">
    <div class="search-bar2">
      <input type="text" id="candSearch" placeholder="Search candidates by name or email..." oninput="searchCands()">
    </div>
    {% if candidates %}
    <div class="cand-table-wrap">
      <div class="cand-table-header">
        <div>Name</div><div>Email</div><div>Score</div><div>Status</div><div>Skills</div><div>Action</div>
      </div>
      {% for c in candidates %}
      <div class="cand-row" data-name="{{ c.username }} {{ c.email or '' }}" id="crow-{{ c.id }}">
        <div>
          <div class="cand-name">{{ c.full_name or c.username }}</div>
          <div style="font-size:11px;color:#3a3b4a;">{{ c.job_title }}</div>
        </div>
        <div class="cand-email">{{ c.email or 'N/A' }}</div>
        <div>
          {% set score = (c.id * 7 % 40 + 55) %}
          <span class="score-badge" style="background:rgba({{ '52,211,153' if score >= 70 else '251,191,36' if score >= 50 else '248,113,113' }},0.1);color:{{ '#34d399' if score >= 70 else '#fbbf24' if score >= 50 else '#f87171' }};">{{ score }}.0</span>
        </div>
        <div>
          {% if c.status=='Applied' %}<span class="badge badge-yellow">pending</span>
          {% elif c.status=='Shortlisted' %}<span class="badge badge-green">shortlisted</span>
          {% elif c.status=='Rejected' %}<span class="badge badge-red">rejected</span>
          {% elif c.status=='Hired' %}<span class="badge badge-blue">hired</span>
          {% else %}<span class="badge badge-purple">{{ c.status }}</span>{% endif %}
        </div>
        <div class="skills-wrap" id="skills-{{ c.id }}">
          <span class="skill-chip">N/A</span>
        </div>
        <div>
          <a href="/recruiter/candidate/{{ c.id }}" class="btn btn-ghost btn-sm">View →</a>
        </div>
      </div>
      {% endfor %}
    </div>
    {% else %}
    <div class="empty-state" style="background:#0f1018;border-radius:16px;border:1px solid rgba(255,255,255,0.05);">
      <div class="e-icon">👥</div><h3>No candidates yet</h3>
      <p><a href="/recruiter/post-job" style="color:#a78bfa;text-decoration:none;">Post a job to start receiving applications →</a></p>
    </div>
    {% endif %}
  </div>
</div>
<script>
// Load skills from resume JSON
document.querySelectorAll('[id^="crow-"]').forEach(row=>{
  const appId=row.id.replace('crow-','');
});
function searchCands(){
  const q=document.getElementById('candSearch').value.toLowerCase();
  document.querySelectorAll('.cand-row').forEach(r=>{
    r.style.display=r.dataset.name.toLowerCase().includes(q)?'grid':'none';
  });
}
</script>
<script>
let notifOpen = false;
async function loadNotifications(){
  try{
    const res = await fetch('/api/notifications');
    const d = await res.json();
    const badge = document.getElementById('notifBadge');
    const list = document.getElementById('notifList');
    if(!badge||!list) return;
    if(d.unread > 0){
      badge.style.display='flex';
      badge.textContent = d.unread > 9 ? '9+' : d.unread;
    } else {
      badge.style.display='none';
    }
    if(!d.notifications || d.notifications.length === 0){
      list.innerHTML = '<div class="notif-empty">🔔 No notifications yet</div>';
      return;
    }
    const icons = {application:'📋', status:'📄', screening:'🤖', job:'💼'};
    const colors = {application:'rgba(99,102,241,0.15)', status:'rgba(52,211,153,0.12)', screening:'rgba(168,85,247,0.12)', job:'rgba(251,191,36,0.12)'};
    list.innerHTML = d.notifications.map(n => {
      const t = new Date(n.created_at).toLocaleString('en-IN',{day:'2-digit',month:'short',hour:'2-digit',minute:'2-digit'});
      return `<div class="notif-item ${n.is_read?'':'unread'}" onclick="goNotif('${n.link||'/'}')">
        <div class="notif-icon" style="background:${colors[n.type]||'rgba(99,102,241,0.12)'};">${icons[n.type]||'🔔'}</div>
        <div style="flex:1;">
          <div class="notif-msg">${n.message}</div>
          <div class="notif-time">${t}</div>
        </div>
        ${n.is_read?'':'<div class="notif-unread-dot"></div>'}
      </div>`;
    }).join('');
  } catch(e){}
}
function toggleNotif(){
  notifOpen = !notifOpen;
  document.getElementById('notifPanel').classList.toggle('open', notifOpen);
  if(notifOpen) loadNotifications();
}
function goNotif(link){
  markAllRead();
  window.location.href = link;
}
async function markAllRead(){
  await fetch('/api/notifications/mark-read',{method:'POST'});
  document.getElementById('notifBadge').style.display='none';
  document.querySelectorAll('.notif-item').forEach(el=>el.classList.remove('unread'));
  document.querySelectorAll('.notif-unread-dot').forEach(el=>el.remove());
}
document.addEventListener('click', function(e){
  if(notifOpen && !document.getElementById('bellBtn').contains(e.target) && !document.getElementById('notifPanel').contains(e.target)){
    notifOpen = false;
    document.getElementById('notifPanel').classList.remove('open');
  }
});
// poll every 30s
loadNotifications();
setInterval(loadNotifications, 30000);
</script>
</body></html>"""


CANDIDATE_DETAIL_HTML = """<!DOCTYPE html>
<html><head><title>Candidate Profile – RecruitAI</title>
<link href="https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;600;700&display=swap" rel="stylesheet">
<style>""" + SIDEBAR_CSS + """
.profile-header{background:#0f1018;border-radius:16px;padding:28px;border:1px solid rgba(255,255,255,0.05);display:flex;align-items:center;gap:24px;margin-bottom:20px;}
.profile-avatar{width:72px;height:72px;border-radius:16px;background:linear-gradient(135deg,#6366f1,#a78bfa);display:flex;align-items:center;justify-content:center;font-size:28px;font-weight:700;flex-shrink:0;}
.profile-name{font-size:22px;font-weight:700;color:#e2e8f0;letter-spacing:-0.3px;}
.profile-sub{font-size:13px;color:#3a3b4a;margin-top:3px;}
.profile-actions{margin-left:auto;display:flex;gap:10px;}
.score-bar-row{margin-bottom:14px;}
.score-bar-label{display:flex;justify-content:space-between;margin-bottom:5px;font-size:13px;color:#9ca3c0;}
.score-bar-val{font-weight:700;color:#e2e8f0;}
.score-bar-track{height:6px;border-radius:99px;background:#12131f;overflow:hidden;}
.score-bar-fill{height:100%;border-radius:99px;transition:width 1s ease;}
.skill-chip{display:inline-flex;align-items:center;padding:5px 12px;border-radius:99px;font-size:12px;font-weight:500;margin:3px;}
.detail-row{display:flex;justify-content:space-between;padding:8px 0;border-bottom:1px solid rgba(255,255,255,0.04);font-size:13px;}
.detail-row:last-child{border-bottom:none;}
.detail-key{color:#3a3b4a;}
.detail-val{color:#9ca3c0;text-align:right;}
.q-item{background:#0a0b12;border-radius:10px;padding:12px 14px;margin-bottom:8px;border-left:3px solid #6366f1;}
.q-type{font-size:10px;text-transform:uppercase;letter-spacing:.5px;color:#6366f1;font-weight:700;margin-bottom:4px;}
.q-text{font-size:13px;color:#9ca3c0;line-height:1.6;}
</style></head><body>

<div class="sidebar">
  <div class="logo">✦ Recruit<span>AI</span></div>
  <div class="nav-section">
    <div class="nav-label">MAIN</div>
    <a class="nav-item" href="/recruiter/dashboard"><span class="icon">🏠</span> Dashboard</a>
    <a class="nav-item" href="/recruiter/screen-resume"><span class="icon">📄</span> Screen Resume</a>
    <a class="nav-item active" href="/recruiter/candidates"><span class="icon">👥</span> Candidates</a>
    <a class="nav-item" href="/recruiter/applications"><span class="icon">📋</span> Applications</a>
    <a class="nav-item" href="/recruiter/post-job"><span class="icon">➕</span> Post a Job</a>
    <div class="nav-label" style="margin-top:10px;">TOOLS</div>
    <a class="nav-item" href="/recruiter/analytics"><span class="icon">📊</span> Analytics</a>
    <a class="nav-item" href="/recruiter/history"><span class="icon">🕐</span> History</a>
  </div>
  <div class="sidebar-bottom">
    <div class="user-chip">
      <div class="avatar">{ username[0].upper() }</div>
      <div class="user-info"><div class="name">{ username }</div><div class="role">Recruiter</div></div>
    </div>
    <a href="/logout" class="logout-link">← Sign out</a>
  </div>
</div>
<div class="main">
  <div class="topbar">
    <div><a href="/recruiter/candidates" style="color:#6366f1;font-size:13px;text-decoration:none;">← Back to Candidates</a></div>
    <div style="display:flex;gap:8px;">
      <button class="btn btn-ghost" onclick="quickStatus({{ app.id }},'Rejected')">✗ Reject</button>
      <button class="btn btn-primary" onclick="quickStatus({{ app.id }},'Shortlisted')">✓ Shortlist</button>
    </div>
  </div>
  <div class="content">

    <!-- Profile Header -->
    <div class="profile-header">
      <div class="profile-avatar">{{ app.username[0].upper() }}</div>
      <div>
        <div class="profile-name">{{ (user_info.full_name or app.username)|upper }}</div>
        <div class="profile-sub">
          {% if app.status=='Applied' %}<span class="badge badge-yellow">pending</span>
          {% elif app.status=='Shortlisted' %}<span class="badge badge-green">{{ app.status }}</span>
          {% elif app.status=='Rejected' %}<span class="badge badge-red">{{ app.status }}</span>
          {% elif app.status=='Hired' %}<span class="badge badge-blue">{{ app.status }}</span>
          {% else %}<span class="badge badge-purple">{{ app.status }}</span>{% endif %}
        </div>
        <div style="font-size:12px;color:#3a3b4a;margin-top:6px;">{{ user_info.email or 'N/A' }}</div>
        <div style="font-size:12px;color:#3a3b4a;margin-top:2px;">{{ user_info.phone or 'N/A' }}</div>
      </div>
      <div class="profile-actions">
        <button class="btn btn-ghost" onclick="generateQuestions()">✦ Generate Interview Questions</button>
        <button class="btn btn-primary" onclick="runAiScreen()">🤖 Re-Screen</button>
      </div>
    </div>

    <div style="display:grid;grid-template-columns:1fr 1fr 300px;gap:16px;margin-bottom:20px;">

      <!-- Scores -->
      <div class="card">
        <div class="card-header"><h2>📊 Scores</h2></div>
        <div style="padding:20px;">
          <div class="score-bar-row" id="ats-row">
            <div class="score-bar-label"><span>ATS Score</span><span class="score-bar-val" id="ats-val">—</span></div>
            <div class="score-bar-track"><div class="score-bar-fill" id="ats-bar" style="width:0%;background:linear-gradient(90deg,#818cf8,#6366f1);"></div></div>
          </div>
          <div class="score-bar-row" id="tech-row">
            <div class="score-bar-label"><span>Technical</span><span class="score-bar-val" id="tech-val">—</span></div>
            <div class="score-bar-track"><div class="score-bar-fill" id="tech-bar" style="width:0%;background:linear-gradient(90deg,#60a5fa,#3b82f6);"></div></div>
          </div>
          <div class="score-bar-row" id="fit-row">
            <div class="score-bar-label"><span>Overall Fit</span><span class="score-bar-val" id="fit-val">—</span></div>
            <div class="score-bar-track"><div class="score-bar-fill" id="fit-bar" style="width:0%;background:linear-gradient(90deg,#34d399,#059669);"></div></div>
          </div>
          <div style="margin-top:16px;text-align:center;">
            <div style="font-size:11px;color:#3a3b4a;text-transform:uppercase;letter-spacing:.5px;margin-bottom:6px;">Decision</div>
            <span class="badge {% if app.status=='Applied' %}badge-yellow{% elif app.status=='Shortlisted' %}badge-green{% elif app.status=='Rejected' %}badge-red{% else %}badge-blue{% endif %}" id="decision-badge" style="font-size:13px;padding:6px 16px;">{{ app.status|upper }}</span>
          </div>
        </div>
      </div>

      <!-- Skills & Contact -->
      <div class="card">
        <div class="card-header"><h2>🛠 Skills</h2></div>
        <div style="padding:16px;">
          <div id="skillsDisplay">
            {% if resume.skills %}
              {% for skill in resume.skills %}
              <span class="skill-chip" style="background:rgba(99,102,241,0.1);color:#818cf8;">{{ skill }}</span>
              {% endfor %}
            {% else %}
              <span style="color:#3a3b4a;font-size:13px;">No skills data — run AI Screen to analyse</span>
            {% endif %}
          </div>
          <div style="margin-top:18px;">
            <div style="font-size:12px;color:#3a3b4a;font-weight:600;text-transform:uppercase;letter-spacing:.5px;margin-bottom:10px;">Contact Info</div>
            <div style="font-size:13px;color:#9ca3c0;margin-bottom:6px;">✉ {{ user_info.email or 'N/A' }}</div>
            <div style="font-size:13px;color:#9ca3c0;">📞 {{ user_info.phone or 'N/A' }}</div>
          </div>
        </div>
      </div>

      <!-- Details panel -->
      <div class="card">
        <div class="card-header"><h2>📋 Details</h2></div>
        <div style="padding:12px 16px 16px;">
          <div class="detail-row"><span class="detail-key">Candidate ID</span><span class="detail-val">#{{ app.id }}</span></div>
          <div class="detail-row"><span class="detail-key">Applied For</span><span class="detail-val" style="color:#818cf8;">{{ app.job_title }}</span></div>
          <div class="detail-row"><span class="detail-key">ATS Score</span><span class="detail-val" id="detail-ats">—</span></div>
          <div class="detail-row"><span class="detail-key">Status</span><span class="detail-val" id="detail-status">{{ app.status }}</span></div>
          <div class="detail-row"><span class="detail-key">Applied On</span><span class="detail-val">{{ app.applied_at[:10] }}</span></div>
          <div style="margin-top:14px;display:flex;flex-direction:column;gap:8px;">
            <button class="btn btn-primary" style="width:100%;justify-content:center;" onclick="generateQuestions()">✦ Generate Interview Questions</button>
            <button class="btn btn-ghost" style="width:100%;justify-content:center;" onclick="runAiScreen()">🤖 Re-Screen AI</button>
          </div>
        </div>
      </div>
    </div>

    <!-- Notes & AI Insight -->
    <div class="card" style="margin-bottom:20px;">
      <div class="card-header"><h2>📝 Notes &amp; AI Insight</h2></div>
      <div style="padding:18px;" id="aiInsight">
        <p style="color:#3a3b4a;font-size:13px;">Click "Re-Screen AI" above to generate AI insights for this candidate.</p>
      </div>
    </div>

    <!-- Interview Questions (hidden until generated) -->
    <div class="card" id="questionsCard" style="display:none;">
      <div class="card-header"><h2>❓ Interview Questions</h2></div>
      <div style="padding:16px;" id="questionsList"></div>
    </div>

  </div>
</div>

<div class="modal-overlay" id="modal">
  <div class="modal">
    <div class="modal-header"><h2 id="modalTitle">Processing...</h2><button class="modal-close" onclick="closeModal()">×</button></div>
    <div id="modalContent"></div>
  </div>
</div>

<script>
const RESUME = {{ resume|tojson }};
const JOB_ID = {{ app.job_id }};
const APP_ID = {{ app.id }};

async function runAiScreen(){
  document.getElementById('modal').classList.add('show');
  document.getElementById('modalTitle').textContent='🤖 AI Screening';
  document.getElementById('modalContent').innerHTML='<p style="color:#555;padding:10px 0;">Analysing candidate against job requirements...</p>';
  try{
    const res=await fetch('/api/ai/job-match',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({resume:RESUME,job_id:JOB_ID})});
    const d=await res.json();
    // Update score bars
    const ats=d.match_percent||0;
    const tech=Math.max(0,ats-8);
    const fit=Math.max(0,ats-15);
    document.getElementById('ats-val').textContent=ats+'%';
    document.getElementById('ats-bar').style.width=ats+'%';
    document.getElementById('tech-val').textContent=tech+'%';
    document.getElementById('tech-bar').style.width=tech+'%';
    document.getElementById('fit-val').textContent=fit+'%';
    document.getElementById('fit-bar').style.width=fit+'%';
    document.getElementById('detail-ats').textContent=ats+'%';
    // Update skills
    if(d.matched_skills&&d.matched_skills.length){
      document.getElementById('skillsDisplay').innerHTML=d.matched_skills.map(s=>`<span class="skill-chip" style="background:rgba(52,211,153,0.1);color:#34d399;">${s}</span>`).join('')+
        (d.missing_skills||[]).map(s=>`<span class="skill-chip" style="background:rgba(248,113,113,0.1);color:#f87171;">${s}</span>`).join('');
    }
    // AI Insight
    document.getElementById('aiInsight').innerHTML=`
      <div style="background:#0a0b12;border-left:3px solid #6366f1;padding:14px;border-radius:10px;font-size:13px;color:#9ca3c0;line-height:1.7;">
        💬 ${d.recommendation||'AI analysis complete.'}
      </div>
      <div style="margin-top:14px;display:flex;gap:10px;flex-wrap:wrap;">
        <div><div style="font-size:11px;color:#3a3b4a;margin-bottom:5px;text-transform:uppercase;letter-spacing:.5px;">✅ Matched Skills</div>
          ${(d.matched_skills||[]).map(s=>`<span class="skill-chip" style="background:rgba(52,211,153,0.1);color:#34d399;">${s}</span>`).join('')||'<span style="color:#3a3b4a;font-size:12px;">None</span>'}
        </div>
        <div><div style="font-size:11px;color:#3a3b4a;margin-bottom:5px;text-transform:uppercase;letter-spacing:.5px;">❌ Missing</div>
          ${(d.missing_skills||[]).map(s=>`<span class="skill-chip" style="background:rgba(248,113,113,0.1);color:#f87171;">${s}</span>`).join('')||'<span style="color:#3a3b4a;font-size:12px;">None</span>'}
        </div>
      </div>`;
    document.getElementById('modalContent').innerHTML=`<div style="text-align:center;padding:10px;"><div style="font-size:36px;font-weight:700;color:#34d399;">${ats}%</div><div style="color:#6b7280;font-size:13px;">Match Score</div><p style="margin-top:12px;color:#9ca3c0;font-size:13px;">${d.recommendation||''}</p></div>`;
  }catch(e){document.getElementById('modalContent').innerHTML='<p style="color:#f87171;">AI screening failed.</p>';}
}

async function generateQuestions(){
  document.getElementById('modal').classList.add('show');
  document.getElementById('modalTitle').textContent='✦ Generating Interview Questions';
  document.getElementById('modalContent').innerHTML='<p style="color:#555;padding:10px 0;">Generating tailored questions...</p>';
  try{
    const res=await fetch('/api/recruiter/interview-questions',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({resume:RESUME,job_title:'{{ app.job_title }}'})});
    const d=await res.json();
    const qs=d.questions||[];
    const html=qs.map(q=>`<div class="q-item"><div class="q-type">${q.type||'General'}</div><div class="q-text">${q.question}</div></div>`).join('');
    document.getElementById('questionsCard').style.display='block';
    document.getElementById('questionsList').innerHTML=html;
    document.getElementById('modalContent').innerHTML=`<p style="color:#34d399;font-size:13px;">✅ ${qs.length} questions generated! See below the profile.</p>`;
    setTimeout(()=>closeModal(),2000);
  }catch(e){document.getElementById('modalContent').innerHTML='<p style="color:#f87171;">Failed to generate questions.</p>';}
}

async function quickStatus(appId,status){
  await fetch('/api/recruiter/update-status',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({app_id:appId,status})});
  const badge=document.getElementById('decision-badge');
  if(badge){badge.textContent=status.toUpperCase();}
  document.getElementById('detail-status').textContent=status;
  alert('Status updated to: '+status);
}

function closeModal(){document.getElementById('modal').classList.remove('show');}
document.getElementById('modal').addEventListener('click',function(e){if(e.target===this)closeModal();});

// Auto-run AI screen on load if resume has data
if(RESUME && RESUME.name){ setTimeout(runAiScreen, 800); }
</script>
<script>
let notifOpen = false;
async function loadNotifications(){
  try{
    const res = await fetch('/api/notifications');
    const d = await res.json();
    const badge = document.getElementById('notifBadge');
    const list = document.getElementById('notifList');
    if(!badge||!list) return;
    if(d.unread > 0){
      badge.style.display='flex';
      badge.textContent = d.unread > 9 ? '9+' : d.unread;
    } else {
      badge.style.display='none';
    }
    if(!d.notifications || d.notifications.length === 0){
      list.innerHTML = '<div class="notif-empty">🔔 No notifications yet</div>';
      return;
    }
    const icons = {application:'📋', status:'📄', screening:'🤖', job:'💼'};
    const colors = {application:'rgba(99,102,241,0.15)', status:'rgba(52,211,153,0.12)', screening:'rgba(168,85,247,0.12)', job:'rgba(251,191,36,0.12)'};
    list.innerHTML = d.notifications.map(n => {
      const t = new Date(n.created_at).toLocaleString('en-IN',{day:'2-digit',month:'short',hour:'2-digit',minute:'2-digit'});
      return `<div class="notif-item ${n.is_read?'':'unread'}" onclick="goNotif('${n.link||'/'}')">
        <div class="notif-icon" style="background:${colors[n.type]||'rgba(99,102,241,0.12)'};">${icons[n.type]||'🔔'}</div>
        <div style="flex:1;">
          <div class="notif-msg">${n.message}</div>
          <div class="notif-time">${t}</div>
        </div>
        ${n.is_read?'':'<div class="notif-unread-dot"></div>'}
      </div>`;
    }).join('');
  } catch(e){}
}
function toggleNotif(){
  notifOpen = !notifOpen;
  document.getElementById('notifPanel').classList.toggle('open', notifOpen);
  if(notifOpen) loadNotifications();
}
function goNotif(link){
  markAllRead();
  window.location.href = link;
}
async function markAllRead(){
  await fetch('/api/notifications/mark-read',{method:'POST'});
  document.getElementById('notifBadge').style.display='none';
  document.querySelectorAll('.notif-item').forEach(el=>el.classList.remove('unread'));
  document.querySelectorAll('.notif-unread-dot').forEach(el=>el.remove());
}
document.addEventListener('click', function(e){
  if(notifOpen && !document.getElementById('bellBtn').contains(e.target) && !document.getElementById('notifPanel').contains(e.target)){
    notifOpen = false;
    document.getElementById('notifPanel').classList.remove('open');
  }
});
// poll every 30s
loadNotifications();
setInterval(loadNotifications, 30000);
</script>
</body></html>"""


POST_JOB_HTML = """<!DOCTYPE html>
<html><head><title>Post a Job – RecruitAI</title>
<link href="https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;600;700&display=swap" rel="stylesheet">
<style>""" + SIDEBAR_CSS + """
.form-card{background:#0f1018;border-radius:16px;padding:28px;border:1px solid rgba(255,255,255,0.05);max-width:640px;}
label{display:block;margin-bottom:5px;margin-top:14px;font-size:12px;color:#555;font-weight:500;}
input,textarea,select{width:100%;padding:11px 13px;border:1px solid rgba(255,255,255,0.07);outline:none;border-radius:11px;background:#0a0b12;color:white;font-size:14px;font-family:'DM Sans',sans-serif;transition:border-color .2s;}
input:focus,textarea:focus,select:focus{border-color:#6366f1;}
select option{background:#0f1018;}
textarea{resize:vertical;}
.row{display:grid;grid-template-columns:1fr 1fr;gap:14px;}
</style></head><body>

<div class="sidebar">
  <div class="logo">✦ Recruit<span>AI</span></div>
  <div class="nav-section">
    <div class="nav-label">MAIN</div>
    <a class="nav-item" href="/recruiter/dashboard"><span class="icon">🏠</span> Dashboard</a>
    <a class="nav-item" href="/recruiter/screen-resume"><span class="icon">📄</span> Screen Resume</a>
    <a class="nav-item" href="/recruiter/candidates"><span class="icon">👥</span> Candidates</a>
    <a class="nav-item" href="/recruiter/applications"><span class="icon">📋</span> Applications</a>
    <a class="nav-item active" href="/recruiter/post-job"><span class="icon">➕</span> Post a Job</a>
    <div class="nav-label" style="margin-top:10px;">TOOLS</div>
    <a class="nav-item" href="/recruiter/analytics"><span class="icon">📊</span> Analytics</a>
    <a class="nav-item" href="/recruiter/history"><span class="icon">🕐</span> History</a>
  </div>
  <div class="sidebar-bottom">
    <div class="user-chip">
      <div class="avatar">{ username[0].upper() }</div>
      <div class="user-info"><div class="name">{ username }</div><div class="role">Recruiter</div></div>
    </div>
    <a href="/logout" class="logout-link">← Sign out</a>
  </div>
</div>
<div class="main">
  <div class="topbar">
    <div><h1>Post a New Job</h1><div class="sub">Fill in the details to attract the right candidates</div></div>
  </div>
  <div class="content">
    <div class="form-card">
      <form method="POST">
        <label>Job Title *</label>
        <input type="text" name="title" placeholder="e.g. Senior Python Developer" required>
        <div class="row">
          <div><label>Company Name *</label><input type="text" name="company" placeholder="e.g. Infosys" required></div>
          <div><label>Location</label><input type="text" name="location" placeholder="e.g. Mumbai / Remote"></div>
        </div>
        <div class="row">
          <div><label>Job Type</label>
            <select name="job_type"><option>Full-time</option><option>Part-time</option><option>Contract</option><option>Internship</option><option>Remote</option></select>
          </div>
          <div><label>Salary / Package</label><input type="text" name="salary" placeholder="e.g. ₹12–18 LPA"></div>
        </div>
        <label>Required Skills (comma separated)</label>
        <input type="text" name="skills_required" placeholder="e.g. Python, Flask, SQL, REST API">
        <label>Job Description *</label>
        <textarea rows="7" name="description" placeholder="Describe the role, responsibilities, and requirements..." required></textarea>
        <div style="margin-top:20px;display:flex;gap:10px;">
          <button type="submit" class="btn btn-primary" style="padding:12px 28px;font-size:14px;">🚀 Post Job</button>
          <a href="/recruiter/dashboard" class="btn btn-ghost" style="padding:12px 20px;font-size:14px;">Cancel</a>
        </div>
      </form>
    </div>
  </div>
</div>
<script>
let notifOpen = false;
async function loadNotifications(){
  try{
    const res = await fetch('/api/notifications');
    const d = await res.json();
    const badge = document.getElementById('notifBadge');
    const list = document.getElementById('notifList');
    if(!badge||!list) return;
    if(d.unread > 0){
      badge.style.display='flex';
      badge.textContent = d.unread > 9 ? '9+' : d.unread;
    } else {
      badge.style.display='none';
    }
    if(!d.notifications || d.notifications.length === 0){
      list.innerHTML = '<div class="notif-empty">🔔 No notifications yet</div>';
      return;
    }
    const icons = {application:'📋', status:'📄', screening:'🤖', job:'💼'};
    const colors = {application:'rgba(99,102,241,0.15)', status:'rgba(52,211,153,0.12)', screening:'rgba(168,85,247,0.12)', job:'rgba(251,191,36,0.12)'};
    list.innerHTML = d.notifications.map(n => {
      const t = new Date(n.created_at).toLocaleString('en-IN',{day:'2-digit',month:'short',hour:'2-digit',minute:'2-digit'});
      return `<div class="notif-item ${n.is_read?'':'unread'}" onclick="goNotif('${n.link||'/'}')">
        <div class="notif-icon" style="background:${colors[n.type]||'rgba(99,102,241,0.12)'};">${icons[n.type]||'🔔'}</div>
        <div style="flex:1;">
          <div class="notif-msg">${n.message}</div>
          <div class="notif-time">${t}</div>
        </div>
        ${n.is_read?'':'<div class="notif-unread-dot"></div>'}
      </div>`;
    }).join('');
  } catch(e){}
}
function toggleNotif(){
  notifOpen = !notifOpen;
  document.getElementById('notifPanel').classList.toggle('open', notifOpen);
  if(notifOpen) loadNotifications();
}
function goNotif(link){
  markAllRead();
  window.location.href = link;
}
async function markAllRead(){
  await fetch('/api/notifications/mark-read',{method:'POST'});
  document.getElementById('notifBadge').style.display='none';
  document.querySelectorAll('.notif-item').forEach(el=>el.classList.remove('unread'));
  document.querySelectorAll('.notif-unread-dot').forEach(el=>el.remove());
}
document.addEventListener('click', function(e){
  if(notifOpen && !document.getElementById('bellBtn').contains(e.target) && !document.getElementById('notifPanel').contains(e.target)){
    notifOpen = false;
    document.getElementById('notifPanel').classList.remove('open');
  }
});
// poll every 30s
loadNotifications();
setInterval(loadNotifications, 30000);
</script>
</body></html>"""


APPLICANTS_HTML = """<!DOCTYPE html>
<html><head><title>Applicants – RecruitAI</title>
<link href="https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;600;700&display=swap" rel="stylesheet">
<style>""" + SIDEBAR_CSS + """
</style></head><body>

<div class="sidebar">
  <div class="logo">✦ Recruit<span>AI</span></div>
  <div class="nav-section">
    <div class="nav-label">MAIN</div>
    <a class="nav-item" href="/recruiter/dashboard"><span class="icon">🏠</span> Dashboard</a>
    <a class="nav-item" href="/recruiter/screen-resume"><span class="icon">📄</span> Screen Resume</a>
    <a class="nav-item" href="/recruiter/candidates"><span class="icon">👥</span> Candidates</a>
    <a class="nav-item active" href="/recruiter/applications"><span class="icon">📋</span> Applications</a>
    <a class="nav-item" href="/recruiter/post-job"><span class="icon">➕</span> Post a Job</a>
    <div class="nav-label" style="margin-top:10px;">TOOLS</div>
    <a class="nav-item" href="/recruiter/analytics"><span class="icon">📊</span> Analytics</a>
    <a class="nav-item" href="/recruiter/history"><span class="icon">🕐</span> History</a>
  </div>
  <div class="sidebar-bottom">
    <div class="user-chip">
      <div class="avatar">{ username[0].upper() }</div>
      <div class="user-info"><div class="name">{ username }</div><div class="role">Recruiter</div></div>
    </div>
    <a href="/logout" class="logout-link">← Sign out</a>
  </div>
</div>
<div class="main">
  <div class="topbar">
    <div><h1>{{ job.title }}</h1><div class="sub">{{ job.company }} &nbsp;•&nbsp; {{ applicants|length }} applicant(s)</div></div>
    <a href="/recruiter/applications" class="btn btn-ghost">← Back</a>
  </div>
  <div class="content">
    <div class="card">
      <div class="card-body">
        {% if applicants %}
        <table>
          <thead><tr><th>Candidate</th><th>Applied</th><th>Status</th><th>AI Screen</th><th>Update</th></tr></thead>
          <tbody>
          {% for app in applicants %}
          <tr>
            <td><strong>{{ app.username }}</strong></td>
            <td>{{ app.applied_at[:10] }}</td>
            <td>
              {% if app.status=='Applied' %}<span class="badge badge-yellow">Pending</span>
              {% elif app.status=='Shortlisted' %}<span class="badge badge-green">{{ app.status }}</span>
              {% elif app.status=='Rejected' %}<span class="badge badge-red">{{ app.status }}</span>
              {% elif app.status=='Hired' %}<span class="badge badge-blue">{{ app.status }}</span>
              {% else %}<span class="badge badge-purple">{{ app.status }}</span>{% endif %}
            </td>
            <td>
              <a href="/recruiter/candidate/{{ app.id }}" class="btn btn-primary btn-sm">View Profile</a>
            </td>
            <td>
              <select onchange="updateStatus({{ app.id }}, this.value)"
                style="background:#0a0b12;color:#aaa;border:1px solid rgba(255,255,255,0.08);border-radius:8px;padding:4px 8px;font-size:11px;cursor:pointer;font-family:'DM Sans',sans-serif;">
                {% for s in ['Applied','Reviewing','Shortlisted','Interviewing','Hired','Rejected'] %}
                <option {% if app.status==s %}selected{% endif %}>{{ s }}</option>
                {% endfor %}
              </select>
            </td>
          </tr>
          {% endfor %}
          </tbody>
        </table>
        {% else %}
        <div class="empty-state"><div class="e-icon">📭</div><h3>No applicants yet</h3><p>Share your job posting to attract candidates</p></div>
        {% endif %}
      </div>
    </div>
  </div>
</div>
<script>
async function updateStatus(appId,status){
  await fetch('/api/recruiter/update-status',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({app_id:appId,status})});
}
</script>
</body></html>"""


# ============================================================
# SHARED RECRUITER SIDEBAR SNIPPET (inline — used in templates below)
# ============================================================
def _rec_nav(active=""):
    items = [
        ("dashboard",     "🏠", "Dashboard",     "/recruiter/dashboard"),
        ("screen-resume", "📄", "Screen Resume",  "/recruiter/screen-resume"),
        ("candidates",    "👥", "Candidates",     "/recruiter/candidates"),
        ("applications",  "📋", "Applications",   "/recruiter/applications"),
        ("post-job",      "➕", "Post a Job",     "/recruiter/post-job"),
    ]
    tools = [
        ("analytics", "📊", "Analytics", "/recruiter/analytics"),
        ("history",   "🕐", "History",   "/recruiter/history"),
    ]
    html = '<div class="nav-label">MAIN</div>'
    for key, icon, label, href in items:
        cls = "nav-item active" if active == key else "nav-item"
        html += f'<a class="{cls}" href="{href}"><span class="icon">{icon}</span> {label}</a>'
    html += '<div class="nav-label" style="margin-top:10px;">TOOLS</div>'
    for key, icon, label, href in tools:
        cls = "nav-item active" if active == key else "nav-item"
        html += f'<a class="{cls}" href="{href}"><span class="icon">{icon}</span> {label}</a>'
    return html


SCREEN_RESUME_HTML = """<!DOCTYPE html>
<html><head><title>Screen Resume – RecruitAI</title>
<link href="https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;600;700&display=swap" rel="stylesheet">
<style>""" + SIDEBAR_CSS + """
.upload-zone{border:2px dashed rgba(99,102,241,0.35);border-radius:16px;padding:40px;text-align:center;cursor:pointer;transition:all .2s;background:#0a0b12;position:relative;}
.upload-zone:hover,.upload-zone.drag{border-color:#6366f1;background:rgba(99,102,241,0.06);}
.upload-zone input[type=file]{position:absolute;inset:0;opacity:0;cursor:pointer;width:100%;height:100%;}
.uz-icon{font-size:40px;margin-bottom:10px;}
.uz-title{font-size:15px;font-weight:600;color:#c4b5fd;margin-bottom:4px;}
.uz-sub{font-size:12px;color:#3a3b4a;}
label{display:block;margin-bottom:5px;margin-top:14px;font-size:12px;color:#555;font-weight:500;}
input[type=text],textarea{width:100%;padding:11px 13px;border:1px solid rgba(255,255,255,0.07);outline:none;border-radius:11px;background:#0a0b12;color:white;font-size:13px;font-family:'DM Sans',sans-serif;transition:border-color .2s;}
input[type=text]:focus,textarea:focus{border-color:#6366f1;}
.score-arc{display:flex;flex-direction:column;align-items:center;justify-content:center;}
.arc-num{font-size:42px;font-weight:700;letter-spacing:-2px;}
.arc-label{font-size:11px;color:#555;text-transform:uppercase;letter-spacing:.5px;margin-top:2px;}
.score-bar-row{margin-bottom:12px;}
.score-bar-label{display:flex;justify-content:space-between;margin-bottom:4px;font-size:13px;color:#9ca3c0;}
.score-bar-track{height:6px;border-radius:99px;background:#12131f;overflow:hidden;}
.score-bar-fill{height:100%;border-radius:99px;transition:width 1.2s ease;}
.chip{display:inline-flex;align-items:center;padding:4px 11px;border-radius:99px;font-size:11px;font-weight:500;margin:3px;}
.rec-pass{background:rgba(52,211,153,0.12);color:#34d399;border:1px solid rgba(52,211,153,0.2);}
.rec-maybe{background:rgba(251,191,36,0.12);color:#fbbf24;border:1px solid rgba(251,191,36,0.2);}
.rec-fail{background:rgba(248,113,113,0.12);color:#f87171;border:1px solid rgba(248,113,113,0.2);}
</style></head><body>
<div class="sidebar">
  <div class="logo">✦ Recruit<span>AI</span></div>
  <div class="nav-section">""" + _rec_nav("screen-resume") + """
  </div>
  <div class="sidebar-bottom">
    <div class="user-chip">
      <div class="avatar">{{ username[0].upper() }}</div>
      <div class="user-info"><div class="name">{{ username }}</div><div class="role">Recruiter</div></div>
    </div>
    <a href="/logout" class="logout-link">← Sign out</a>
  </div>
</div>

<div class="main">
  <div class="topbar">
    <div><h1>Screen Resume</h1><div class="sub">Upload a candidate's PDF resume for instant AI screening</div></div>
  </div>
  <div class="content">
    <div style="display:grid;grid-template-columns:420px 1fr;gap:20px;align-items:start;">

      <!-- Upload Form -->
      <div>
        <div class="card" style="padding:24px;">
          <div style="font-size:14px;font-weight:600;color:#c4b5fd;margin-bottom:16px;">📤 Upload Resume</div>
          <form method="POST" enctype="multipart/form-data" id="screenForm">
            <div class="upload-zone" id="dropZone">
              <input type="file" name="resume_pdf" id="resumeFile" accept=".pdf" required onchange="onFile(this)">
              <div class="uz-icon">📄</div>
              <div class="uz-title" id="uzTitle">Drop PDF here or click to browse</div>
              <div class="uz-sub">PDF only • Max 5MB</div>
            </div>
            <label>Candidate Name</label>
            <input type="text" name="candidate_name" placeholder="e.g. Rahul Sharma" value="{{ result.candidate_name if result else '' }}">
            <label>Job Description (optional — for tailored scoring)</label>
            <textarea name="job_description" rows="4" placeholder="Paste the job description here..."></textarea>
            <button type="submit" class="btn btn-primary" style="width:100%;justify-content:center;margin-top:16px;padding:13px;font-size:14px;" id="submitBtn">
              🤖 Screen Resume
            </button>
          </form>
        </div>
      </div>

      <!-- Results -->
      <div id="resultsPanel">
        {% if result and not result.error %}
        <div class="card" style="margin-bottom:16px;">
          <div style="padding:24px;">
            <div style="display:flex;align-items:center;gap:20px;margin-bottom:20px;">
              <div style="width:80px;height:80px;border-radius:50%;border:4px solid {{ '#34d399' if result.ats_score >= 70 else '#fbbf24' if result.ats_score >= 50 else '#f87171' }};display:flex;align-items:center;justify-content:center;flex-shrink:0;">
                <div class="score-arc">
                  <div class="arc-num" style="color:{{ '#34d399' if result.ats_score >= 70 else '#fbbf24' if result.ats_score >= 50 else '#f87171' }};">{{ result.ats_score }}</div>
                </div>
              </div>
              <div style="flex:1;">
                <div style="font-size:18px;font-weight:700;color:#e2e8f0;margin-bottom:4px;">{{ result.candidate_name }}</div>
                <div style="font-size:12px;color:#3a3b4a;margin-bottom:8px;">{{ result.filename }}</div>
                <span class="chip {% if result.recommendation == 'Hire' %}rec-pass{% elif result.recommendation == 'Maybe' %}rec-maybe{% else %}rec-fail{% endif %}">
                  {{ '✓ Hire' if result.recommendation == 'Hire' else '~ Maybe' if result.recommendation == 'Maybe' else '✗ Pass' }}
                </span>
              </div>
              <div style="text-align:right;">
                <div style="font-size:22px;font-weight:700;color:#e2e8f0;">{{ result.grade }}</div>
                <div style="font-size:11px;color:#3a3b4a;">Grade</div>
              </div>
            </div>

            <!-- Score bars -->
            <div class="score-bar-row">
              <div class="score-bar-label"><span>ATS Score</span><span style="font-weight:700;color:#818cf8;">{{ result.ats_score }}%</span></div>
              <div class="score-bar-track"><div class="score-bar-fill" style="width:{{ result.ats_score }}%;background:linear-gradient(90deg,#818cf8,#6366f1);"></div></div>
            </div>
            <div class="score-bar-row">
              <div class="score-bar-label"><span>Technical</span><span style="font-weight:700;color:#60a5fa;">{{ result.technical_score }}%</span></div>
              <div class="score-bar-track"><div class="score-bar-fill" style="width:{{ result.technical_score }}%;background:linear-gradient(90deg,#60a5fa,#3b82f6);"></div></div>
            </div>
            <div class="score-bar-row">
              <div class="score-bar-label"><span>Overall Fit</span><span style="font-weight:700;color:#34d399;">{{ result.overall_fit }}%</span></div>
              <div class="score-bar-track"><div class="score-bar-fill" style="width:{{ result.overall_fit }}%;background:linear-gradient(90deg,#34d399,#059669);"></div></div>
            </div>
            <div style="font-size:12px;color:#555;margin-top:6px;">Keyword match: <span style="color:#a78bfa;font-weight:600;">{{ result.keyword_match }}%</span></div>
          </div>
        </div>

        <!-- Summary + Skills -->
        <div class="card" style="margin-bottom:16px;">
          <div style="padding:20px;">
            {% if result.summary %}
            <div style="background:#0a0b12;border-left:3px solid #6366f1;padding:12px 14px;border-radius:10px;font-size:13px;color:#9ca3c0;line-height:1.7;margin-bottom:16px;">
              💬 {{ result.summary }}
            </div>
            {% endif %}
            <div style="display:grid;grid-template-columns:1fr 1fr;gap:14px;">
              <div>
                <div style="font-size:11px;color:#3a3b4a;text-transform:uppercase;letter-spacing:.5px;margin-bottom:8px;font-weight:600;">✅ Strengths</div>
                {% for s in result.strengths %}
                <div style="font-size:12px;color:#9ca3c0;margin-bottom:5px;display:flex;gap:6px;"><span style="color:#34d399;">•</span>{{ s }}</div>
                {% endfor %}
              </div>
              <div>
                <div style="font-size:11px;color:#3a3b4a;text-transform:uppercase;letter-spacing:.5px;margin-bottom:8px;font-weight:600;">⚠️ Improvements</div>
                {% for s in result.improvements %}
                <div style="font-size:12px;color:#9ca3c0;margin-bottom:5px;display:flex;gap:6px;"><span style="color:#fbbf24;">•</span>{{ s }}</div>
                {% endfor %}
              </div>
            </div>
            <div style="margin-top:14px;">
              <div style="font-size:11px;color:#3a3b4a;text-transform:uppercase;letter-spacing:.5px;margin-bottom:6px;font-weight:600;">🛠 Skills Found</div>
              {% for s in result.skills_found %}
              <span class="chip" style="background:rgba(52,211,153,0.1);color:#34d399;">{{ s }}</span>
              {% endfor %}
              {% for s in result.skills_missing %}
              <span class="chip" style="background:rgba(248,113,113,0.1);color:#f87171;">{{ s }}</span>
              {% endfor %}
            </div>
            <div style="margin-top:12px;">
              <div style="font-size:11px;color:#3a3b4a;text-transform:uppercase;letter-spacing:.5px;margin-bottom:6px;font-weight:600;">📋 Sections</div>
              {% for s in result.sections_found %}
              <span class="chip" style="background:rgba(99,102,241,0.1);color:#818cf8;">{{ s }}</span>
              {% endfor %}
              {% for s in result.sections_missing %}
              <span class="chip" style="background:rgba(248,113,113,0.08);color:#f87171;text-decoration:line-through;">{{ s }}</span>
              {% endfor %}
            </div>
          </div>
        </div>
        {% elif result and result.error %}
        <div class="card" style="padding:24px;">
          <p style="color:#f87171;">{{ result.error }}</p>
        </div>
        {% else %}
        <div class="card" style="padding:48px;text-align:center;">
          <div style="font-size:48px;margin-bottom:12px;opacity:.3;">🤖</div>
          <div style="font-size:15px;color:#3a3b4a;margin-bottom:6px;">AI Screening Results</div>
          <div style="font-size:13px;color:#2a2b3a;">Upload a PDF resume and click Screen Resume to see the analysis here</div>
        </div>
        {% endif %}
      </div>
    </div>
  </div>
</div>
<script>
const dz=document.getElementById('dropZone');
function onFile(input){if(input.files[0])document.getElementById('uzTitle').textContent='✅ '+input.files[0].name;}
dz.addEventListener('dragover',e=>{e.preventDefault();dz.classList.add('drag');});
dz.addEventListener('dragleave',()=>dz.classList.remove('drag'));
dz.addEventListener('drop',e=>{
  e.preventDefault();dz.classList.remove('drag');
  const f=e.dataTransfer.files[0];
  if(f&&f.type==='application/pdf'){
    document.getElementById('resumeFile').files=e.dataTransfer.files;
    document.getElementById('uzTitle').textContent='✅ '+f.name;
  }
});
document.getElementById('screenForm').onsubmit=()=>{
  document.getElementById('submitBtn').textContent='⏳ Analysing...';
  document.getElementById('submitBtn').disabled=true;
};
</script>
<script>
let notifOpen = false;
async function loadNotifications(){
  try{
    const res = await fetch('/api/notifications');
    const d = await res.json();
    const badge = document.getElementById('notifBadge');
    const list = document.getElementById('notifList');
    if(!badge||!list) return;
    if(d.unread > 0){
      badge.style.display='flex';
      badge.textContent = d.unread > 9 ? '9+' : d.unread;
    } else {
      badge.style.display='none';
    }
    if(!d.notifications || d.notifications.length === 0){
      list.innerHTML = '<div class="notif-empty">🔔 No notifications yet</div>';
      return;
    }
    const icons = {application:'📋', status:'📄', screening:'🤖', job:'💼'};
    const colors = {application:'rgba(99,102,241,0.15)', status:'rgba(52,211,153,0.12)', screening:'rgba(168,85,247,0.12)', job:'rgba(251,191,36,0.12)'};
    list.innerHTML = d.notifications.map(n => {
      const t = new Date(n.created_at).toLocaleString('en-IN',{day:'2-digit',month:'short',hour:'2-digit',minute:'2-digit'});
      return `<div class="notif-item ${n.is_read?'':'unread'}" onclick="goNotif('${n.link||'/'}')">
        <div class="notif-icon" style="background:${colors[n.type]||'rgba(99,102,241,0.12)'};">${icons[n.type]||'🔔'}</div>
        <div style="flex:1;">
          <div class="notif-msg">${n.message}</div>
          <div class="notif-time">${t}</div>
        </div>
        ${n.is_read?'':'<div class="notif-unread-dot"></div>'}
      </div>`;
    }).join('');
  } catch(e){}
}
function toggleNotif(){
  notifOpen = !notifOpen;
  document.getElementById('notifPanel').classList.toggle('open', notifOpen);
  if(notifOpen) loadNotifications();
}
function goNotif(link){
  markAllRead();
  window.location.href = link;
}
async function markAllRead(){
  await fetch('/api/notifications/mark-read',{method:'POST'});
  document.getElementById('notifBadge').style.display='none';
  document.querySelectorAll('.notif-item').forEach(el=>el.classList.remove('unread'));
  document.querySelectorAll('.notif-unread-dot').forEach(el=>el.remove());
}
document.addEventListener('click', function(e){
  if(notifOpen && !document.getElementById('bellBtn').contains(e.target) && !document.getElementById('notifPanel').contains(e.target)){
    notifOpen = false;
    document.getElementById('notifPanel').classList.remove('open');
  }
});
// poll every 30s
loadNotifications();
setInterval(loadNotifications, 30000);
</script>
</body></html>"""


RECRUITER_ANALYTICS_HTML = """<!DOCTYPE html>
<html><head><title>Analytics – RecruitAI</title>
<link href="https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;600;700&display=swap" rel="stylesheet">
<style>""" + SIDEBAR_CSS + """
.prog-bar-wrap{background:#12131f;border-radius:99px;height:7px;overflow:hidden;margin-top:5px;}
.prog-bar{height:100%;border-radius:99px;background:linear-gradient(90deg,#6366f1,#a78bfa);transition:width 1s ease;}
</style></head><body>
<div class="sidebar">
  <div class="logo">✦ Recruit<span>AI</span></div>
  <div class="nav-section">""" + _rec_nav("analytics") + """
  </div>
  <div class="sidebar-bottom">
    <div class="user-chip">
      <div class="avatar">{{ username[0].upper() }}</div>
      <div class="user-info"><div class="name">{{ username }}</div><div class="role">Recruiter</div></div>
    </div>
    <a href="/logout" class="logout-link">← Sign out</a>
  </div>
</div>
<div class="main">
  <div class="topbar">
    <div><h1>Analytics</h1><div class="sub">Hiring funnel and performance overview</div></div>
  </div>
  <div class="content">
    <div class="stats-grid" style="grid-template-columns:repeat(5,1fr);margin-bottom:24px;">
      <div class="stat-card"><div class="s-icon">👥</div><div class="s-val">{{ total_apps }}</div><div class="s-label">Total Applications</div></div>
      <div class="stat-card"><div class="s-icon">⏳</div><div class="s-val" style="color:#fbbf24;">{{ pending }}</div><div class="s-label">Pending</div></div>
      <div class="stat-card"><div class="s-icon">⭐</div><div class="s-val" style="color:#818cf8;">{{ shortlisted }}</div><div class="s-label">Shortlisted</div></div>
      <div class="stat-card"><div class="s-icon">✅</div><div class="s-val" style="color:#34d399;">{{ hired }}</div><div class="s-label">Hired</div></div>
      <div class="stat-card"><div class="s-icon">❌</div><div class="s-val" style="color:#f87171;">{{ rejected }}</div><div class="s-label">Rejected</div></div>
    </div>

    <div style="display:grid;grid-template-columns:1fr 1fr;gap:16px;margin-bottom:20px;">
      <!-- Funnel -->
      <div class="card">
        <div class="card-header"><h2>📊 Hiring Funnel</h2></div>
        <div style="padding:20px;">
          {% set stages = [('Applications', total_apps, '#6366f1'), ('Shortlisted', shortlisted, '#818cf8'), ('Hired', hired, '#34d399'), ('Rejected', rejected, '#f87171')] %}
          {% for label, val, color in stages %}
          <div style="margin-bottom:14px;">
            <div style="display:flex;justify-content:space-between;font-size:12px;color:#9ca3c0;margin-bottom:4px;">
              <span>{{ label }}</span>
              <span style="font-weight:700;color:{{ color }};">{{ val }}</span>
            </div>
            <div class="prog-bar-wrap">
              <div class="prog-bar" style="width:{% if total_apps %}{{ ((val / total_apps) * 100)|int }}{% else %}0{% endif %}%;background:{{ color }};"></div>
            </div>
          </div>
          {% endfor %}
        </div>
      </div>

      <!-- Donut -->
      <div class="card">
        <div class="card-header"><h2>🍩 Status Breakdown</h2></div>
        <div style="padding:20px;display:flex;align-items:center;gap:20px;">
          <canvas id="donut" width="140" height="140"></canvas>
          <div style="display:flex;flex-direction:column;gap:10px;">
            <div style="display:flex;align-items:center;gap:8px;font-size:12px;color:#9ca3c0;"><span style="width:10px;height:10px;border-radius:50%;background:#fbbf24;flex-shrink:0;display:inline-block;"></span>Pending ({{ pending }})</div>
            <div style="display:flex;align-items:center;gap:8px;font-size:12px;color:#9ca3c0;"><span style="width:10px;height:10px;border-radius:50%;background:#818cf8;flex-shrink:0;display:inline-block;"></span>Shortlisted ({{ shortlisted }})</div>
            <div style="display:flex;align-items:center;gap:8px;font-size:12px;color:#9ca3c0;"><span style="width:10px;height:10px;border-radius:50%;background:#34d399;flex-shrink:0;display:inline-block;"></span>Hired ({{ hired }})</div>
            <div style="display:flex;align-items:center;gap:8px;font-size:12px;color:#9ca3c0;"><span style="width:10px;height:10px;border-radius:50%;background:#f87171;flex-shrink:0;display:inline-block;"></span>Rejected ({{ rejected }})</div>
          </div>
        </div>
      </div>
    </div>

    <!-- Per-job breakdown -->
    <div class="card">
      <div class="card-header"><h2>💼 Jobs Performance</h2></div>
      <div class="card-body">
        {% if job_stats %}
        <table>
          <thead><tr><th>Job Title</th><th>Total Applicants</th><th>Pipeline</th><th>Actions</th></tr></thead>
          <tbody>
          {% for js in job_stats %}
          <tr>
            <td><strong>{{ js.title }}</strong></td>
            <td><span class="badge badge-purple">{{ js.count }}</span></td>
            <td>
              <div class="prog-bar-wrap" style="width:200px;">
                <div class="prog-bar" style="width:{% if js.count %}{{ [js.count * 15, 100]|min if false else (100 if js.count >= 7 else js.count * 14) }}{% else %}0{% endif %}%;"></div>
              </div>
            </td>
            <td><a href="/recruiter/job/{{ js.id }}/applicants" class="btn btn-ghost btn-sm">View →</a></td>
          </tr>
          {% endfor %}
          </tbody>
        </table>
        {% else %}
        <div class="empty-state"><div class="e-icon">💼</div><h3>No jobs posted yet</h3></div>
        {% endif %}
      </div>
    </div>
  </div>
</div>
<script>
window.addEventListener('load',()=>{
  const c=document.getElementById('donut');if(!c)return;
  const ctx=c.getContext('2d');
  const data=[{{ pending }},{{ shortlisted }},{{ hired }},{{ rejected }}];
  const colors=['#fbbf24','#818cf8','#34d399','#f87171'];
  const sum=data.reduce((a,b)=>a+b,0)||1;
  let start=-Math.PI/2;const cx=70,cy=70,r=60,inner=35;
  data.forEach((v,i)=>{
    const angle=(v/sum)*Math.PI*2;
    ctx.beginPath();ctx.moveTo(cx,cy);ctx.arc(cx,cy,r,start,start+angle);ctx.closePath();
    ctx.fillStyle=colors[i];ctx.fill();start+=angle;
  });
  ctx.beginPath();ctx.arc(cx,cy,inner,0,Math.PI*2);ctx.fillStyle='#0f1018';ctx.fill();
  ctx.fillStyle='#6b7280';ctx.font='bold 13px DM Sans,sans-serif';
  ctx.textAlign='center';ctx.textBaseline='middle';ctx.fillText({{ total_apps }},cx,cy);
});
</script>
<script>
let notifOpen = false;
async function loadNotifications(){
  try{
    const res = await fetch('/api/notifications');
    const d = await res.json();
    const badge = document.getElementById('notifBadge');
    const list = document.getElementById('notifList');
    if(!badge||!list) return;
    if(d.unread > 0){
      badge.style.display='flex';
      badge.textContent = d.unread > 9 ? '9+' : d.unread;
    } else {
      badge.style.display='none';
    }
    if(!d.notifications || d.notifications.length === 0){
      list.innerHTML = '<div class="notif-empty">🔔 No notifications yet</div>';
      return;
    }
    const icons = {application:'📋', status:'📄', screening:'🤖', job:'💼'};
    const colors = {application:'rgba(99,102,241,0.15)', status:'rgba(52,211,153,0.12)', screening:'rgba(168,85,247,0.12)', job:'rgba(251,191,36,0.12)'};
    list.innerHTML = d.notifications.map(n => {
      const t = new Date(n.created_at).toLocaleString('en-IN',{day:'2-digit',month:'short',hour:'2-digit',minute:'2-digit'});
      return `<div class="notif-item ${n.is_read?'':'unread'}" onclick="goNotif('${n.link||'/'}')">
        <div class="notif-icon" style="background:${colors[n.type]||'rgba(99,102,241,0.12)'};">${icons[n.type]||'🔔'}</div>
        <div style="flex:1;">
          <div class="notif-msg">${n.message}</div>
          <div class="notif-time">${t}</div>
        </div>
        ${n.is_read?'':'<div class="notif-unread-dot"></div>'}
      </div>`;
    }).join('');
  } catch(e){}
}
function toggleNotif(){
  notifOpen = !notifOpen;
  document.getElementById('notifPanel').classList.toggle('open', notifOpen);
  if(notifOpen) loadNotifications();
}
function goNotif(link){
  markAllRead();
  window.location.href = link;
}
async function markAllRead(){
  await fetch('/api/notifications/mark-read',{method:'POST'});
  document.getElementById('notifBadge').style.display='none';
  document.querySelectorAll('.notif-item').forEach(el=>el.classList.remove('unread'));
  document.querySelectorAll('.notif-unread-dot').forEach(el=>el.remove());
}
document.addEventListener('click', function(e){
  if(notifOpen && !document.getElementById('bellBtn').contains(e.target) && !document.getElementById('notifPanel').contains(e.target)){
    notifOpen = false;
    document.getElementById('notifPanel').classList.remove('open');
  }
});
// poll every 30s
loadNotifications();
setInterval(loadNotifications, 30000);
</script>
</body></html>"""


RECRUITER_HISTORY_HTML = """<!DOCTYPE html>
<html><head><title>History – RecruitAI</title>
<link href="https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;600;700&display=swap" rel="stylesheet">
<style>""" + SIDEBAR_CSS + """
.timeline-item{display:flex;gap:14px;padding:14px 0;border-bottom:1px solid rgba(255,255,255,0.04);}
.timeline-item:last-child{border-bottom:none;}
.tl-dot{width:36px;height:36px;border-radius:9px;display:flex;align-items:center;justify-content:center;font-size:16px;flex-shrink:0;margin-top:2px;}
.tl-content{flex:1;}
.tl-title{font-size:13px;font-weight:600;color:#e2e8f0;margin-bottom:3px;}
.tl-sub{font-size:12px;color:#3a3b4a;}
.tl-time{font-size:11px;color:#2a2b3a;white-space:nowrap;}
.tab-row{display:flex;gap:0;background:#0a0b12;border-radius:12px;padding:3px;border:1px solid rgba(255,255,255,0.05);margin-bottom:18px;width:fit-content;}
.tab-btn{padding:8px 20px;border:none;border-radius:9px;background:transparent;color:#555;font-size:13px;font-weight:600;cursor:pointer;transition:all .2s;font-family:'DM Sans',sans-serif;}
.tab-btn.active{background:#16172a;color:#c4b5fd;}
</style></head><body>
<div class="sidebar">
  <div class="logo">✦ Recruit<span>AI</span></div>
  <div class="nav-section">""" + _rec_nav("history") + """
  </div>
  <div class="sidebar-bottom">
    <div class="user-chip">
      <div class="avatar">{{ username[0].upper() }}</div>
      <div class="user-info"><div class="name">{{ username }}</div><div class="role">Recruiter</div></div>
    </div>
    <a href="/logout" class="logout-link">← Sign out</a>
  </div>
</div>
<div class="main">
  <div class="topbar">
    <div><h1>History</h1><div class="sub">All screening activity and application events</div></div>
  </div>
  <div class="content">
    <div class="tab-row">
      <button class="tab-btn active" onclick="showTab('screening',this)">🤖 Screened Resumes</button>
      <button class="tab-btn" onclick="showTab('applications',this)">📋 Application Events</button>
    </div>

    <!-- Screened Resumes -->
    <div id="tab-screening">
      {% if history %}
      <div class="card">
        <div class="card-header"><h2>Recently Screened Candidates</h2><span style="font-size:11px;color:#3a3b4a;">{{ history|length }} total</span></div>
        <div class="card-body">
          <table>
            <thead><tr><th>Candidate</th><th>File</th><th>ATS Score</th><th>Status</th><th>Screened</th></tr></thead>
            <tbody>
            {% for h in history %}
            {% set res = h.result_json|fromjson if h.result_json else {} %}
            <tr>
              <td><strong>{{ h.candidate_name }}</strong></td>
              <td style="color:#3a3b4a;font-size:12px;">{{ h.filename }}</td>
              <td>
                <span style="font-weight:700;color:{% if h.ats_score >= 70 %}#34d399{% elif h.ats_score >= 50 %}#fbbf24{% else %}#f87171{% endif %};">{{ h.ats_score }}%</span>
                <div style="background:#12131f;border-radius:99px;height:3px;width:80px;margin-top:3px;overflow:hidden;">
                  <div style="width:{{ h.ats_score }}%;height:100%;background:{% if h.ats_score >= 70 %}#34d399{% elif h.ats_score >= 50 %}#fbbf24{% else %}#f87171{% endif %};border-radius:99px;"></div>
                </div>
              </td>
              <td>
                {% if h.result_json %}
                  {% set rec = h.result_json %}
                  {% if 'Hire' in rec %}<span class="badge badge-green">Hire</span>
                  {% elif 'Maybe' in rec %}<span class="badge badge-yellow">Maybe</span>
                  {% else %}<span class="badge badge-red">Pass</span>{% endif %}
                {% else %}<span class="badge badge-purple">Screened</span>{% endif %}
              </td>
              <td style="color:#3a3b4a;font-size:12px;">{{ h.screened_at[:16].replace('T',' ') }}</td>
            </tr>
            {% endfor %}
            </tbody>
          </table>
        </div>
      </div>
      {% else %}
      <div class="empty-state" style="background:#0f1018;border-radius:16px;border:1px solid rgba(255,255,255,0.05);">
        <div class="e-icon">📄</div>
        <h3>No screened resumes yet</h3>
        <p><a href="/recruiter/screen-resume" style="color:#a78bfa;text-decoration:none;">Screen your first resume →</a></p>
      </div>
      {% endif %}
    </div>

    <!-- Application Events -->
    <div id="tab-applications" style="display:none;">
      <div class="card">
        <div class="card-header"><h2>Application Activity</h2><span style="font-size:11px;color:#3a3b4a;">{{ apps|length }} total</span></div>
        <div style="padding:8px 20px 16px;">
          {% if apps %}
            {% for app in apps %}
            <div class="timeline-item">
              <div class="tl-dot" style="background:{% if app.status=='Hired' %}rgba(52,211,153,0.12){% elif app.status=='Rejected' %}rgba(248,113,113,0.12){% elif app.status=='Shortlisted' %}rgba(99,102,241,0.12){% else %}rgba(251,191,36,0.12){% endif %};">
                {% if app.status=='Hired' %}✅{% elif app.status=='Rejected' %}❌{% elif app.status=='Shortlisted' %}⭐{% else %}📋{% endif %}
              </div>
              <div class="tl-content">
                <div class="tl-title">{{ app.username }} applied for <span style="color:#818cf8;">{{ app.job_title }}</span></div>
                <div class="tl-sub">Status:
                  {% if app.status=='Applied' %}<span class="badge badge-yellow">Pending</span>
                  {% elif app.status=='Shortlisted' %}<span class="badge badge-green">{{ app.status }}</span>
                  {% elif app.status=='Rejected' %}<span class="badge badge-red">{{ app.status }}</span>
                  {% elif app.status=='Hired' %}<span class="badge badge-blue">{{ app.status }}</span>
                  {% else %}<span class="badge badge-purple">{{ app.status }}</span>{% endif %}
                </div>
              </div>
              <div class="tl-time">{{ app.applied_at[:10] }}</div>
            </div>
            {% endfor %}
          {% else %}
          <div style="text-align:center;padding:40px;color:#2a2b3a;font-size:13px;">No application activity yet</div>
          {% endif %}
        </div>
      </div>
    </div>

  </div>
</div>
<script>
function showTab(id,btn){
  document.getElementById('tab-screening').style.display=id==='screening'?'block':'none';
  document.getElementById('tab-applications').style.display=id==='applications'?'block':'none';
  document.querySelectorAll('.tab-btn').forEach(b=>b.classList.remove('active'));
  btn.classList.add('active');
}
</script>
<script>
let notifOpen = false;
async function loadNotifications(){
  try{
    const res = await fetch('/api/notifications');
    const d = await res.json();
    const badge = document.getElementById('notifBadge');
    const list = document.getElementById('notifList');
    if(!badge||!list) return;
    if(d.unread > 0){
      badge.style.display='flex';
      badge.textContent = d.unread > 9 ? '9+' : d.unread;
    } else {
      badge.style.display='none';
    }
    if(!d.notifications || d.notifications.length === 0){
      list.innerHTML = '<div class="notif-empty">🔔 No notifications yet</div>';
      return;
    }
    const icons = {application:'📋', status:'📄', screening:'🤖', job:'💼'};
    const colors = {application:'rgba(99,102,241,0.15)', status:'rgba(52,211,153,0.12)', screening:'rgba(168,85,247,0.12)', job:'rgba(251,191,36,0.12)'};
    list.innerHTML = d.notifications.map(n => {
      const t = new Date(n.created_at).toLocaleString('en-IN',{day:'2-digit',month:'short',hour:'2-digit',minute:'2-digit'});
      return `<div class="notif-item ${n.is_read?'':'unread'}" onclick="goNotif('${n.link||'/'}')">
        <div class="notif-icon" style="background:${colors[n.type]||'rgba(99,102,241,0.12)'};">${icons[n.type]||'🔔'}</div>
        <div style="flex:1;">
          <div class="notif-msg">${n.message}</div>
          <div class="notif-time">${t}</div>
        </div>
        ${n.is_read?'':'<div class="notif-unread-dot"></div>'}
      </div>`;
    }).join('');
  } catch(e){}
}
function toggleNotif(){
  notifOpen = !notifOpen;
  document.getElementById('notifPanel').classList.toggle('open', notifOpen);
  if(notifOpen) loadNotifications();
}
function goNotif(link){
  markAllRead();
  window.location.href = link;
}
async function markAllRead(){
  await fetch('/api/notifications/mark-read',{method:'POST'});
  document.getElementById('notifBadge').style.display='none';
  document.querySelectorAll('.notif-item').forEach(el=>el.classList.remove('unread'));
  document.querySelectorAll('.notif-unread-dot').forEach(el=>el.remove());
}
document.addEventListener('click', function(e){
  if(notifOpen && !document.getElementById('bellBtn').contains(e.target) && !document.getElementById('notifPanel').contains(e.target)){
    notifOpen = false;
    document.getElementById('notifPanel').classList.remove('open');
  }
});
// poll every 30s
loadNotifications();
setInterval(loadNotifications, 30000);
</script>
</body></html>"""


JOB_BOARD_HTML = """<!DOCTYPE html>
<html><head><title>Browse Jobs – RecruitAI</title>
<link href="https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;600;700&display=swap" rel="stylesheet">
<style>""" + SIDEBAR_CSS + """
.search-bar{display:flex;gap:10px;margin-bottom:22px;}
.search-bar input{flex:1;padding:12px 16px;border:1px solid rgba(255,255,255,0.07);outline:none;border-radius:12px;background:#0f1018;color:white;font-size:14px;font-family:'DM Sans',sans-serif;}
.search-bar input:focus{border-color:#6366f1;}
.search-bar button{padding:12px 22px;border:none;border-radius:12px;background:linear-gradient(135deg,#6366f1,#8b5cf6);color:white;font-weight:600;font-size:14px;cursor:pointer;font-family:'DM Sans',sans-serif;}
.job-card{background:#0f1018;border-radius:16px;padding:22px;margin-bottom:12px;border:1px solid rgba(255,255,255,0.05);transition:border-color .2s;}
.job-card:hover{border-color:rgba(99,102,241,0.3);}
.job-card .jc-top{display:flex;justify-content:space-between;align-items:flex-start;gap:16px;}
.job-card h3{font-size:16px;font-weight:600;color:#e2e8f0;margin-bottom:3px;}
.job-card .jc-company{font-size:13px;color:#6366f1;margin-bottom:8px;}
.job-card .jc-meta{display:flex;gap:14px;flex-wrap:wrap;font-size:12px;color:#2a2b3a;margin-bottom:10px;}
.job-card .jc-skills{display:flex;flex-wrap:wrap;gap:6px;margin-bottom:12px;}
.skill-tag{background:rgba(99,102,241,0.1);color:#818cf8;padding:4px 10px;border-radius:99px;font-size:11px;font-weight:500;}
.jc-actions{display:flex;flex-direction:column;gap:8px;align-items:flex-end;flex-shrink:0;}
.apply-btn{padding:10px 20px;border:none;border-radius:11px;background:linear-gradient(135deg,#6366f1,#8b5cf6);color:white;font-weight:700;font-size:13px;cursor:pointer;font-family:'DM Sans',sans-serif;white-space:nowrap;}
.apply-btn.applied{background:#34d399;cursor:default;}
.match-btn{padding:8px 14px;border:1px solid rgba(99,102,241,0.25);border-radius:10px;background:transparent;color:#818cf8;font-size:12px;font-weight:600;cursor:pointer;font-family:'DM Sans',sans-serif;white-space:nowrap;}
.match-btn:hover{background:rgba(99,102,241,0.1);}
.jc-desc{font-size:13px;color:#2a2b3a;line-height:1.6;margin-bottom:10px;}
</style></head><body>

<div class="sidebar">
  <div class="logo">✦ Recruit<span>AI</span></div>
  <div class="nav-section">
    <div class="nav-label">Candidate</div>
    <a class="nav-item" href="/dashboard"><span class="icon">🏠</span> Dashboard</a>
    <a class="nav-item" href="/resume-builder"><span class="icon">📄</span> Resume Builder</a>
    <a class="nav-item active" href="/jobs"><span class="icon">💼</span> Browse Jobs</a>
    <a class="nav-item" href="/analytics"><span class="icon">📊</span> Analytics</a>
  </div>
  <div class="sidebar-bottom">
    <div class="user-chip">
      <div class="avatar">C</div>
      <div class="user-info"><div class="name">Candidate</div><div class="role">Job Seeker</div></div>
    </div>
    <a href="/logout" class="logout-link">← Sign out</a>
  </div>
</div>

<div class="main">
  <div class="topbar">
    <div><h1>Browse Jobs</h1><div class="sub">Find your next opportunity</div></div>
  </div>
  <div class="content">
    <form method="GET" class="search-bar">
      <input name="q" value="{{ query }}" placeholder="Search by title, company, or skill...">
      <button type="submit">🔍 Search</button>
    </form>

    {% if jobs %}
      {% for job in jobs %}
      <div class="job-card">
        <div class="jc-top">
          <div style="flex:1;">
            <h3>{{ job.title }}</h3>
            <div class="jc-company">{{ job.company }}</div>
            <div class="jc-meta">
              <span>📍 {{ job.location or 'Not specified' }}</span>
              <span>💼 {{ job.job_type }}</span>
              {% if job.salary %}<span>💰 {{ job.salary }}</span>{% endif %}
              <span>📅 {{ job.posted_at[:10] }}</span>
            </div>
            {% if job.skills_required %}
            <div class="jc-skills">
              {% for skill in job.skills_required.split(',') %}
              <span class="skill-tag">{{ skill.strip() }}</span>
              {% endfor %}
            </div>
            {% endif %}
            {% if job.description %}
            <div class="jc-desc">{{ job.description[:180] }}{% if job.description|length > 180 %}...{% endif %}</div>
            {% endif %}
          </div>
          <div class="jc-actions">
            <button class="apply-btn" id="apply-{{ job.id }}" onclick="applyJob({{ job.id }}, this)">Apply Now</button>
            <button class="match-btn" onclick="checkMatch({{ job.id }})">🤖 Match %</button>
          </div>
        </div>
      </div>
      {% endfor %}
    {% else %}
    <div class="empty-state" style="background:#0f1018;border-radius:16px;border:1px solid rgba(255,255,255,0.05);">
      <div class="e-icon">💼</div>
      <h3>{% if query %}No jobs found for "{{ query }}"{% else %}No jobs posted yet{% endif %}</h3>
      <p>{% if query %}Try a different search term{% else %}Check back soon!{% endif %}</p>
    </div>
    {% endif %}
  </div>
</div>

<div class="toast" id="toast"></div>

<div class="modal-overlay" id="modal">
  <div class="modal">
    <div class="modal-header">
      <h2>🤖 Job Match Analysis</h2>
      <button class="modal-close" onclick="closeModal()">×</button>
    </div>
    <div id="modalContent"></div>
  </div>
</div>

<script>
function showToast(msg){const t=document.getElementById('toast');t.textContent=msg;t.classList.add('show');setTimeout(()=>t.classList.remove('show'),2800);}
async function applyJob(jobId,btn){
  if(btn.classList.contains('applied'))return;
  try{
    const res=await fetch('/api/jobs/apply',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({job_id:jobId})});
    const d=await res.json();
    if(d.success){btn.textContent='✅ Applied';btn.classList.add('applied');showToast('✅ Application submitted!');}
    else showToast('⚠️ '+(d.error||'Failed to apply'));
  }catch(e){showToast('❌ Error applying');}
}
async function checkMatch(jobId){
  document.getElementById('modal').classList.add('show');
  document.getElementById('modalContent').innerHTML='<p style="color:#555;padding:12px 0;">🔄 Analysing match...</p>';
  try{
    const draftRes=await fetch('/api/resume-builder/load');
    const resume=await draftRes.json();
    const res=await fetch('/api/ai/job-match',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({resume,job_id:jobId})});
    const d=await res.json();
    document.getElementById('modalContent').innerHTML=`
      <div style="display:flex;align-items:center;gap:16px;margin-bottom:18px;">
        <div class="score-ring">${d.match_percent}%</div>
        <div><div style="font-size:13px;color:#6b7280;">Match with this job</div><div style="font-size:22px;font-weight:700;color:#e2e8f0;margin-top:4px;">${d.match_percent >= 70 ? 'Strong Fit ✨' : d.match_percent >= 50 ? 'Moderate Fit' : 'Needs Work'}</div></div>
      </div>
      <div style="margin-bottom:12px;"><div style="font-size:11px;color:#2a2b3a;margin-bottom:6px;text-transform:uppercase;letter-spacing:.5px;">✅ You have</div>${(d.matched_skills||[]).map(s=>`<span class="tag tag-g">${s}</span>`).join('')||'<span style="color:#2a2b3a;font-size:12px;">Load your resume first</span>'}</div>
      <div style="margin-bottom:16px;"><div style="font-size:11px;color:#2a2b3a;margin-bottom:6px;text-transform:uppercase;letter-spacing:.5px;">❌ You're missing</div>${(d.missing_skills||[]).map(s=>`<span class="tag tag-r">${s}</span>`).join('')||'<span style="color:#2a2b3a;font-size:12px;">None detected</span>'}</div>
      <div style="background:#0a0b12;padding:14px;border-radius:12px;font-size:13px;color:#6b7280;line-height:1.65;border:1px solid rgba(255,255,255,0.04);">💡 ${d.recommendation||''}</div>`;
  }catch(e){document.getElementById('modalContent').innerHTML='<p style="color:#f87171;">Match failed. Make sure your resume is saved.</p>';}
}
function closeModal(){document.getElementById('modal').classList.remove('show');}
document.getElementById('modal').addEventListener('click',function(e){if(e.target===this)closeModal();});
</script>
</body></html>"""


ANALYTICS_HTML = """<!DOCTYPE html>
<html><head><title>Analytics – RecruitAI</title>
<link href="https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;600;700&display=swap" rel="stylesheet">
<style>""" + SIDEBAR_CSS + """
.input-row{display:flex;gap:10px;}
.input-row input,.input-row textarea{flex:1;padding:11px 13px;border:1px solid rgba(255,255,255,0.07);outline:none;border-radius:11px;background:#0a0b12;color:white;font-size:13px;font-family:'DM Sans',sans-serif;}
.input-row input:focus,.input-row textarea:focus{border-color:#6366f1;}
.skill-tag{display:inline-flex;align-items:center;padding:5px 12px;border-radius:99px;font-size:12px;margin:3px;background:rgba(99,102,241,0.1);color:#818cf8;cursor:pointer;border:1px solid transparent;transition:all .2s;}
.skill-tag:hover{border-color:#6366f1;background:rgba(99,102,241,0.2);}
.score-circle{display:inline-flex;align-items:center;justify-content:center;width:88px;height:88px;border-radius:50%;border:3px solid #6366f1;font-size:26px;font-weight:700;color:#a78bfa;flex-shrink:0;}
.upload-zone{border:2px dashed rgba(99,102,241,0.3);border-radius:14px;padding:28px;text-align:center;cursor:pointer;transition:all .2s;background:#0a0b12;position:relative;}
.upload-zone:hover,.upload-zone.drag{border-color:#6366f1;background:rgba(99,102,241,0.05);}
.upload-zone input[type=file]{position:absolute;inset:0;opacity:0;cursor:pointer;width:100%;height:100%;}
.upload-zone .uz-icon{font-size:32px;margin-bottom:8px;}
.upload-zone .uz-title{font-size:14px;font-weight:600;color:#c4b5fd;margin-bottom:4px;}
.upload-zone .uz-sub{font-size:12px;color:#3a3b4a;}
.tab-row{display:flex;gap:0;background:#0a0b12;border-radius:12px;padding:3px;border:1px solid rgba(255,255,255,0.05);margin-bottom:16px;}
.tab-btn{flex:1;padding:9px;border:none;border-radius:9px;background:transparent;color:#555;font-size:13px;font-weight:600;cursor:pointer;transition:all .2s;font-family:'DM Sans',sans-serif;}
.tab-btn.active{background:#16172a;color:#c4b5fd;}
.prog-bar-wrap{background:#12131f;border-radius:99px;height:8px;overflow:hidden;margin-top:6px;}
.prog-bar{height:100%;border-radius:99px;background:linear-gradient(90deg,#6366f1,#a78bfa);transition:width .8s ease;}
</style></head><body>

<div class="sidebar">
  <div class="logo">✦ Recruit<span>AI</span></div>
  <div class="nav-section">
    <div class="nav-label">Candidate</div>
    <a class="nav-item" href="/dashboard"><span class="icon">🏠</span> Dashboard</a>
    <a class="nav-item" href="/resume-builder"><span class="icon">📄</span> Resume Builder</a>
    <a class="nav-item" href="/jobs"><span class="icon">💼</span> Browse Jobs</a>
    <a class="nav-item active" href="/analytics"><span class="icon">📊</span> Analytics</a>
  </div>
  <div class="sidebar-bottom">
    <div class="user-chip">
      <div class="avatar">{{ username[0].upper() }}</div>
      <div class="user-info"><div class="name">{{ username }}</div><div class="role">Job Seeker</div></div>
    </div>
    <a href="/logout" class="logout-link">← Sign out</a>
  </div>
</div>

<div class="main">
  <div class="topbar">
    <div><h1>Analytics</h1><div class="sub">Track your resume performance and job activity</div></div>
  </div>
  <div class="content">

    <div class="stats-grid">
      <div class="stat-card"><div class="s-icon">📥</div><div class="s-val">{{ stats.downloads }}</div><div class="s-label">Resume Downloads</div></div>
      <div class="stat-card"><div class="s-icon">🎯</div><div class="s-val">{{ stats.ats_score or '—' }}</div><div class="s-label">Last ATS Score</div></div>
      <div class="stat-card"><div class="s-icon">📨</div><div class="s-val">{{ applications|length }}</div><div class="s-label">Total Applications</div></div>
    </div>

    <!-- ATS SCORE CARD (full width, prominent) -->
    <div class="card" style="margin-bottom:20px;">
      <div class="card-header">
        <h2>📊 Resume ATS Score & Feedback</h2>
        <span style="font-size:11px;color:#3a3b4a;">Upload your existing resume PDF or use your builder resume</span>
      </div>
      <div style="padding:20px;">
        <div class="tab-row">
          <button class="tab-btn active" id="tab-upload" onclick="switchTab('upload')">📤 Upload PDF Resume</button>
          <button class="tab-btn" id="tab-builder" onclick="switchTab('builder')">✏️ Use Builder Resume</button>
        </div>

        <!-- Upload tab -->
        <div id="panel-upload">
          <div class="upload-zone" id="dropZone">
            <input type="file" id="resumeFile" accept=".pdf" onchange="onFileSelect(this)">
            <div class="uz-icon">📄</div>
            <div class="uz-title" id="uzTitle">Drop your resume PDF here</div>
            <div class="uz-sub">or click to browse &nbsp;•&nbsp; PDF only &nbsp;•&nbsp; Max 5MB</div>
          </div>
          <textarea id="jobDescUpload" rows="3" placeholder="Paste a job description here for a tailored score (optional)..."
            style="width:100%;margin-top:12px;padding:11px 13px;border:1px solid rgba(255,255,255,0.07);outline:none;border-radius:11px;background:#0a0b12;color:white;font-size:13px;font-family:'DM Sans',sans-serif;resize:vertical;"></textarea>
          <button class="btn btn-primary" style="width:100%;justify-content:center;margin-top:12px;padding:13px;" onclick="scoreUploadedResume()">🎯 Analyse Uploaded Resume</button>
        </div>

        <!-- Builder tab -->
        <div id="panel-builder" style="display:none;">
          <p style="font-size:13px;color:#3a3b4a;margin-bottom:12px;">Uses your saved Resume Builder data for the ATS analysis.</p>
          <textarea id="jobDescBuilder" rows="3" placeholder="Paste a job description here for a tailored score (optional)..."
            style="width:100%;padding:11px 13px;border:1px solid rgba(255,255,255,0.07);outline:none;border-radius:11px;background:#0a0b12;color:white;font-size:13px;font-family:'DM Sans',sans-serif;resize:vertical;"></textarea>
          <button class="btn btn-primary" style="width:100%;justify-content:center;margin-top:12px;padding:13px;" onclick="scoreBuilderResume()">🎯 Analyse Builder Resume</button>
        </div>

        <!-- Results area -->
        <div id="atsResult" style="margin-top:20px;display:none;"></div>
      </div>
    </div>

    <!-- AI SKILLS + APPLICATIONS ROW -->
    <div style="display:grid;grid-template-columns:1fr 1fr;gap:16px;margin-bottom:20px;">
      <div class="card">
        <div class="card-header"><h2>🤖 AI Skill Suggestions</h2></div>
        <div style="padding:18px;">
          <div class="input-row" style="margin-bottom:14px;">
            <input type="text" id="jobTitleInput" placeholder="Enter job title (e.g. Data Analyst)">
            <button class="btn btn-primary" onclick="suggestSkills()">Get Skills</button>
          </div>
          <div id="suggestedSkills" style="min-height:40px;"></div>
        </div>
      </div>
      <div class="card">
        <div class="card-header"><h2>📈 Application Summary</h2></div>
        <div style="padding:18px;">
          {% set total = applications|length %}
          {% set shortlisted = applications|selectattr('status','in',['Shortlisted','Reviewing'])|list|length %}
          {% set rejected = applications|selectattr('status','equalto','Rejected')|list|length %}
          {% set hired = applications|selectattr('status','equalto','Hired')|list|length %}
          <div style="margin-bottom:14px;">
            <div style="display:flex;justify-content:space-between;font-size:12px;color:#6b7280;margin-bottom:4px;"><span>Applications Sent</span><span style="color:#c4b5fd;font-weight:600;">{{ total }}</span></div>
            <div class="prog-bar-wrap"><div class="prog-bar" style="width:100%;"></div></div>
          </div>
          <div style="margin-bottom:14px;">
            <div style="display:flex;justify-content:space-between;font-size:12px;color:#6b7280;margin-bottom:4px;"><span>Shortlisted</span><span style="color:#34d399;font-weight:600;">{{ shortlisted }}</span></div>
            <div class="prog-bar-wrap"><div class="prog-bar" style="width:{{ ((shortlisted/total*100)|int if total else 0) }}%;background:linear-gradient(90deg,#34d399,#059669);"></div></div>
          </div>
          <div style="margin-bottom:14px;">
            <div style="display:flex;justify-content:space-between;font-size:12px;color:#6b7280;margin-bottom:4px;"><span>Rejected</span><span style="color:#f87171;font-weight:600;">{{ rejected }}</span></div>
            <div class="prog-bar-wrap"><div class="prog-bar" style="width:{{ ((rejected/total*100)|int if total else 0) }}%;background:linear-gradient(90deg,#f87171,#dc2626);"></div></div>
          </div>
          <div>
            <div style="display:flex;justify-content:space-between;font-size:12px;color:#6b7280;margin-bottom:4px;"><span>Hired</span><span style="color:#60a5fa;font-weight:600;">{{ hired }}</span></div>
            <div class="prog-bar-wrap"><div class="prog-bar" style="width:{{ ((hired/total*100)|int if total else 0) }}%;background:linear-gradient(90deg,#60a5fa,#2563eb);"></div></div>
          </div>
        </div>
      </div>
    </div>

    <div class="card">
      <div class="card-header">
        <h2>My Applications</h2>
        <a href="/jobs" class="btn btn-ghost btn-sm">Browse more →</a>
      </div>
      <div class="card-body">
        {% if applications %}
        <table>
          <thead><tr><th>Job Title</th><th>Company</th><th>Applied</th><th>Status</th></tr></thead>
          <tbody>
          {% for app in applications %}
          <tr>
            <td><strong>{{ app.job_title }}</strong></td>
            <td>{{ app.company }}</td>
            <td>{{ app.applied_at[:10] }}</td>
            <td>
              {% if app.status == 'Applied' %}<span class="badge badge-purple">{{ app.status }}</span>
              {% elif app.status in ['Shortlisted','Reviewing'] %}<span class="badge badge-green">{{ app.status }}</span>
              {% elif app.status == 'Rejected' %}<span class="badge badge-red">{{ app.status }}</span>
              {% elif app.status == 'Interviewing' %}<span class="badge badge-yellow">{{ app.status }}</span>
              {% elif app.status == 'Hired' %}<span class="badge badge-blue">{{ app.status }}</span>
              {% else %}<span class="badge badge-purple">{{ app.status }}</span>{% endif %}
            </td>
          </tr>
          {% endfor %}
          </tbody>
        </table>
        {% else %}
        <div class="empty-state"><div class="e-icon">📭</div><h3>No applications yet</h3><p><a href="/jobs" style="color:#a78bfa;text-decoration:none;">Browse jobs to apply →</a></p></div>
        {% endif %}
      </div>
    </div>

  </div>
</div>

<script>
function switchTab(t){
  document.getElementById('tab-upload').classList.toggle('active',t==='upload');
  document.getElementById('tab-builder').classList.toggle('active',t==='builder');
  document.getElementById('panel-upload').style.display=t==='upload'?'block':'none';
  document.getElementById('panel-builder').style.display=t==='builder'?'block':'none';
  document.getElementById('atsResult').style.display='none';
}

function onFileSelect(input){
  const f=input.files[0];
  if(f) document.getElementById('uzTitle').textContent='📄 '+f.name;
}

const dz=document.getElementById('dropZone');
dz.addEventListener('dragover',e=>{e.preventDefault();dz.classList.add('drag');});
dz.addEventListener('dragleave',()=>dz.classList.remove('drag'));
dz.addEventListener('drop',e=>{
  e.preventDefault();dz.classList.remove('drag');
  const f=e.dataTransfer.files[0];
  if(f&&f.type==='application/pdf'){
    document.getElementById('resumeFile').files=e.dataTransfer.files;
    document.getElementById('uzTitle').textContent='📄 '+f.name;
  }
});

function renderAtsResult(d){
  const resultEl=document.getElementById('atsResult');
  resultEl.style.display='block';
  const scoreColor=d.score>=75?'#34d399':d.score>=50?'#fbbf24':'#f87171';
  const sectionsFound=(d.sections_found||[]).map(s=>`<span class="tag tag-g">${s}</span>`).join('');
  const sectionsMissing=(d.sections_missing||[]).map(s=>`<span class="tag tag-r">${s}</span>`).join('');
  resultEl.innerHTML=`
    <div style="border-top:1px solid rgba(255,255,255,0.05);padding-top:20px;">
      <div style="display:flex;align-items:center;gap:22px;margin-bottom:20px;background:#0a0b12;padding:18px;border-radius:14px;border:1px solid rgba(255,255,255,0.04);">
        <div class="score-circle" style="border-color:${scoreColor};color:${scoreColor};">${d.score}</div>
        <div style="flex:1;">
          <div style="display:flex;align-items:baseline;gap:10px;margin-bottom:6px;">
            <div style="font-size:28px;font-weight:700;color:#e2e8f0;">${d.grade||'—'}</div>
            <div style="font-size:13px;color:#6b7280;">ATS Score / 100</div>
          </div>
          <div style="background:#12131f;border-radius:99px;height:6px;margin-bottom:8px;overflow:hidden;">
            <div style="width:${d.score}%;height:100%;border-radius:99px;background:linear-gradient(90deg,${scoreColor},${scoreColor}88);transition:width 1s ease;"></div>
          </div>
          <div style="font-size:12px;color:#555;">Keyword match: <span style="color:#a78bfa;font-weight:600;">${d.keyword_match||0}%</span></div>
        </div>
      </div>
      ${d.summary?`<div style="background:#0a0b12;padding:14px 16px;border-radius:12px;font-size:13px;color:#9ca3c0;line-height:1.7;margin-bottom:16px;border-left:3px solid #6366f1;">💬 ${d.summary}</div>`:''}
      <div style="display:grid;grid-template-columns:1fr 1fr;gap:14px;margin-bottom:14px;">
        <div><div style="font-size:11px;color:#2a2b3a;margin-bottom:7px;text-transform:uppercase;letter-spacing:.5px;font-weight:600;">✅ Strengths</div>${(d.strengths||[]).map(s=>`<div style="display:flex;align-items:flex-start;gap:8px;margin-bottom:6px;font-size:12px;color:#9ca3c0;"><span style="color:#34d399;margin-top:2px;">•</span>${s}</div>`).join('')}</div>
        <div><div style="font-size:11px;color:#2a2b3a;margin-bottom:7px;text-transform:uppercase;letter-spacing:.5px;font-weight:600;">⚠️ Improvements</div>${(d.improvements||[]).map(s=>`<div style="display:flex;align-items:flex-start;gap:8px;margin-bottom:6px;font-size:12px;color:#9ca3c0;"><span style="color:#fbbf24;margin-top:2px;">•</span>${s}</div>`).join('')}</div>
      </div>
      ${sectionsFound||sectionsMissing?`<div style="display:flex;gap:16px;flex-wrap:wrap;">
        ${sectionsFound?`<div><div style="font-size:11px;color:#2a2b3a;margin-bottom:5px;text-transform:uppercase;letter-spacing:.5px;">Sections Found</div>${sectionsFound}</div>`:''}
        ${sectionsMissing?`<div><div style="font-size:11px;color:#2a2b3a;margin-bottom:5px;text-transform:uppercase;letter-spacing:.5px;">Sections Missing</div>${sectionsMissing}</div>`:''}
      </div>`:''}
    </div>`;
}

async function scoreUploadedResume(){
  const fileInput=document.getElementById('resumeFile');
  if(!fileInput.files.length){
    alert('Please select a PDF file first.');return;
  }
  const resultEl=document.getElementById('atsResult');
  resultEl.style.display='block';
  resultEl.innerHTML='<p style="color:#555;padding:12px 0;">🔄 Extracting text and calculating ATS score...</p>';
  const formData=new FormData();
  formData.append('resume_pdf',fileInput.files[0]);
  formData.append('job_description',document.getElementById('jobDescUpload').value);
  try{
    const res=await fetch('/api/ai/ats-score-upload',{method:'POST',body:formData});
    const d=await res.json();
    if(d.error){resultEl.innerHTML=`<p style="color:#f87171;">${d.error}</p>`;return;}
    renderAtsResult(d);
  }catch(e){resultEl.innerHTML='<p style="color:#f87171;font-size:13px;">Analysis failed. Try again.</p>';}
}

async function scoreBuilderResume(){
  const resultEl=document.getElementById('atsResult');
  resultEl.style.display='block';
  resultEl.innerHTML='<p style="color:#555;padding:12px 0;">🔄 Loading your resume and calculating ATS score...</p>';
  try{
    const draftRes=await fetch('/api/resume-builder/load');
    const resume=await draftRes.json();
    if(!resume||!resume.name){
      resultEl.innerHTML='<p style="color:#f87171;font-size:13px;">No saved resume found. Please build and save your resume first in the <a href="/resume-builder" style="color:#a78bfa;">Resume Builder</a>.</p>';
      return;
    }
    const res=await fetch('/api/ai/ats-score',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({resume,job_description:document.getElementById('jobDescBuilder').value})});
    const d=await res.json();
    renderAtsResult({...d,sections_found:[],sections_missing:[]});
  }catch(e){resultEl.innerHTML='<p style="color:#f87171;font-size:13px;">Score failed. Save your resume first.</p>';}
}

async function suggestSkills(){
  const title=document.getElementById('jobTitleInput').value.trim();
  if(!title)return;
  const el=document.getElementById('suggestedSkills');
  el.innerHTML='<span style="color:#2a2b3a;font-size:13px;">🔄 Getting suggestions...</span>';
  try{
    const res=await fetch('/api/ai/suggest-skills',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({job_title:title})});
    const d=await res.json();
    el.innerHTML=(d.skills||[]).map(s=>`<span class="skill-tag" title="Click to copy">✦ ${s}</span>`).join('');
    el.querySelectorAll('.skill-tag').forEach(t=>{
      t.onclick=()=>{navigator.clipboard.writeText(t.textContent.replace('✦ ',''));t.style.borderColor='#34d399';t.style.color='#34d399';}
    });
  }catch(e){el.innerHTML='<span style="color:#f87171;font-size:13px;">AI unavailable</span>';}
}
</script>
</body></html>"""


RESUME_BUILDER_HTML = """<!DOCTYPE html>
<html lang="en"><head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Resume Builder – RecruitAI</title>
<link href="https://fonts.googleapis.com/css2?family=DM+Sans:wght@300;400;500;600;700&display=swap" rel="stylesheet">
<style>
*{margin:0;padding:0;box-sizing:border-box;}
body{font-family:'DM Sans',sans-serif;background:#08090e;color:white;display:flex;min-height:100vh;}
.sidebar{width:240px;min-width:240px;background:#0c0d15;border-right:1px solid rgba(255,255,255,0.05);display:flex;flex-direction:column;padding:24px 0;height:100vh;position:fixed;left:0;top:0;}
.logo{font-size:18px;font-weight:700;color:#a78bfa;padding:0 20px 24px;letter-spacing:-0.3px;border-bottom:1px solid rgba(255,255,255,0.05);}
.logo span{color:#6366f1;}
.nav-section{padding:20px 12px 8px;flex:1;}
.nav-label{font-size:10px;color:#2a2b3a;text-transform:uppercase;letter-spacing:1px;font-weight:600;padding:0 8px;margin-bottom:6px;}
.nav-item{display:flex;align-items:center;gap:10px;padding:10px 12px;border-radius:10px;color:#4a4b5e;font-size:13px;font-weight:500;text-decoration:none;transition:all .18s;margin-bottom:2px;}
.nav-item:hover{background:#12131f;color:#9ca3c0;}
.nav-item.active{background:#16172a;color:#c4b5fd;}
.nav-item .icon{width:18px;text-align:center;font-size:15px;}
.sidebar-bottom{padding:16px 12px;border-top:1px solid rgba(255,255,255,0.05);}
.user-chip{display:flex;align-items:center;gap:10px;padding:10px 12px;border-radius:10px;background:#12131f;}
.avatar{width:32px;height:32px;border-radius:8px;background:linear-gradient(135deg,#6366f1,#a78bfa);display:flex;align-items:center;justify-content:center;font-weight:700;font-size:13px;flex-shrink:0;}
.user-info .name{font-size:13px;font-weight:600;color:#ddd;}
.user-info .role{font-size:11px;color:#444;}
.logout-link{display:block;text-align:center;margin-top:8px;font-size:12px;color:#333;text-decoration:none;padding:8px;border-radius:8px;transition:all .2s;}
.logout-link:hover{background:#12131f;color:#f87171;}
.editor{width:380px;min-width:380px;background:#0c0d15;border-right:1px solid rgba(255,255,255,0.05);padding:20px;overflow-y:auto;height:100vh;position:fixed;left:240px;top:0;}
.preview-pane{margin-left:620px;flex:1;padding:36px 40px;overflow-y:auto;min-height:100vh;background:#0e0f18;}
.editor-title{font-size:14px;font-weight:700;color:#c4b5fd;margin-bottom:16px;letter-spacing:.3px;}
.e-card{background:#0f1018;border-radius:14px;padding:16px;margin-bottom:12px;border:1px solid rgba(255,255,255,0.04);}
.e-card h3{font-size:11px;font-weight:600;color:#6366f1;text-transform:uppercase;letter-spacing:.6px;margin-bottom:10px;}
label{display:block;margin-top:8px;margin-bottom:3px;font-size:11px;color:#444;font-weight:500;}
input,textarea{width:100%;padding:9px 11px;border:1px solid rgba(255,255,255,0.06);outline:none;border-radius:9px;background:#0a0b12;color:white;font-size:12px;font-family:'DM Sans',sans-serif;transition:border-color .2s;}
input:focus,textarea:focus{border-color:#6366f1;}
textarea{resize:vertical;}
.add-btn{width:100%;margin-top:8px;padding:8px;border:none;border-radius:9px;background:#6366f1;color:white;font-weight:600;font-size:12px;cursor:pointer;font-family:'DM Sans',sans-serif;transition:opacity .2s;}
.add-btn:hover{opacity:.9;}
.ai-btn{width:100%;margin-top:5px;padding:8px;border:1px solid rgba(99,102,241,0.3);border-radius:9px;background:transparent;color:#818cf8;font-weight:600;font-size:11px;cursor:pointer;font-family:'DM Sans',sans-serif;transition:all .2s;}
.ai-btn:hover{background:rgba(99,102,241,0.1);}
.entry{background:#0a0b12;padding:8px 11px;border-radius:9px;margin-top:5px;font-size:12px;display:flex;justify-content:space-between;align-items:center;gap:8px;border:1px solid rgba(255,255,255,0.04);}
.entry span{flex:1;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;color:#9ca3c0;}
.del-btn{background:none;border:none;color:#3a3b4a;font-size:16px;cursor:pointer;flex-shrink:0;line-height:1;padding:0 2px;}
.del-btn:hover{color:#f87171;}
.action-bar{display:flex;gap:8px;margin-top:12px;padding-top:12px;border-top:1px solid rgba(255,255,255,0.05);}
.action-bar button{flex:1;padding:10px 6px;border:none;border-radius:10px;font-weight:700;font-size:11px;cursor:pointer;font-family:'DM Sans',sans-serif;transition:all .2s;}
.a-save{background:#12131f;color:#9ca3c0;border:1px solid rgba(255,255,255,0.06);}
.a-save:hover{border-color:#6366f1;color:#c4b5fd;}
.a-pdf{background:linear-gradient(135deg,#6366f1,#8b5cf6);color:white;}
.a-jobs{background:#16172a;color:#a78bfa;border:1px solid rgba(99,102,241,0.2);}
.a-jobs:hover{background:#1d1e35;}
.preview{max-width:780px;margin:auto;background:white;color:#111;border-radius:20px;padding:50px 55px;min-height:900px;box-shadow:0 24px 80px rgba(0,0,0,0.5);}
.resume-name{font-size:30px;font-weight:700;color:#111;letter-spacing:-0.5px;}
.resume-title{font-size:15px;color:#666;margin-top:5px;}
.resume-contact{margin-top:7px;color:#888;font-size:12px;}
.divider{border:none;border-top:2px solid #6366f1;margin:14px 0 0;}
.section{margin-top:20px;}
.section h3{color:#6366f1;margin-bottom:7px;font-size:13px;font-weight:700;text-transform:uppercase;letter-spacing:.6px;border-bottom:1px solid #ede9ff;padding-bottom:4px;}
.item{margin-bottom:11px;}
.item-title{font-weight:700;font-size:14px;color:#111;}
.item-sub{color:#888;margin-top:2px;font-size:12px;}
.item p{margin-top:4px;font-size:12px;line-height:1.55;color:#444;}
.skill-tag{display:inline-block;background:#ede9ff;color:#5b3fcf;padding:3px 10px;border-radius:99px;margin:3px;font-size:11px;font-weight:500;}
.toast{position:fixed;bottom:20px;right:20px;background:#1a1b2e;color:#c4b5fd;padding:11px 16px;border-radius:11px;font-weight:500;font-size:12px;opacity:0;transition:opacity .3s;pointer-events:none;z-index:9999;border:1px solid rgba(99,102,241,0.3);}
.toast.show{opacity:1;}
</style></head><body>

<div class="sidebar">
  <div class="logo">✦ Recruit<span>AI</span></div>
  <div class="nav-section">
    <div class="nav-label">Candidate</div>
    <a class="nav-item" href="/dashboard"><span class="icon">🏠</span> Dashboard</a>
    <a class="nav-item active" href="/resume-builder"><span class="icon">📄</span> Resume Builder</a>
    <a class="nav-item" href="/jobs"><span class="icon">💼</span> Browse Jobs</a>
    <a class="nav-item" href="/analytics"><span class="icon">📊</span> Analytics</a>
  </div>
  <div class="sidebar-bottom">
    <div class="user-chip">
      <div class="avatar">C</div>
      <div class="user-info"><div class="name">Candidate</div><div class="role">Job Seeker</div></div>
    </div>
    <a href="/logout" class="logout-link">← Sign out</a>
  </div>
</div>

<div class="editor">
  <div class="editor-title">✏️ Edit Resume</div>

  <div class="e-card" style="border-color:rgba(99,102,241,0.2);">
    <h3>⬆ Upload Existing Resume</h3>
    <div id="uploadZone" style="border:2px dashed rgba(99,102,241,0.3);border-radius:10px;padding:16px;text-align:center;cursor:pointer;transition:all .2s;position:relative;" onclick="document.getElementById('resumeUploadInput').click()">
      <div style="font-size:22px;">📄</div>
      <div id="uploadLabel" style="font-size:12px;color:#6b7280;margin-top:4px;">Click to upload PDF or DOCX</div>
      <input type="file" id="resumeUploadInput" accept=".pdf,.docx,.doc" style="display:none;" onchange="handleResumeUpload(this)">
    </div>
    <div id="uploadStatus" style="font-size:11px;margin-top:6px;color:#6b7280;"></div>
  </div>

  <div class="e-card">
    <h3>Personal Info</h3>
    <label>Full Name</label><input type="text" id="name" placeholder="Rahul Sharma" oninput="renderResume()">
    <label>Headline</label><input type="text" id="headline" placeholder="Full Stack Developer" oninput="renderResume()">
    <label>Email</label><input type="email" id="email" placeholder="you@email.com" oninput="renderResume()">
    <label>Phone</label><input type="text" id="phone" placeholder="+91 98765 43210" oninput="renderResume()">
    <label>Location</label><input type="text" id="location" placeholder="Mumbai, India" oninput="renderResume()">
    <label>Summary</label><textarea rows="3" id="summary" placeholder="Brief professional summary..." oninput="renderResume()"></textarea>
  </div>

  <div class="e-card">
    <h3>Skills</h3>
    <input type="text" id="skillInput" placeholder="Type a skill and press Enter">
    <button class="add-btn" onclick="addSkill()">+ Add Skill</button>
    <button class="ai-btn" onclick="aiSuggestSkills()">🤖 AI Suggest Skills</button>
    <div id="skillsList"></div>
  </div>

  <div class="e-card">
    <h3>Experience</h3>
    <input type="text" id="jobTitle" placeholder="Job Title">
    <input type="text" id="company" placeholder="Company Name" style="margin-top:6px;">
    <input type="text" id="jobStart" placeholder="Start (Jan 2022)" style="margin-top:6px;">
    <input type="text" id="jobEnd" placeholder="End / Present" style="margin-top:6px;">
    <textarea id="jobDesc" rows="2" placeholder="Describe your role..." style="margin-top:6px;"></textarea>
    <button class="add-btn" onclick="addExperience()">+ Add Experience</button>
    <div id="experienceList"></div>
  </div>

  <div class="e-card">
    <h3>Education</h3>
    <input type="text" id="degree" placeholder="B.Tech CSE">
    <input type="text" id="college" placeholder="College / University" style="margin-top:6px;">
    <input type="text" id="eduFrom" placeholder="From (2018)" style="margin-top:6px;">
    <input type="text" id="eduTo" placeholder="To (2022)" style="margin-top:6px;">
    <button class="add-btn" onclick="addEducation()">+ Add Education</button>
    <div id="educationList"></div>
  </div>

  <div class="e-card">
    <h3>Projects</h3>
    <input type="text" id="projectName" placeholder="Project Name">
    <input type="text" id="projectTech" placeholder="Technologies" style="margin-top:6px;">
    <textarea id="projectDesc" rows="2" placeholder="What did you build?" style="margin-top:6px;"></textarea>
    <button class="add-btn" onclick="addProject()">+ Add Project</button>
    <div id="projectList"></div>
  </div>

  <div class="e-card">
    <h3>Languages</h3>
    <input type="text" id="languageInput" placeholder="e.g. English, Hindi">
    <button class="add-btn" onclick="addLanguage()">+ Add Language</button>
    <div id="languageList"></div>
  </div>

  <div class="e-card">
    <h3>Certifications</h3>
    <input type="text" id="certInput" placeholder="e.g. AWS Certified Developer">
    <button class="add-btn" onclick="addCertification()">+ Add Certification</button>
    <div id="certList"></div>
  </div>

  <div class="e-card">
    <h3>🤖 AI Career Tools</h3>
    <label>Target Role</label>
    <input type="text" id="careerRole" placeholder="e.g. Software Engineer">
    <button class="add-btn" style="margin-top:8px;" onclick="getProjects()">🚀 Suggest Portfolio Projects</button>
    <button class="add-btn" style="margin-top:6px;background:#059669;" onclick="getRoadmap()">🗺 Generate Career Roadmap</button>
    <div id="aiOutput" style="margin-top:10px;background:#0a0b12;border:1px solid rgba(255,255,255,0.06);border-radius:9px;padding:12px;font-size:11px;color:#6b7280;line-height:1.7;min-height:50px;display:none;"></div>
  </div>

  <div class="action-bar">
    <button class="a-save" onclick="saveDraft()">💾 SAVE</button>
    <button class="a-pdf" onclick="downloadResume()">📄 PDF</button>
    <button class="a-jobs" onclick="window.location='/jobs'">💼 JOBS</button>
  </div>
</div>

<div class="preview-pane">
  <div class="preview" id="preview"></div>
</div>

<div class="toast" id="toast"></div>

<script>
let resume={skills:[],experiences:[],education:[],projects:[],languages:[],certifications:[]};

function showToast(msg){const t=document.getElementById('toast');t.textContent=msg;t.classList.add('show');setTimeout(()=>t.classList.remove('show'),2800);}

function renderResume(){
  const name=document.getElementById('name').value||'Your Name';
  const headline=document.getElementById('headline').value;
  const email=document.getElementById('email').value;
  const phone=document.getElementById('phone').value;
  const location=document.getElementById('location').value;
  const summary=document.getElementById('summary').value;
  const contactParts=[email,phone,location].filter(Boolean);
  let html=`<div class="resume-name">${name}</div>
    ${headline?`<div class="resume-title">${headline}</div>`:''}
    ${contactParts.length?`<div class="resume-contact">${contactParts.join(' &nbsp;•&nbsp; ')}</div>`:''}
    <hr class="divider">`;
  if(summary)html+=`<div class="section"><h3>Summary</h3><p style="font-size:12px;line-height:1.65;color:#444;">${summary}</p></div>`;
  if(resume.skills.length){html+=`<div class="section"><h3>Skills</h3>`;resume.skills.forEach(s=>{html+=`<span class="skill-tag">${s}</span>`;});html+=`</div>`;}
  if(resume.experiences.length){html+=`<div class="section"><h3>Experience</h3>`;resume.experiences.forEach(e=>{html+=`<div class="item"><div class="item-title">${e.title}</div><div class="item-sub">${e.company}${e.start?' &nbsp;|&nbsp; '+e.start+' – '+(e.end||'Present'):''}</div>${e.desc?`<p>${e.desc}</p>`:''}</div>`;});html+=`</div>`;}
  if(resume.education.length){html+=`<div class="section"><h3>Education</h3>`;resume.education.forEach(e=>{html+=`<div class="item"><div class="item-title">${e.degree}</div><div class="item-sub">${e.college}${e.from?' &nbsp;|&nbsp; '+e.from+' – '+(e.to||''):''}</div></div>`;});html+=`</div>`;}
  if(resume.projects.length){html+=`<div class="section"><h3>Projects</h3>`;resume.projects.forEach(p=>{html+=`<div class="item"><div class="item-title">${p.name}</div>${p.tech?`<div class="item-sub">${p.tech}</div>`:''} ${p.desc?`<p>${p.desc}</p>`:''}</div>`;});html+=`</div>`;}
  if(resume.languages.length){html+=`<div class="section"><h3>Languages</h3>`;resume.languages.forEach(l=>{html+=`<span class="skill-tag">${l}</span>`;});html+=`</div>`;}
  if(resume.certifications.length){html+=`<div class="section"><h3>Certifications</h3>`;resume.certifications.forEach(c=>{html+=`<div class="item" style="font-size:12px;color:#333;">• ${c}</div>`;});html+=`</div>`;}
  document.getElementById('preview').innerHTML=html;
}

async function aiSuggestSkills(){
  const title=document.getElementById('headline').value.trim()||'Software Developer';
  showToast('🤖 Getting AI suggestions...');
  try{
    const res=await fetch('/api/ai/suggest-skills',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({job_title:title})});
    const d=await res.json();
    if(d.skills){d.skills.forEach(s=>{if(!resume.skills.includes(s))resume.skills.push(s);});renderSkillList();renderResume();showToast(`✅ Added ${d.skills.length} skills!`);}
  }catch(e){showToast('❌ AI unavailable');}
}

function addSkill(){const v=document.getElementById('skillInput').value.trim();if(!v)return;resume.skills.push(v);document.getElementById('skillInput').value='';renderSkillList();renderResume();}
function renderSkillList(){document.getElementById('skillsList').innerHTML=resume.skills.map((s,i)=>`<div class="entry"><span>${s}</span><button class="del-btn" onclick="removeItem('skills',${i})">×</button></div>`).join('');}
function addExperience(){const title=document.getElementById('jobTitle').value.trim();if(!title){showToast('⚠️ Job Title required');return;}resume.experiences.push({title,company:document.getElementById('company').value,start:document.getElementById('jobStart').value,end:document.getElementById('jobEnd').value,desc:document.getElementById('jobDesc').value});['jobTitle','company','jobStart','jobEnd','jobDesc'].forEach(id=>document.getElementById(id).value='');renderExperienceList();renderResume();}
function renderExperienceList(){document.getElementById('experienceList').innerHTML=resume.experiences.map((e,i)=>`<div class="entry"><span>${e.title} @ ${e.company}</span><button class="del-btn" onclick="removeItem('experiences',${i})">×</button></div>`).join('');}
function addEducation(){const degree=document.getElementById('degree').value.trim();if(!degree){showToast('⚠️ Degree required');return;}resume.education.push({degree,college:document.getElementById('college').value,from:document.getElementById('eduFrom').value,to:document.getElementById('eduTo').value});['degree','college','eduFrom','eduTo'].forEach(id=>document.getElementById(id).value='');renderEducationList();renderResume();}
function renderEducationList(){document.getElementById('educationList').innerHTML=resume.education.map((e,i)=>`<div class="entry"><span>${e.degree} — ${e.college}</span><button class="del-btn" onclick="removeItem('education',${i})">×</button></div>`).join('');}
function addProject(){const name=document.getElementById('projectName').value.trim();if(!name){showToast('⚠️ Project name required');return;}resume.projects.push({name,tech:document.getElementById('projectTech').value,desc:document.getElementById('projectDesc').value});['projectName','projectTech','projectDesc'].forEach(id=>document.getElementById(id).value='');renderProjectList();renderResume();}
function renderProjectList(){document.getElementById('projectList').innerHTML=resume.projects.map((p,i)=>`<div class="entry"><span>${p.name}</span><button class="del-btn" onclick="removeItem('projects',${i})">×</button></div>`).join('');}
function addLanguage(){const v=document.getElementById('languageInput').value.trim();if(!v)return;resume.languages.push(v);document.getElementById('languageInput').value='';renderLanguageList();renderResume();}
function renderLanguageList(){document.getElementById('languageList').innerHTML=resume.languages.map((l,i)=>`<div class="entry"><span>${l}</span><button class="del-btn" onclick="removeItem('languages',${i})">×</button></div>`).join('');}
function addCertification(){const v=document.getElementById('certInput').value.trim();if(!v)return;resume.certifications.push(v);document.getElementById('certInput').value='';renderCertList();renderResume();}
function renderCertList(){document.getElementById('certList').innerHTML=resume.certifications.map((c,i)=>`<div class="entry"><span>${c}</span><button class="del-btn" onclick="removeItem('certifications',${i})">×</button></div>`).join('');}
function removeItem(section,index){resume[section].splice(index,1);const r={skills:renderSkillList,experiences:renderExperienceList,education:renderEducationList,projects:renderProjectList,languages:renderLanguageList,certifications:renderCertList};r[section]();renderResume();}

document.getElementById('skillInput').addEventListener('keydown',e=>{if(e.key==='Enter')addSkill();});
document.getElementById('languageInput').addEventListener('keydown',e=>{if(e.key==='Enter')addLanguage();});
document.getElementById('certInput').addEventListener('keydown',e=>{if(e.key==='Enter')addCertification();});

async function saveDraft(){
  const data={name:document.getElementById('name').value,headline:document.getElementById('headline').value,email:document.getElementById('email').value,phone:document.getElementById('phone').value,location:document.getElementById('location').value,summary:document.getElementById('summary').value,...resume};
  try{const res=await fetch('/api/resume-builder/save',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(data)});const j=await res.json();showToast(j.success?'✅ Saved!':'❌ Save failed');}catch(e){showToast('❌ Save failed');}
}

async function handleResumeUpload(input){
  const file = input.files[0];
  if(!file) return;
  document.getElementById('uploadLabel').textContent = file.name;
  document.getElementById('uploadStatus').textContent = '⏳ Parsing resume with AI...';
  const fd = new FormData();
  fd.append('resume', file);
  try {
    const res = await fetch('/api/resume-builder/parse', {method:'POST', body:fd});
    const d = await res.json();
    if(!d.success){ document.getElementById('uploadStatus').textContent = '❌ ' + (d.error||'Parse failed'); return; }
    const p = d.data;
    // Personal info
    if(p.name) document.getElementById('name').value = p.name;
    if(p.title) document.getElementById('headline').value = p.title;
    if(p.email) document.getElementById('email').value = p.email;
    if(p.phone) document.getElementById('phone').value = p.phone;
    if(p.location) document.getElementById('location').value = p.location;
    if(p.summary) document.getElementById('summary').value = p.summary;
    // Skills
    if(p.skill && p.skill.length){
      p.skill.forEach(s=>{ const name=typeof s==='string'?s:s.name; if(name && !resume.skills.includes(name)) resume.skills.push(name); });
      renderSkillList();
    }
    // Experience
    if(p.exp && p.exp.length){
      resume.experiences = p.exp.map(e=>({title:e.title||'',company:e.company||'',start:e.start||'',end:e.end||'',desc:e.desc||''}));
      renderExperienceList();
    }
    // Education — parser returns `institution`, form uses `college`
    if(p.edu && p.edu.length){
      resume.education = p.edu.map(e=>({degree:e.degree||'',college:e.institution||e.college||'',from:e.start||'',to:e.end||''}));
      renderEducationList();
    }
    // Projects
    if(p.proj && p.proj.length){
      resume.projects = p.proj.map(pr=>({name:pr.title||pr.name||'',tech:pr.tech||'',desc:pr.desc||''}));
      renderProjectList();
    }
    // Languages
    if(p.lang && p.lang.length){
      resume.languages = p.lang.map(l=>typeof l==='string'?l:l.name).filter(Boolean);
      renderLanguageList();
    }
    // Certifications
    if(p.cert && p.cert.length){
      resume.certifications = p.cert.map(c=>typeof c==='string'?c:(c.title+(c.issuer?' — '+c.issuer:''))).filter(Boolean);
      renderCertList();
    }
    renderResume();
    document.getElementById('uploadStatus').textContent = '✅ Resume loaded! Review and edit below.';
    showToast('✅ Resume parsed successfully!');
  } catch(e){
    document.getElementById('uploadStatus').textContent = '❌ Upload failed. Try again.';
  }
}

async function downloadResume(){
  const data={name:document.getElementById('name').value,headline:document.getElementById('headline').value,email:document.getElementById('email').value,phone:document.getElementById('phone').value,location:document.getElementById('location').value,summary:document.getElementById('summary').value,skills:(state.skill||[]).map(s=>typeof s==='string'?s:s.name),experiences:(state.exp||[]),education:(state.edu||[]).map(e=>({degree:e.degree||'',college:e.institution||e.college||'',from:e.start||'',to:e.end||'',desc:e.desc||''})),projects:(state.proj||[]).map(p=>({name:p.title||p.name||'',tech:p.tech||'',desc:p.desc||''})),languages:(state.lang||[]).map(l=>typeof l==='string'?l:l.name),certifications:(state.cert||[]).map(c=>typeof c==='string'?c:(c.title+(c.issuer?' — '+c.issuer:'')))};
  showToast('⏳ Generating PDF...');
  try{
    const response=await fetch('/api/resume-builder/download',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(data)});
    if(!response.ok){showToast('❌ PDF failed');return;}
    const blob=await response.blob();const url=window.URL.createObjectURL(blob);const a=document.createElement('a');a.href=url;a.download='resume.pdf';a.click();window.URL.revokeObjectURL(url);
    showToast('✅ PDF Downloaded!');
  }catch(e){showToast('❌ PDF failed');}
}

async function loadDraft(){
  try{
    const res=await fetch('/api/resume-builder/load');const data=await res.json();
    if(!data||!data.name)return;
    ['name','headline','email','phone','location','summary'].forEach(id=>{document.getElementById(id).value=data[id]||'';});
    ['skills','experiences','education','projects','languages','certifications'].forEach(k=>{if(data[k])resume[k]=data[k];});
    renderSkillList();renderExperienceList();renderEducationList();renderProjectList();renderLanguageList();renderCertList();renderResume();
    showToast('📄 Draft loaded!');
  }catch(e){}
}
loadDraft();renderResume();

async function getProjects(){
  const role=document.getElementById('careerRole').value.trim()||document.getElementById('headline').value.trim()||'Software Developer';
  const out=document.getElementById('aiOutput');
  out.style.display='block';out.innerHTML='🔄 Getting AI project suggestions...';out.style.color='#555';
  try{
    const res=await fetch('/api/ai/project-ideas',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({role})});
    const d=await res.json();
    out.style.color='#9ca3c0';
    out.innerHTML=(d.projects||[]).map(p=>`<div style="margin-bottom:8px;padding:8px;background:#12131f;border-radius:6px;border-left:2px solid #6366f1;">${p}</div>`).join('');
  }catch(e){out.innerHTML='❌ Failed to get suggestions';out.style.color='#f87171';}
}

async function getRoadmap(){
  const role=document.getElementById('careerRole').value.trim()||document.getElementById('headline').value.trim()||'Software Developer';
  const out=document.getElementById('aiOutput');
  out.style.display='block';out.innerHTML='🔄 Generating your career roadmap...';out.style.color='#555';
  try{
    const res=await fetch('/api/ai/career-roadmap',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({role})});
    const d=await res.json();
    out.style.color='#9ca3c0';out.style.whiteSpace='pre-wrap';
    out.textContent=d.roadmap||'No roadmap generated';
  }catch(e){out.innerHTML='❌ Failed to generate roadmap';out.style.color='#f87171';}
}
</script></body></html>"""

# ============================================================
# RUN
# ============================================================

init_tables()

# ── HELPER: extract raw text from uploaded file ──
def extract_text_from_file(filepath, filename):
    ext = filename.rsplit('.', 1)[-1].lower()
    print(f"=== EXTRACTING: {filename}, ext={ext}, path={filepath} ===")
    if ext == 'pdf':
        doc = fitz.open(filepath)
        print(f"=== PDF PAGES: {len(doc)} ===")
        text = "\n".join(page.get_text() for page in doc)
        print(f"=== TEXT LENGTH: {len(text)} ===")
        print(f"=== TEXT PREVIEW: {repr(text[:300])} ===")
        if not text.strip():
            text = ""
            for page in doc:
                blocks = page.get_text("blocks")
                for block in blocks:
                    if block[6] == 0:
                        text += block[4] + "\n"
        if not text.strip():
            try:
                from pdfminer.high_level import extract_text
                text = extract_text(filepath)
            except:
                pass
        return text.strip()
    elif ext in ('docx', 'doc'):
        d = docxlib.Document(filepath)
        parts = []
        for p in d.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)
    return ""
# ── ROUTE: parse uploaded resume ──
@app.route('/api/resume-builder/parse', methods=['POST'])
def parse_resume():
    if 'resume' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400

    file = request.files['resume']
    if not file.filename:
        return jsonify({'error': 'Empty filename'}), 400

    allowed = {'pdf', 'docx', 'doc'}
    ext = file.filename.rsplit('.', 1)[-1].lower()
    if ext not in allowed:
        return jsonify({'error': 'Only PDF or DOCX allowed'}), 400

    import tempfile
    with tempfile.NamedTemporaryFile(delete=False, suffix='.' + ext) as tmp:
        file.save(tmp.name)
        tmp_path = tmp.name

    try:
        raw_text = extract_text_from_file(tmp_path, file.filename)
    finally:
        os.unlink(tmp_path)

    if not raw_text.strip():
        return jsonify({'error': 'Could not extract text. Please use a DOCX file or a text-based PDF (not scanned/image).'}), 400

    prompt = f"""
You are a resume parser. Extract all information from the resume text below and return ONLY valid JSON — no markdown, no explanation, just the raw JSON object.

Return this exact structure:
{{
  "name": "",
  "title": "",
  "email": "",
  "phone": "",
  "location": "",
  "linkedin": "",
  "github": "",
  "website": "",
  "summary": "",
  "exp": [
    {{"title": "", "company": "", "start": "", "end": "", "desc": ""}}
  ],
  "edu": [
    {{"degree": "", "institution": "", "start": "", "end": "", "desc": ""}}
  ],
  "cert": [
    {{"title": "", "issuer": "", "date": "", "id": ""}}
  ],
  "skill": [
    {{"name": "", "level": ""}}
  ],
  "lang": [
    {{"name": "", "level": ""}}
  ],
  "proj": [
    {{"title": "", "tech": "", "desc": ""}}
  ],
  "extra": [
    {{"title": "", "org": "", "type": "✨ Other", "year": "", "desc": ""}}
  ]
}}

Resume text:
{raw_text[:6000]}
"""

    try:
        raw = gemini(prompt)
        raw = re.sub(r'^```[a-z]*\n?', '', raw.strip())
        raw = re.sub(r'\n?```$', '', raw)
        parsed = json.loads(raw)
        return jsonify({'success': True, 'data': parsed})
    except Exception as e:
        return jsonify({'error': f'AI parse failed: {str(e)}'}), 500

# ── ADD THESE ROUTES TO app.py (just above if __name__ == '__main__':) ──

@app.route('/job-finder')
def job_finder():
    return render_template('jobseeker/job-finder.html')


@app.route('/api/job-finder/search', methods=['POST'])
def job_finder_search():
    if "user" not in session:
        return jsonify({'error': 'Not logged in'}), 401
    data = request.get_json()
    query = data.get('query', 'developer')
    location = data.get('location', '').strip() or 'India'
    query_words = [w.lower() for w in query.split() if len(w) > 2]

    def score(title, desc):
        t, d = title.lower(), desc.lower()
        matches = sum(1 for w in query_words if w in t or w in d)
        return min(95, 50 + (matches * 18)) if matches > 0 else 35

    jobs = []

    # JSearch — aggregates Indeed, LinkedIn, Glassdoor, Naukri etc.
    try:
        jsearch_key = os.environ.get('JSEARCH_KEY', '')
        resp = requests.get(
            'https://jsearch.p.rapidapi.com/search',
            headers={
                'X-RapidAPI-Key': jsearch_key,
                'X-RapidAPI-Host': 'jsearch.p.rapidapi.com'
            },
            params={
                'query': f"{query} in {location}",
                'num_pages': '2',
                'date_posted': 'month'
            },
            timeout=15
        )
        for j in resp.json().get('data', []):
            title = j.get('job_title', '')
            desc  = j.get('job_description', '')
            salary = ''
            if j.get('job_min_salary') and j.get('job_max_salary'):
                salary = f"{j['job_currency'] or '₹'}{int(j['job_min_salary']):,} - {int(j['job_max_salary']):,}"
            jobs.append({
                'title': title,
                'company': j.get('employer_name', ''),
                'location': j.get('job_city', '') or j.get('job_country', '') or location,
                'url': j.get('job_apply_link', '') or j.get('job_google_link', ''),
                'description': desc[:300],
                'created': (j.get('job_posted_at_datetime_utc') or '')[:10],
                'salary': salary,
                'fit_score': score(title, desc),
                'source': j.get('job_publisher', '')
            })
    except Exception as e:
        print(f"JSearch error: {e}")

    # Fallback — Remotive for remote jobs
    try:
        resp2 = requests.get(
            'https://remotive.com/api/remote-jobs',
            params={'search': query, 'limit': 10},
            timeout=15
        )
        for j in resp2.json().get('jobs', []):
            title = j.get('title', '')
            desc  = j.get('description', '')
            jobs.append({
                'title': title,
                'company': j.get('company_name', ''),
                'location': j.get('candidate_required_location', 'Worldwide Remote'),
                'url': j.get('url', ''),
                'description': desc[:300],
                'created': j.get('publication_date', '')[:10],
                'salary': j.get('salary', ''),
                'fit_score': score(title, desc),
                'source': 'Remotive'
            })
    except Exception as e:
        print(f"Remotive error: {e}")

    if not jobs:
        return jsonify({'error': 'No jobs found. Try different keywords.'}), 404

    jobs.sort(key=lambda x: x['fit_score'], reverse=True)
    return jsonify({'success': True, 'jobs': jobs[:30], 'total': len(jobs)})

@app.route('/api/job-finder/apply', methods=['POST'])
def job_finder_apply():
    if "user" not in session:
        return jsonify({'error': 'Not logged in'}), 401
    data = request.get_json()
    title = data.get('title', 'Unknown Role')
    company = data.get('company', 'Unknown Company')
    location = data.get('location', '')
    url = data.get('url', '')
    conn = get_db()
    # avoid duplicate entries
    existing = conn.execute(
        "SELECT id FROM applications WHERE username=? AND job_title=? AND company=?",
        (session["user"], title, company)
    ).fetchone()
    if not existing:
        conn.execute(
    "INSERT INTO applications (username, job_title, company, status, applied_at) VALUES (?,?,?,?,?)",
    (session["user"], title, company, 'Applied', datetime.now().isoformat())
)
        conn.commit()
    conn.close()
    return jsonify({'success': True})
  
@app.route('/api/job-finder/extract-skills', methods=['POST'])
def extract_skills_from_resume():
    if 'resume' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400

    file = request.files['resume']
    ext = file.filename.rsplit('.', 1)[-1].lower()
    if ext not in {'pdf', 'docx', 'doc'}:
        return jsonify({'error': 'Only PDF or DOCX allowed'}), 400

    import tempfile
    with tempfile.NamedTemporaryFile(delete=False, suffix='.' + ext) as tmp:
        file.save(tmp.name)
        tmp_path = tmp.name

    try:
        raw_text = extract_text_from_file(tmp_path, file.filename)
    finally:
        os.unlink(tmp_path)

    if not raw_text.strip():
        return jsonify({'error': 'Could not extract text. Use DOCX for best results.'}), 400

    prompt = f"""
From this resume text, extract:
1. The best job title/role this person should search for
2. Top 5 skills

Return ONLY JSON like this (no markdown, no explanation):
{{"role": "Python Developer", "skills": ["Python", "Flask", "SQL", "Machine Learning", "Django"], "location": "India"}}

Resume:
{raw_text[:3000]}
"""
    try:
        raw = gemini(prompt)
        if not raw or not raw.strip():
            return jsonify({'error': 'AI returned empty response. Check your GROQ key.'}), 500
        if raw.startswith('AI error:'):
            return jsonify({'error': raw}), 500
        raw = re.sub(r'^```[a-z]*\n?', '', raw.strip())
        raw = re.sub(r'\n?```$', '', raw)
        # extract just the JSON object
        match = re.search(r'\{.*\}', raw, re.DOTALL)
        if not match:
            return jsonify({'error': f'No JSON found in AI response: {raw[:200]}'}), 500
        parsed = json.loads(match.group())
        return jsonify({'success': True, 'data': parsed})
    except Exception as e:
        return jsonify({'error': f'AI extraction failed: {str(e)}'}), 500
 
@app.route('/api/resume-builder/export-pdf', methods=['POST'])
def resume_export_pdf():
    if "user" not in session:
        return jsonify({'error': 'Not logged in'}), 401
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.enums import TA_CENTER, TA_LEFT

        data = request.get_json()
        name = data.get('name', 'Resume')
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
            leftMargin=20*mm, rightMargin=20*mm,
            topMargin=20*mm, bottomMargin=20*mm)

        blue = colors.HexColor('#2563eb')
        gray = colors.HexColor('#555555')

        styles = getSampleStyleSheet()
        s_name     = ParagraphStyle('n', fontSize=22, textColor=blue, alignment=TA_CENTER, spaceAfter=2)
        s_title    = ParagraphStyle('t', fontSize=12, textColor=gray, alignment=TA_CENTER, spaceAfter=2)
        s_contact  = ParagraphStyle('c', fontSize=10, textColor=gray, alignment=TA_CENTER, spaceAfter=8)
        s_heading  = ParagraphStyle('h', fontSize=11, textColor=blue, fontName='Helvetica-Bold', spaceBefore=10, spaceAfter=4, borderPadding=(0,0,2,0))
        s_body     = ParagraphStyle('b', fontSize=10, spaceAfter=2)
        s_bold     = ParagraphStyle('bb', fontSize=10, fontName='Helvetica-Bold', spaceAfter=1)
        s_sub      = ParagraphStyle('sub', fontSize=9, textColor=gray, spaceAfter=3)

        story = []

        story.append(Paragraph(data.get('name',''), s_name))
        if data.get('title'): story.append(Paragraph(data['title'], s_title))
        contact = ' | '.join(filter(None,[data.get('email',''), data.get('phone',''), data.get('location','')]))
        if contact: story.append(Paragraph(contact, s_contact))
        links = ' | '.join(filter(None,[data.get('linkedin',''), data.get('github','')]))
        if links: story.append(Paragraph(links, s_contact))

        if data.get('summary'):
            story.append(Paragraph('PROFILE', s_heading))
            story.append(Spacer(1, 1))
            story.append(Paragraph(data['summary'], s_body))

        if data.get('exp'):
            story.append(Paragraph('EXPERIENCE', s_heading))
            for e in data['exp']:
                story.append(Paragraph(f"{e.get('title','')} — {e.get('company','')}", s_bold))
                story.append(Paragraph(f"{e.get('start','')} — {e.get('end','')}", s_sub))
                if e.get('desc'): story.append(Paragraph(e['desc'], s_body))

        if data.get('edu'):
            story.append(Paragraph('EDUCATION', s_heading))
            for e in data['edu']:
                story.append(Paragraph(e.get('degree',''), s_bold))
                pct = (' | ' + e['percent']) if e.get('percent') else ''
                story.append(Paragraph(f"{e.get('college','')} | {e.get('from','')} — {e.get('to','')}{pct}", s_sub))

        if data.get('skill'):
            story.append(Paragraph('SKILLS', s_heading))
            story.append(Paragraph(', '.join([s.get('name','') for s in data['skill']]), s_body))

        if data.get('proj'):
            story.append(Paragraph('PROJECTS', s_heading))
            for p in data['proj']:
                story.append(Paragraph(p.get('title',''), s_bold))
                if p.get('desc'): story.append(Paragraph(p['desc'], s_body))

        if data.get('cert'):
            story.append(Paragraph('CERTIFICATIONS', s_heading))
            for c in data['cert']:
                story.append(Paragraph(f"{c.get('title','')} — {c.get('issuer','')} {c.get('date','')}", s_body))

        if data.get('lang'):
            story.append(Paragraph('LANGUAGES', s_heading))
            story.append(Paragraph(', '.join([l.get('name','') + (' (' + l.get('level','') + ')' if l.get('level') else '') for l in data['lang']]), s_body))

        if data.get('extra'):
            story.append(Paragraph('EXTRA CURRICULAR', s_heading))
            for x in data['extra']:
                story.append(Paragraph(f"{x.get('title','')} — {x.get('org','')}", s_bold))
                if x.get('desc'): story.append(Paragraph(x['desc'], s_body))

        doc.build(story)
        buffer.seek(0)
        return send_file(buffer, mimetype='application/pdf', as_attachment=True,
                         download_name=f"{name.replace(' ','_')}_resume.pdf")
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5001, use_reloader=False)